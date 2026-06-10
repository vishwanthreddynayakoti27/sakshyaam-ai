"""
FIXED-LAYOUT renderers reproducing the EXACT human-written templates from the
user's reference samples (Telangana Police, Makthal PS / Narayanpet District):

  1. Charge Sheet — Form-VII (18-row 2-col table, U/s 193 BNSS)
  2. Case Diary Part-I — header table + 8 numbered fields + narrative
  3. Remand Case Diary — Part-I letter format addressed to JMFC

DESIGN PRINCIPLE
================
The skeleton is HARD-CODED to match the real station-issued documents byte-
for-byte (layout, numbering, captions, prayer clauses, footers). NO LLM
re-arranges the structure. Missing field values render as `_____` so the
officer fills them by hand or in Word.

References used (extracted via extract_file_tool):
  * 57-26 Chargesheet.pdf  (FIR 57/2026 Makthal PS)
  * 236 remand.pdf         (FIR 236/2021 Makthal PS)
  * 57-26 CD 1.pdf         (Case Diary Part-I, FIR 57/2026)
"""
from __future__ import annotations

import io
import logging
import re
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

logger = logging.getLogger(__name__)

BLANK = "NOT FOUND IN DOCUMENTS"


# ============================================================
# Low-level helpers
# ============================================================
def _val(d: Optional[Dict[str, Any]], *keys: str, default: str = BLANK) -> str:
    """Walk nested dict; return BLANK when missing/empty."""
    cur: Any = d or {}
    for k in keys:
        if not isinstance(cur, dict):
            return default
        cur = cur.get(k)
        if cur is None or cur == "":
            return default
    if isinstance(cur, (list, dict)):
        return default
    s = str(cur).strip()
    return s if s else default


def _set_borders(cell, sz: str = "6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:color"), "000000")
        borders.append(b)
    tc_pr.append(borders)


def _shade_cell(cell, fill: str = "DDDDDD"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_run(run, *, bold: bool = False, italic: bool = False, size: int = 11,
             color: Optional[Tuple[int, int, int]] = None):
    run.font.name = "Times New Roman"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def _para(doc, text: str = "", *, bold: bool = False, italic: bool = False,
          size: int = 11, align: int = WD_ALIGN_PARAGRAPH.LEFT,
          space_before: int = 0, space_after: int = 0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        _set_run(run, bold=bold, italic=italic, size=size)
    return p


def _cell_para(cell, text: str, *, bold: bool = False, size: int = 11,
               align: int = WD_ALIGN_PARAGRAPH.LEFT, clear_first: bool = True):
    """Replace cell content with a single paragraph styled consistently."""
    if clear_first:
        # Replace existing default paragraph
        cell.text = ""
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text or "")
    _set_run(run, bold=bold, size=size)
    return p


def _make_table(doc, rows: int, cols: int, widths_cm: Optional[List[float]] = None):
    t = doc.add_table(rows=rows, cols=cols)
    t.style = "Table Grid"
    t.autofit = False
    if widths_cm:
        for i, w in enumerate(widths_cm):
            for cell in t.columns[i].cells:
                cell.width = Cm(w)
    for row in t.rows:
        for c in row.cells:
            _set_borders(c)
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    return t


def _set_margins(doc, top: float = 1.5, right: float = 1.5,
                 bottom: float = 1.5, left: float = 1.5):
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.right_margin = Cm(right)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)


# ============================================================
# Person / witness formatting helpers (matches station style)
# ============================================================
def _pick_relation(p: Dict[str, Any]) -> str:
    """Choose the correct parentage/spouse prefix from gender + marital hints.

    Returns exactly ONE of: 'W/o', 'S/o', 'D/o'.
    Never combines them (fixes the 'S/o W/o' bug).
    Priority:
      1. Explicit p["relation"] if it cleanly matches one of the three.
      2. Female + (married | salutation starts 'Smt.') → W/o
      3. Female + unmarried → D/o
      4. Male / default                              → S/o
    """
    explicit = (p.get("relation") or "").strip()
    # Strip combined garbage like "S/o W/o" → fallback to gender-based pick
    valid = {"w/o": "W/o", "s/o": "S/o", "d/o": "D/o"}
    low = explicit.lower()
    if low in valid:
        return valid[low]
    gender = (p.get("gender") or "").strip().lower()
    salutation = (p.get("salutation") or "").strip().lower()
    marital = (p.get("marital_status") or "").strip().lower()
    is_female = gender.startswith("f") or salutation.startswith(("smt", "kum"))
    is_married = salutation.startswith("smt") or marital == "married"
    if is_female:
        return "W/o" if is_married else "D/o"
    return "S/o"


def _format_person_block(p: Dict[str, Any]) -> str:
    """
    Compose a station-style person block:
      "Smt./Sri./Kum. <Name> <W/o|S/o|D/o> <Parent/Spouse>,
       Age: <N> years, Caste: <X>, Occ: <Y>, R/o <Address>, Ph.<phone>"
    Missing parts collapse to '_____' (or 'NOT FOUND IN DOCUMENTS' upstream).
    """
    if not isinstance(p, dict):
        return BLANK
    name = (p.get("name") or "").strip()
    # Gender → salutation default
    gender = (p.get("gender") or "").strip().lower()
    is_female = gender.startswith("f")
    marital = (p.get("marital_status") or "").strip().lower()
    if p.get("salutation"):
        salutation = p["salutation"]
    elif is_female and marital == "married":
        salutation = "Smt."
    elif is_female:
        salutation = "Kum."
    else:
        salutation = "Sri."
    relation = _pick_relation(p)
    father = (p.get("father") or p.get("guardian") or "").strip()
    age = (p.get("age") or "").strip()
    caste = (p.get("caste") or "").strip()
    occ = (p.get("occupation") or p.get("occ") or "").strip()
    addr = (p.get("address") or p.get("permanent_address") or "").strip()
    phone = (p.get("phone") or "").strip()

    def f(v): return v if v else BLANK
    pieces = [
        f"{salutation} {f(name)}",
        f"{relation} {f(father)}",
        f"Age: {f(age)} years" if age else f"Age: {BLANK} years",
        f"Caste: {f(caste)}",
        f"Occ: {f(occ)}",
        f"R/o {f(addr)}",
    ]
    if phone:
        pieces.append(f"Ph.{phone}")
    else:
        pieces.append(f"Ph.{BLANK}")
    aadhaar = p.get("aadhaar_number")
    if aadhaar:
        digits = re.sub(r"\D", "", str(aadhaar))
        if len(digits) == 12:
            aadhaar_disp = f"{digits[:4]} {digits[4:8]} {digits[8:]}"
        else:
            aadhaar_disp = aadhaar
        pieces.append(f"Aadhaar: {aadhaar_disp}")
    return ", ".join(pieces)


def _format_io_block(io: Dict[str, Any], ps: str) -> str:
    """Sri. <Name>, <Rank> of Police PS <PS> (IO & Filed Charge sheet)"""
    if not isinstance(io, dict):
        io = {}
    name = io.get("name") or BLANK
    rank = io.get("designation") or io.get("rank") or "SI"
    return f"Sri. {name}, {rank} of Police PS {ps or BLANK} (IO & Filed Charge sheet)"


# ============================================================
# Aadhaar auto-extraction (mechanical, no AI)
# ============================================================
_AADHAAR_NUM_RE = re.compile(r"\b\d{4}[\s-]?\d{4}[\s-]?\d{4}\b")
_DOB_RE = re.compile(r"\b(?:DOB|D\.O\.B|Date of Birth|Year of Birth|YOB)[:\s]*([\d/.\-]+)", re.I)
_GENDER_RE = re.compile(r"\b(Male|Female|MALE|FEMALE)\b")
_NAME_RE = re.compile(r"^[A-Z][a-zA-Z. ]{2,40}\s+[A-Z][a-zA-Z. ]{2,40}$")

# Keywords commonly found at the top of an Aadhaar card that look like a name
# but ARE NOT a person — used to filter false positives in name detection.
_AADHAAR_NAME_BLOCKLIST = {
    "Government Of India",
    "Government of India",
    "Unique Identification Authority Of India",
    "Unique Identification Authority of India",
    "Aadhaar Number",
    "Date Of Birth",
    "Date of Birth",
    "Address",
    "Father Name",
    "Mother Name",
    "MAIL DAK",
}


def _looks_like_aadhaar_header(line: str) -> bool:
    cleaned = re.sub(r"\s+", " ", line).strip()
    if cleaned in _AADHAAR_NAME_BLOCKLIST:
        return True
    low = cleaned.lower()
    return any(b.lower() in low for b in [
        "government", "unique identification", "authority", "aadhaar",
        "uidai", "issue date", "vid:", "enrolment",
    ])


def extract_aadhaar_from_files(files_meta: List[Dict[str, Any]]) -> Dict[str, str]:
    """Return a dict of aadhaar_* fields, BLANK strings when not detected."""
    out = {
        "aadhaar_number": "",
        "aadhaar_name": "",
        "aadhaar_dob": "",
        "aadhaar_gender": "",
        "aadhaar_address": "",
    }
    for f in (files_meta or []):
        text = (f.get("ocr_text") or f.get("text") or "")
        if not text and isinstance(f.get("extracted_data"), dict):
            text = f["extracted_data"].get("raw_text", "") or ""
        if not text:
            continue
        if ("Aadhaar" not in text and "AADHAAR" not in text and "Government of India" not in text
                and "UIDAI" not in text and not _AADHAAR_NUM_RE.search(text)):
            continue
        m = _AADHAAR_NUM_RE.search(text)
        if m and not out["aadhaar_number"]:
            out["aadhaar_number"] = re.sub(r"\s|-", "", m.group(0))
        m = _DOB_RE.search(text)
        if m and not out["aadhaar_dob"]:
            out["aadhaar_dob"] = m.group(1).strip()
        m = _GENDER_RE.search(text)
        if m and not out["aadhaar_gender"]:
            out["aadhaar_gender"] = m.group(1).title()
        for line in text.splitlines()[:30]:
            line = line.strip()
            if _looks_like_aadhaar_header(line):
                continue
            if _NAME_RE.match(line):
                if not out["aadhaar_name"]:
                    out["aadhaar_name"] = line
                    break
        for ln in text.splitlines():
            ls = ln.strip()
            if ls.startswith(("Address", "S/O", "D/O", "W/O", "C/O", "S/o", "D/o", "W/o", "C/o")) \
                    and not out["aadhaar_address"]:
                out["aadhaar_address"] = ls[:300]
                break
    return out


# ============================================================
# 1. CHARGE SHEET — Telangana Form-VII (18-section layout)
# ============================================================
# This matches reference samples 156.2025 CS.docx and 13.2025 CS.docx
# (Makthal PS) section-by-section. See render_charge_sheet() below for
# the verbatim section list.


def render_charge_sheet(case: Dict[str, Any]) -> bytes:
    """
    Reproduces the EXACT 18-section Telangana Police charge sheet layout
    from reference samples 156.2025 CS.docx and 13.2025 CS.docx (Makthal PS).

    Sections (HARD-CODED — never renamed, never skipped):
      01. Dist / PS / FIR No / Date (compound row)
      02. Charge Sheet No.
      03. Date of Charge
      04. Act and Section of Law
      05. Type of the final report
      06. If final report is un-occurred
      07. If charge sheet is original or supplementary.
      08. Name and rank of the I.O (s)
      09. Name and Address of the complainant or informant
      10. Details of property seized during the course of investigation.
      11. Particulars of accused persons charge sheeted (+ sub-rows a/b/c/d)
      12. Particulars of the accused persons not charge sheeted
      13. Particulars of witnesses to be examined  (+ witness sub-table)
      14. If F.R. is false, indicate action taken U/S 217/238 BNS
      15. Result of Laboratory Analysis
      16. Brief facts of the case  (+ narrative paragraphs)
      17. Is ack. copy of notice to complainant is enclosed
      18. Dispatched on
      → "Hence the charge sheet." → Signature block

    Missing data is rendered literally as `_____` so the officer fills
    the value by hand.
    """
    doc = Document()
    _set_margins(doc, top=1.2, right=1.5, bottom=1.5, left=1.5)

    # ── Title block (matches sample exactly: "C H A R G E – S H E E T") ──
    _para(doc, "C H A R G E – S H E E T", bold=True, size=16,
          align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, "(UNDER SECTION 193 BNSS.)", size=11,
          align=WD_ALIGN_PARAGRAPH.CENTER)
    court_name = _val(case, "court_name",
                      default="ADDL. JUDICIAL FIRST CLASS MAGISTRATE")
    court_place = _val(case, "court_place",
                       default=_val(case, "place", default="MAKTHAL"))
    _para(doc, f"IN THE COURT OF {court_name.upper()}", size=11,
          align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, f"AT {court_place.upper()}", size=11,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    # ── 5-column body table (Sno | Field | : | Value | Value-merge) ──
    body = _make_table(doc, 0, 5, widths_cm=[1.0, 6.5, 0.4, 4.0, 4.1])

    def _add_row(sno: str, label: str, value: str, sep: str = ":"):
        row = body.add_row()
        c = row.cells
        _cell_para(c[0], sno, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(c[1], label, size=10)
        _cell_para(c[2], sep, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        # merge cols 3+4 for the value cell
        merged = c[3].merge(c[4])
        _cell_para(merged, value, size=10)
        for cell in row.cells:
            _set_borders(cell)
        return row

    # 01 — combined Dist/PS/FIR row (matches sample: special compound layout)
    row01 = body.add_row()
    cells = row01.cells
    _cell_para(cells[0], "01", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_para(
        cells[1],
        f"Dist: {_val(case, 'district')}\nPS. : {_val(case, 'police_station')}",
        bold=True, size=10,
    )
    _cell_para(cells[2], ":", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    merged = cells[3].merge(cells[4])
    _cell_para(
        merged,
        f"FIR. No. {_val(case, 'fir_number')},   "
        f"Dtd: {_val(case, 'fir_date')}",
        bold=True, size=10,
    )
    for cell in row01.cells:
        _set_borders(cell)

    # 02 — Charge Sheet No.
    cs_no = _val(case, "charge_sheet_no", default="")
    if cs_no == BLANK or not cs_no:
        cs_no = f"{BLANK}/{datetime.now().year}"
    _add_row("02", "Charge Sheet No.", cs_no)

    # 03 — Date of Charge
    _add_row("03", "Date of Charge",
             _val(case, "today_date",
                  default=datetime.now().strftime("%d.%m.%Y")))

    # 04 — Act and Section of Law
    sections_val = _val(case, "sections")
    if sections_val and sections_val != BLANK and not sections_val.upper().startswith("U/S"):
        sections_val = f"U/S {sections_val}"
    _add_row("04", "Act and Section of Law", sections_val)

    # 05 — Type of final report
    _add_row("05", "Type of the final report",
             _val(case, "final_report_type", default="Charge Sheet."))

    # 06 — If final report is un-occurred
    _add_row("06", "If final report is un-occurred",
             _val(case, "fr_unoccurred", default="----"))

    # 07 — Original / supplementary
    _add_row("07", "If charge sheet is original or supplementary.",
             _val(case, "charge_sheet_kind", default="Original."))

    # 08 — IO name & rank (verbatim sample format)
    io_dict = case.get("io") or {}
    io_name = _val(io_dict, "name") if isinstance(io_dict, dict) else BLANK
    io_rank = _val(io_dict, "designation",
                   default=_val(io_dict, "rank", default="Sub inspector of Police"))
    io_value = (
        f"1. Sri. {io_name}, {io_rank}, PS {_val(case, 'police_station')}.\n"
        f"                                     (IO & filed Charge sheet)"
    )
    _add_row("08", "Name and rank of the I.O (s)", io_value)

    # 09 — Complainant
    _add_row("09", "Name and Address of the complainant or informant",
             _format_person_block(case.get("complainant") or {}))

    # 10 — Property seized
    _add_row("10", "Details of property seized during the course of investigation.",
             _val(case, "properties_seized", default="---"))

    # 11 — Particulars of accused persons charge sheeted
    accused = case.get("accused") or [{}]
    if not isinstance(accused, list) or not accused:
        accused = [{}]
    accused_block = "Particulars of accused persons charge sheeted ::-\n" + "\n\n".join(
        f"A{i}. {_format_person_block(a)}" for i, a in enumerate(accused, 1)
    )
    # 11 row uses a different visual: label cell holds the entire accused block
    row11 = body.add_row()
    _cell_para(row11.cells[0], "11", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    merged = row11.cells[1].merge(row11.cells[2]).merge(row11.cells[3]).merge(row11.cells[4])
    _cell_para(merged, accused_block, size=10)
    for cell in row11.cells:
        _set_borders(cell)

    # 11(a) Date of arrest, release / forwarded to court
    _add_row("", "a). Date of arrest, release\n      forwarded to court.",
             _val(case, "arrest_release", default="--"))
    # 11(b) Sureties
    _add_row("", "b) Particulars of sureties if\n      Released on bail.",
             _val(case, "sureties", default="--"))
    # 11(c) Previous convictions
    _add_row("", "c) Previous convictions if any.",
             _val(case, "previous_convictions", default="---"))
    # 11(d) Absconding
    _add_row("", "d) Particulars of accused\n      Persons absconding.",
             _val(case, "absconding", default="---"))

    # 12 — Accused not charge sheeted
    _add_row("12", "Particulars of the accused persons not charge sheeted",
             _val(case, "accused_not_chargesheeted", default="Nil"))

    # ── Section 13 — heading paragraph BEFORE the witness sub-table ──
    doc.add_paragraph("")
    _para(doc, "13. Particulars of witnesses to be examined: - Noted Below",
          size=11, space_before=4, space_after=4)

    # Witness sub-table (4 cols: LW-N | Name+Addr | : | Role)
    witnesses = case.get("witnesses") or [{}, {}, {}]
    if not isinstance(witnesses, list) or not witnesses:
        witnesses = [{}, {}, {}]
    wt = _make_table(doc, 0, 4, widths_cm=[1.7, 10.0, 0.4, 3.9])
    for i, w in enumerate(witnesses, 1):
        wrow = wt.add_row()
        _cell_para(wrow.cells[0], f"LW-{i}", bold=True, size=10,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(wrow.cells[1], _format_person_block(w), size=10)
        _cell_para(wrow.cells[2], ":", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(wrow.cells[3],
                   _val(w, "type", default=_val(w, "role", default=BLANK)),
                   size=10)
        for cell in wrow.cells:
            _set_borders(cell)

    # ── Sections 14, 15, 16 — second 3-row table (3 cols: Sno | Label | Value) ──
    doc.add_paragraph("")
    tbl_b = _make_table(doc, 0, 3, widths_cm=[1.0, 10.5, 4.5])
    for sno, label, val in [
        ("14", "If F.R. is false, indicate action taken or proposed to be taken  U/S 217/238 BNS",
         _val(case, "fr_false_action", default="--")),
        ("15", "Result of Laboratory Analysis",
         _val(case, "lab_result", default="--")),
        ("16", "Brief facts of the case", ""),  # value sits below in narrative
    ]:
        row = tbl_b.add_row()
        _cell_para(row.cells[0], sno, bold=True, size=10,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(row.cells[1], label, size=10)
        _cell_para(row.cells[2], val, size=10)
        for cell in row.cells:
            _set_borders(cell)

    # ── Section 16 narrative — the "Honoured Sir," letter ──
    doc.add_paragraph("")
    _para(doc, "Honoured Sir,", size=11, space_after=6)

    # Use brief_facts if provided as a single block, OR an array of paragraphs
    brief_paragraphs: List[str] = []
    if case.get("brief_facts_paragraphs") and isinstance(
            case["brief_facts_paragraphs"], list):
        brief_paragraphs = [str(p).strip() for p in case["brief_facts_paragraphs"] if p]
    elif case.get("brief_facts"):
        brief_paragraphs = [str(case["brief_facts"]).strip()]
    else:
        brief_paragraphs = [
            f"This is a case of {BLANK}. The incident occurred on {BLANK} at {BLANK} hours at {BLANK}. "
            f"The said location falls within the local limits of {_val(case, 'police_station')} Police Station.",
            f"The brief facts of the case are that on {BLANK} at {BLANK} hours complainant "
            f"{_format_person_block(case.get('complainant') or {})} came to "
            f"{_val(case, 'police_station')} PS and lodged a written petition.",
            f"Based on the above complaint, LW-{len(witnesses) or 'N'} registered a case in "
            f"Cr.No.{_val(case, 'fir_number')} {sections_val} and dispatched copies of the FIR to "
            f"the Hon'ble Court and all concerned offices as per procedure and took up the investigation.",
            f"During the course of investigation, LW-{len(witnesses) or 'N'} examined and recorded the "
            f"statement of LW-1 U/s 180(3) BNSS at the Police Station. The statement was duly documented "
            f"in Part-II CDs. Subsequently, LW-{len(witnesses) or 'N'} visited the scene of offence.",
        ]
        # Per-accused notice paragraphs
        for i, a in enumerate(accused, 1):
            brief_paragraphs.append(
                f"On {BLANK}, LW-{len(witnesses) or 'N'} served notice under Section 35(3) of the BNS, 2023, "
                f"to the accused A{i} (as listed in Column No. 11), informing of the allegations and directed "
                f"appearance for enquiry on {BLANK} along with ID proofs and documents."
            )
            brief_paragraphs.append(
                f"In compliance with the notice, on {BLANK}, the accused A{i} appeared before "
                f"LW-{len(witnesses) or 'N'} at PS {_val(case, 'police_station')} and voluntarily admitted guilt. "
                f"As such LW-{len(witnesses) or 'N'} obtained acknowledgement and instructed appearance before "
                f"the Hon'ble Court as required and released after collecting address proof."
            )
        brief_paragraphs.append(
            f"Based on the evidence collected during the course of investigation, it is clearly established "
            f"that the offence is committed by the accused person{'s' if len(accused) > 1 else ''} as cited."
        )
        n_acc = len(accused)
        last_label = f"A1 to A{n_acc}" if n_acc > 1 else "A1"
        brief_paragraphs.append(
            f"Therefore, the Hon'ble court is prayed that the accused person{'s' if n_acc > 1 else ''} "
            f"{last_label} mentioned in column No. 11 of this charge sheet may be tried and dealt suitably "
            f"as per law."
        )

    for p_text in brief_paragraphs:
        _para(doc, p_text, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              space_after=6)

    _para(doc, "Hence the charge sheet.", bold=True, size=11,
          space_before=4, space_after=8)

    # ── Sections 17, 18 ─────────────────────────────────────────────
    _para(doc,
          f"17. Is ack. copy of notice to complainant is enclosed: "
          f"{_val(case, 'ack_notice_enclosed', default='No.')}",
          size=11, space_after=2)
    _para(doc,
          f"18. Dispatched on: {_val(case, 'dispatch_date', default=BLANK)}.",
          size=11, space_after=20)

    # ── Signature block (matches sample exactly) ────────────────────
    _para(doc, "Signature of the Investigation officer", size=11,
          align=WD_ALIGN_PARAGRAPH.RIGHT)
    _para(doc, "Submitting chargesheet", size=11,
          align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=4)
    _para(doc, f"({_val(io_dict, 'name')})", size=11,
          align=WD_ALIGN_PARAGRAPH.RIGHT)
    _para(doc, f"{io_rank},", size=11, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _para(doc, f"PS {_val(case, 'police_station')}.", size=11,
          align=WD_ALIGN_PARAGRAPH.RIGHT)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ============================================================
# 2. CASE DIARY PART-I — Telangana CD format
# ============================================================
# Matches the user's sample 57-26 CD 1.pdf exactly:
#   header strip (PS / Dist / FIR / Date-place / CD Dt / Offence)
#   numbered table 1..8
#   narrative body (paragraphs)
#   "Closed the CD for the day; further progress follows."
#   "Copy submitted to the SDPO ..., through CI of Police ... f.f.i."
# ============================================================
def render_case_diary_part1(case: Dict[str, Any]) -> bytes:
    doc = Document()
    _set_margins(doc, top=1.2, right=1.5, bottom=1.5, left=1.5)

    _para(doc, "CASE DIARY", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, "(Part-I)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, "(Under Section 193(8) BNSS / Section 172 CrPC)",
          italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    # Header strip (single 2-col table with 6 rows)
    head = _make_table(doc, 0, 2, widths_cm=[5.5, 10.5])
    for label, key in [
        ("Police Station", "police_station"),
        ("Dist.", "district"),
        ("F.I.R No.", "fir_number"),
        ("Date, Time & Place of occurrence", "occurrence_dtp"),
        ("CD Dt.", "cd_date"),
        ("Offence U/s", "sections"),
    ]:
        row = head.add_row()
        _cell_para(row.cells[0], label, bold=True, size=11)
        _cell_para(row.cells[1], _val(case, key), size=11)
        _shade_cell(row.cells[0], "F2F2F2")
        for c in row.cells:
            _set_borders(c)

    doc.add_paragraph("")

    # Numbered fields (1..8)
    body = _make_table(doc, 0, 2, widths_cm=[7.5, 8.5])
    for num, label, key, default in [
        ("1", "Date and time of report", "report_datetime", BLANK),
        ("2", "Name of the Complainant / Informant", "_complainant_block", None),
        ("3", "Name and address of accused", "_accused_block", None),
        ("4", "Property Lost", "property_lost", "Nil"),
        ("5", "Property recovered", "property_recovered", "Nil"),
        ("6", "Date of Last Case Diary", "last_cd_date", "First CD"),
        ("7", "Name and address of deceased", "deceased", "Nil"),
        ("8", "Name and address of witnesses examined", "_witnesses_block", None),
    ]:
        row = body.add_row()
        _cell_para(row.cells[0], f"{num}.  {label}", bold=True, size=10)
        if key == "_complainant_block":
            value = _format_person_block(case.get("complainant") or {})
        elif key == "_accused_block":
            accused = case.get("accused") or [{}]
            if not isinstance(accused, list) or not accused:
                accused = [{}]
            value = "\n".join(f"A{i}. {_format_person_block(a)}" for i, a in enumerate(accused, 1))
        elif key == "_witnesses_block":
            wts = case.get("witnesses_examined") or case.get("witnesses") or []
            if not wts:
                wts = [{}, {}, {}]
            value = "\n".join(f"LW-{i}. {_format_person_block(w)}" for i, w in enumerate(wts, 1))
        else:
            value = _val(case, key, default=default or BLANK)
        _cell_para(row.cells[1], value, size=10)
        for c in row.cells:
            _set_borders(c)

    doc.add_paragraph("")

    # Narrative body (chronological investigation) — paragraphs
    _para(doc, "Brief Facts and Investigation Details:", bold=True, size=11, space_before=6)
    _para(doc, _val(case, "brief_facts",
                    default="(To be filled. Describe the complaint, registration of FIR, "
                            "examination of witnesses, scene visit, panchanama, sketch, "
                            "search & seizure, accused notice/appearance.)"),
          size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)

    # Step list
    _para(doc, "Steps Taken:", bold=True, size=11, space_before=6)
    for step_num, step in enumerate(case.get("investigation_steps") or [], 1):
        _para(doc, f"   {step_num}.  {step}", size=11,
              align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
    if not case.get("investigation_steps"):
        for tmpl in [
            "Visited the scene of offence and observed the surroundings.",
            "Conducted Scene of Offence Panchanama in the presence of mediators.",
            "Recorded statements of witnesses LW-1 to LW-N u/s 180 BNSS / 161 CrPC.",
            "Issued notice u/s 35(3) BNSS to the accused for appearance.",
            "Collected ID proofs and acknowledgement receipts from the accused.",
        ]:
            _para(doc, f"   •  {tmpl}", italic=True, size=10,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2,
                  )

    # Closing statement
    doc.add_paragraph("")
    _para(doc, "Closed the CD for the day; further progress follows.", italic=True,
          size=11, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=18)

    # Signature block
    _para(doc, "Signature of the Investigation Officer", bold=True, size=11,
          align=WD_ALIGN_PARAGRAPH.RIGHT)
    _para(doc, f"({_val(case, 'io', 'name')})", size=11, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _para(doc, f"{_val(case, 'io', 'designation')}, PS {_val(case, 'police_station')}",
          size=11, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=14)

    # Distribution footer
    _para(
        doc,
        f"Copy submitted to the SDPO {_val(case, 'district')}, through CI of Police "
        f"{_val(case, 'circle', default=_val(case, 'place'))} f.f.i.",
        italic=True, size=10,
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ============================================================
# 3. REMAND CASE DIARY (Part-I) — letter-format to Magistrate
# ============================================================
# Matches the user's sample 236 remand.pdf exactly:
#   Title: "REMAND CASE DIARY", "Part-I"
#   "IN THE COURT OF JUDICIAL MAGISTRATE OF FIRST CLASS AT ..."
#   "P.S.: <PS>   Dist: <DIST>"   "FIR No. ../....   Dated: ..."
#   "Honoured Sir,"
#   1..8 numbered fields (Investigating Officer / occurrence / offence /
#                         action / complainant / accused / property / witnesses)
#   Narrative paragraphs
#   "Reasons for arrest:" + para
#   "Hence the remand report."
#   Standard prayer clause (verbatim)
#   Signature, Encl, Escort
# ============================================================
def render_remand_report(case: Dict[str, Any]) -> bytes:
    doc = Document()
    _set_margins(doc, top=1.2, right=1.5, bottom=1.5, left=1.5)

    _para(doc, "REMAND CASE DIARY", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, "Part-I", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    court_place = _val(case, "court_place", default=_val(case, "district", default="NARAYANPET"))
    _para(doc, f"IN THE COURT OF JUDICIAL MAGISTRATE OF FIRST CLASS AT {court_place.upper()},",
          bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    # Top strip: PS / Dist / FIR No / Dated
    strip = _make_table(doc, 1, 2, widths_cm=[8.0, 8.0])
    _cell_para(strip.rows[0].cells[0],
               f"P.S.: {_val(case, 'police_station')}\nDist: {_val(case, 'district')}",
               bold=True, size=11)
    _cell_para(strip.rows[0].cells[1],
               f"FIR No. {_val(case, 'fir_number')}\nDated: {_val(case, 'fir_date')}",
               bold=True, size=11, align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.add_paragraph("")
    _para(doc, "Honoured Sir,", size=11, space_after=6)

    # Numbered fields 1..8 in a 2-col table (matches sample)
    body = _make_table(doc, 0, 2, widths_cm=[7.5, 8.5])

    def addrow(num: str, label: str, value: str, bold_label: bool = True):
        row = body.add_row()
        _cell_para(row.cells[0], f"{num}. {label}", bold=bold_label, size=10)
        _cell_para(row.cells[1], value, size=10)
        for c in row.cells:
            _set_borders(c)

    addrow("1", "Name of the Investigating Officer",
           _format_io_block(case.get("io") or {}, _val(case, "police_station", default="")))
    addrow("2", "Date and place of occurrence", _val(case, "occurrence_dtp"))
    addrow("3", "Offence U/s", _val(case, "sections"))
    addrow("4", "Date on which action was taken",
           _val(case, "action_taken_datetime", default=_val(case, "fir_date")))
    addrow("5", "Name of the complainant", _format_person_block(case.get("complainant") or {}))

    # Accused — multi-line A1..AN
    accused = case.get("accused") or [{}]
    if not isinstance(accused, list) or not accused:
        accused = [{}]
    acc_text = "\n".join(f"A{i}. {_format_person_block(a)}" for i, a in enumerate(accused, 1))
    addrow("6", "Name of the accused", acc_text)
    addrow("7", "Property lost", _val(case, "property_lost", default="Nil"))
    addrow("8", "Property recovered", _val(case, "property_recovered", default="Nil"))
    addrow("9", "Name of the deceased", _val(case, "deceased", default="Nil"))

    # Witnesses (renumbered list 1..N)
    witnesses = case.get("witnesses") or [{}, {}, {}]
    wt_text = "\n".join(f"{i}. {_format_person_block(w)}" for i, w in enumerate(witnesses, 1))
    addrow("10", "Name of the witnesses", wt_text)

    doc.add_paragraph("")

    # Narrative — Brief facts
    _para(doc, "Brief Facts of the Case:", bold=True, size=11, space_before=6)
    _para(doc, _val(case, "brief_facts",
                    default="(To be filled. Describe the date, time, place of occurrence, "
                            "the act of the accused, and how the case was registered.)"),
          size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)

    _para(doc, "Investigation Done So Far:", bold=True, size=11, space_before=6)
    _para(doc, _val(case, "investigation_done",
                    default="(To be filled. Witness statements u/s 180 BNSS, scene visit, "
                            "panchanama, recovery, arrest of accused.)"),
          size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)

    # Reasons for arrest
    _para(doc, "Reasons for arrest:", bold=True, size=11, space_before=6)
    _para(doc, _val(case, "grounds_of_arrest",
                    default="(To be filled. Grounds and necessity of arrest u/s 35 BNSS.)"),
          size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)

    _para(doc, "Hence the remand report.", bold=True, size=11, space_before=6, space_after=4)

    # Standard prayer (verbatim from sample)
    n_acc = len(accused)
    last = f"A{n_acc}" if n_acc > 1 else "A1"
    prayer = (
        f"The arrested accused person{'s' if n_acc > 1 else ''} "
        f"{'A1 to ' + last if n_acc > 1 else 'A1'} herewith produced before the Hon'ble Court "
        f"under proper escort with a pray to send {'them' if n_acc > 1 else 'him/her'} for "
        f"{_val(case, 'remand_type', default='judicial')} remand custody as the court deems fit."
    )
    _para(doc, prayer, italic=True, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=14)

    # Signature
    _para(doc, "", size=4)
    _para(doc, "Signature of the Investigating Officer", bold=True, size=11,
          align=WD_ALIGN_PARAGRAPH.RIGHT)
    _para(doc, f"({_val(case, 'io', 'name')})", size=11, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _para(doc, f"{_val(case, 'io', 'designation')}", size=11, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _para(doc, f"{_val(case, 'police_station')} P.S.", size=11,
          align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=14)

    # Encl & Escort
    _para(doc, "Encl:", bold=True, size=11)
    encl_items = case.get("enclosures") or [
        "Copy of FIR",
        "Statements of LW-1 to LW-N u/s 180 BNSS",
        "Scene of Offence Panchanama",
        "Wound certificate / MLC",
        "Notice u/s 35(3) BNSS to accused",
        "Aadhaar/ID proof of accused",
    ]
    for it in encl_items:
        _para(doc, f"   • {it}", size=10)
    _para(doc, "", size=4)
    _para(doc, f"Escort: {_val(case, 'escort', default=BLANK)}", size=11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ============================================================
# Top-level dispatcher
# ============================================================
RENDERERS = {
    "charge_sheet": render_charge_sheet,
    "case_diary_part1": render_case_diary_part1,
    "remand_report": render_remand_report,
}


def render_fixed_doc(doc_type: str, case_data: Dict[str, Any]) -> Tuple[bytes, str]:
    """Returns (docx_bytes, suggested_filename)."""
    if doc_type not in RENDERERS:
        raise ValueError(f"Unknown doc_type: {doc_type}. Must be one of {list(RENDERERS)}")
    bytes_ = RENDERERS[doc_type](case_data)
    safe_fir = re.sub(r"[^A-Za-z0-9_-]", "_", _val(case_data, "fir_number", default="UNKNOWN"))
    fname = f"{doc_type}_{safe_fir}_{datetime.now().strftime('%Y%m%d')}.docx"
    return bytes_, fname
