"""
Station-format Case Diary Part-I DOCX renderer.
Takes the structured JSON from `intelligent_case_diary.py` and produces
a print-ready DOCX matching the real station format.
"""
from __future__ import annotations

import io as _io
from typing import Any, Dict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

BLANK = "__________"


def _blank(val: Any, width: str = BLANK) -> str:
    if val is None:
        return width
    s = str(val).strip()
    return s if s else width


def render_case_diary_docx(data: Dict[str, Any]) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Header
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("C A S E   D I A R Y   —   P A R T  I")
    r.bold = True
    r.font.size = Pt(14)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("(UNDER SECTION 192 BNSS.)")
    r.font.size = Pt(10)

    # FIR line
    fir = doc.add_paragraph()
    r = fir.add_run(
        f"Dist.:- {_blank(data.get('district'))}    "
        f"Police Station: {_blank(data.get('police_station'))}    "
        f"FIR. No: {_blank(data.get('fir_number'))}    "
        f"Dated: {_blank(data.get('fir_date'))}"
    )
    r.bold = True
    r.font.size = Pt(10)

    meta = doc.add_paragraph()
    meta.add_run(
        f"U/s: {_blank(data.get('sections'))}    "
        f"Complainant: {_blank(data.get('complainant_name'))}    "
        f"Accused: {_blank(data.get('accused_list'))}"
    ).font.size = Pt(10)

    doc.add_paragraph()

    # Entries table
    entries = data.get("entries") or []
    if not entries:
        # empty template for manual entry
        entries = [{"date": "", "time": "", "entry": ""} for _ in range(4)]

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(1.1)
    table.columns[1].width = Inches(0.9)
    table.columns[2].width = Inches(5.2)

    hdr = table.rows[0].cells
    hdr[0].text = "Date"
    hdr[1].text = "Time"
    hdr[2].text = "Entry"
    for c in hdr:
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        shd = parse_xml(rf'<w:shd {nsdecls("w")} w:fill="BFBFBF"/>')
        c._tc.get_or_add_tcPr().append(shd)
        for p in c.paragraphs:
            for rn in p.runs:
                rn.bold = True

    for ent in entries:
        row = table.add_row().cells
        row[0].text = _blank(ent.get("date"), "__________")
        row[1].text = _blank(ent.get("time"), "________")
        body = (ent.get("entry") or "").strip()
        row[2].text = body if body else ("_" * 120)

    # Closing
    doc.add_paragraph()
    close = doc.add_paragraph()
    close.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = close.add_run(data.get("closing") or "")
    r.font.size = Pt(11)

    # Signature
    doc.add_paragraph()
    doc.add_paragraph()
    io = data.get("io") or {}
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = sig.add_run("Signature of the Investigating Officer\n")
    r.bold = True
    r.font.size = Pt(10)
    sig2 = doc.add_paragraph()
    sig2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = sig2.add_run(
        f"{_blank(io.get('name'), '__________________')},\n"
        f"{_blank(io.get('rank'), 'SI of Police')}, of {_blank(io.get('station'), 'PS __________')}"
    )
    r.font.size = Pt(10)

    out = _io.BytesIO()
    doc.save(out)
    return out.getvalue()
