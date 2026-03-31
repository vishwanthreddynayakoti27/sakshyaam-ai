"""
Word Document Generator for Charge Sheet, Case Diary Part-I, and Remand Case Diary
Generates exact Makthal PS format documents in .docx with stable tables
"""
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import io


def set_cell_border(cell, **kwargs):
    """Set cell border properties."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            element.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            element.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def add_borders_to_table(table):
    """Add borders to all cells in a table."""
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top={'val': 'single', 'sz': 4, 'color': '000000'},
                bottom={'val': 'single', 'sz': 4, 'color': '000000'},
                left={'val': 'single', 'sz': 4, 'color': '000000'},
                right={'val': 'single', 'sz': 4, 'color': '000000'})


def generate_chargesheet_docx(data: dict, case_info: dict) -> bytes:
    """
    Generate Charge Sheet in Word document format (.docx)
    Exact replication of Makthal PS 18-column format
    """
    doc = Document()
    
    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("C H A R G E – S H E E T")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("(UNDER SECTION 193 BNSS.)")
    subtitle_run.font.size = Pt(11)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    court = doc.add_paragraph()
    court_run = court.add_run(f"IN THE COURT OF ADDL. JUDICIAL FIRST CLASS MAGISTRATE AT {case_info.get('police_station', 'MAKTHAL').upper()}")
    court_run.bold = True
    court_run.font.size = Pt(11)
    court.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Create main table (18 rows)
    table = doc.add_table(rows=18, cols=3)
    
    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(1.2)
        row.cells[1].width = Cm(5)
        row.cells[2].width = Cm(10)
    
    add_borders_to_table(table)
    
    # Extract data
    complainant = data.get('complainant', {})
    accused_list = data.get('accused_persons', [])
    witnesses = data.get('witnesses', [])
    
    # Row data
    rows_data = [
        ("01", "Dist / PS / FIR", f"Dist.:- {case_info.get('district', 'Narayanpet')} Police Station: {case_info.get('police_station', 'Makthal')} FIR. No: {case_info.get('fir_number', '')} Dated: {case_info.get('fir_date', '')}"),
        ("02", "Final Report/Charge Sheet No.", "          /2026"),
        ("03", "Date", datetime.now().strftime("%d.%m.%Y")),
        ("04", "Act/Sections.", case_info.get('sections', '')),
        ("05", "Type of Final Report", "Charge sheet"),
        ("06", "F.R. Un occurred", "---"),
        ("07", "Original/supplementary", "Original"),
        ("08", "Names of I.O.", f"Sri. {case_info.get('io_name', '')}, {case_info.get('io_rank', 'SI of Police')} PS {case_info.get('police_station', '')}"),
        ("09", "Name of complainant", format_complainant(complainant)),
        ("10", "Property Recovered/Seized", data.get('property_recovered', '---')),
        ("11", "Particulars of accused", format_accused_list(accused_list, data.get('notice_date', ''))),
        ("12", "Sureties/Convictions/Absconding", "---"),
        ("13", "Accused not charge sheeted", "---"),
        ("14", "Witnesses to be examined", format_witnesses_list(witnesses)),
        ("15", "Result of Lab Analysis", "---"),
        ("16", "Brief facts of the case", data.get('brief_facts', '')),
        ("17", "Notice to complainant", "---"),
        ("18", "Dispatched on", datetime.now().strftime("%d.%m.%Y")),
    ]
    
    # Fill table
    for i, (num, label, value) in enumerate(rows_data):
        row = table.rows[i]
        row.cells[0].text = num
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[1].text = label
        row.cells[2].text = str(value)
        
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    
    # Prayer
    doc.add_paragraph()
    prayer = doc.add_paragraph()
    prayer.add_run("PRAYER: ").bold = True
    prayer.add_run("Therefore, the Hon'ble Court is prayed that the accused persons mentioned in column No. 11 may be tried and dealt with suitably as per law.")
    
    # Signature
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig.add_run(f"Signature of Investigation Officer\n{case_info.get('io_name', '')}\n{case_info.get('io_rank', 'SI of Police')}\nPS {case_info.get('police_station', '')}")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_case_diary_docx(data: dict, case_info: dict) -> bytes:
    """
    Generate Case Diary Part-I in Word format.
    Exact Makthal PS format with 8-point header.
    """
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.add_run("CASE DIARY Part - I")
    header_run.bold = True
    header_run.font.size = Pt(14)
    
    # Case Info Line
    info_para = doc.add_paragraph()
    info_para.add_run("Police Station: ").bold = True
    info_para.add_run(f"{case_info.get('police_station', 'MAKTHAL')}    ")
    info_para.add_run("Dist: ").bold = True
    info_para.add_run(f"{case_info.get('district', 'NARAYANPET')}.\n")
    info_para.add_run("F.I.R. No.: ").bold = True
    info_para.add_run(f"{case_info.get('fir_number', '')}    ")
    info_para.add_run("CD Dt: ").bold = True
    info_para.add_run(f"{datetime.now().strftime('%d.%m.%Y')}\n")
    info_para.add_run("Offence u/s ").bold = True
    info_para.add_run(f"{case_info.get('sections', '')}")
    
    doc.add_paragraph()
    
    # 8-Point Header Table
    complainant = data.get('complainant', {})
    accused_list = data.get('accused_persons', [])
    witnesses = data.get('witnesses', [])
    offense = data.get('offense_details', {})
    
    eight_points = [
        ("1.", "Date and time of report:", f"{case_info.get('fir_date', '')} at {offense.get('time', '')}"),
        ("2.", "Name of the Complainant/Informant:", format_complainant(complainant)),
        ("3.", "Name and address of accused:", format_accused_for_cd(accused_list)),
        ("4.", "Property Lost:", data.get('property_lost', '---')),
        ("5.", "Property recovered:", data.get('property_recovered', '---')),
        ("6.", "Date of Last Case Diary:", "First CD"),
        ("7.", "Name and address of deceased:", "---"),
        ("8.", "Name and address of witnesses examined:", format_witnesses_for_cd(witnesses)),
    ]
    
    table = doc.add_table(rows=8, cols=3)
    add_borders_to_table(table)
    
    for i, (num, label, value) in enumerate(eight_points):
        row = table.rows[i]
        row.cells[0].text = num
        row.cells[0].width = Cm(0.8)
        row.cells[1].text = label
        row.cells[1].width = Cm(5)
        row.cells[2].text = str(value)
        row.cells[2].width = Cm(10)
        
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Investigation Narrative
    narrative_header = doc.add_paragraph()
    narrative_header.add_run("INVESTIGATION NARRATIVE:").bold = True
    
    narrative = doc.add_paragraph()
    narrative.add_run(f"""On this day I resumed further investigation into this case.

{data.get('brief_facts', '')}

Closed the C.D. for the day.
Further progress follows through my next CD.""")
    
    # Signature
    doc.add_paragraph()
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig.add_run(f"({case_info.get('io_name', '')})\n{case_info.get('io_rank', 'Sub Inspector of Police')}, PS {case_info.get('police_station', '')}\n")
    sig.add_run(f"Copy submitted to the SDPO {case_info.get('district', '')}, Through CI of Police {case_info.get('police_station', '')} f.f.i.")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_remand_case_diary_docx(data: dict, case_info: dict) -> bytes:
    """
    Generate Remand Case Diary in Word format.
    Format based on 236 remand.pdf reference.
    """
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.add_run("REMAND CASE DIARY")
    header_run.bold = True
    header_run.font.size = Pt(14)
    
    subheader = doc.add_paragraph()
    subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subheader.add_run("Part-I")
    
    # Case Info
    info_para = doc.add_paragraph()
    info_para.add_run("Police Station: ").bold = True
    info_para.add_run(f"{case_info.get('police_station', 'MAKTHAL')}    ")
    info_para.add_run("Dist: ").bold = True
    info_para.add_run(f"{case_info.get('district', 'NARAYANPET')}.\n")
    info_para.add_run("Crime No.: ").bold = True
    info_para.add_run(f"{case_info.get('fir_number', '')}\n")
    info_para.add_run("U/S: ").bold = True
    info_para.add_run(f"{case_info.get('sections', '')}\n")
    info_para.add_run("Remand CD Dt: ").bold = True
    info_para.add_run(f"{datetime.now().strftime('%d.%m.%Y')}")
    
    doc.add_paragraph()
    
    complainant = data.get('complainant', {})
    accused_list = data.get('accused_persons', [])
    witnesses = data.get('witnesses', [])
    
    # Previous CD reference
    prev_cd = doc.add_paragraph()
    prev_cd.add_run("Previous case diary: ").bold = True
    prev_cd.add_run("This is the first Remand Case Diary.")
    
    # Deceased (if applicable)
    deceased = doc.add_paragraph()
    deceased.add_run("Name of the deceased: ").bold = True
    deceased.add_run("---")
    
    # Witnesses
    wit_para = doc.add_paragraph()
    wit_para.add_run("Name of the witnesses examined: ").bold = True
    wit_para.add_run(format_witnesses_for_cd(witnesses))
    
    doc.add_paragraph()
    
    # Address to Court
    court_address = doc.add_paragraph()
    court_address.add_run("Honoured Sir,").bold = True
    
    doc.add_paragraph()
    
    # Brief Facts
    facts = doc.add_paragraph()
    facts.add_run("Brief facts of the case:\n\n").bold = True
    facts.add_run(data.get('brief_facts', ''))
    
    doc.add_paragraph()
    
    # REASONS FOR ARREST
    reasons_header = doc.add_paragraph()
    reasons_header.add_run("REASONS FOR ARREST:").bold = True
    reasons_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    reasons = doc.add_paragraph()
    arrest_details = data.get('arrest_details', {})
    grounds = f"""
1. The accused has committed a cognizable offence punishable under sections {case_info.get('sections', '')} of BNS.

2. There is a reasonable suspicion that the accused has committed the said offence based on the evidence collected during investigation.

3. The arrest is necessary to prevent the accused from:
   a) Committing any further offence
   b) Tampering with evidence
   c) Influencing witnesses
   d) Absconding

4. The accused was served notice U/s 35(3) BNSS on {data.get('notice_date', '')} and appeared before the investigating officer.

5. The investigation is still ongoing and the accused's custody is required for further investigation.
"""
    reasons.add_run(grounds)
    
    doc.add_paragraph()
    
    # Prayer
    prayer_header = doc.add_paragraph()
    prayer_header.add_run("HENCE THE REMAND REPORT:").bold = True
    prayer_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    prayer = doc.add_paragraph()
    prayer.add_run(f"""
It is therefore prayed that the Hon'ble Court may kindly remand the accused:

{format_accused_for_cd(accused_list)}

to judicial custody for a period of 15 days to enable the investigating officer to complete the investigation.
""")
    
    doc.add_paragraph()
    
    # Enclosures
    encl = doc.add_paragraph()
    encl.add_run("Encl:").bold = True
    encl.add_run("\n1. Remand application\n2. Case diary copies\n3. Section 35(3) BNSS notice copies")
    
    # Signature
    doc.add_paragraph()
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig.add_run(f"({case_info.get('io_name', '')})\n{case_info.get('io_rank', 'Sub Inspector of Police')}\nPS {case_info.get('police_station', '')}")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# === Helper Functions ===

def format_complainant(comp: dict) -> str:
    """Format complainant details."""
    if not comp:
        return "[ ]"
    name = comp.get('name', '[ ]')
    father = comp.get('father_name', '[ ]')
    age = comp.get('age', '[ ]')
    caste = comp.get('caste', '[ ]')
    occupation = comp.get('occupation', '[ ]')
    address = comp.get('address', '[ ]')
    phone = comp.get('phone', '')
    
    result = f"Sri. {name} S/o {father}, Age: {age} years, Caste: {caste}, Occ: {occupation}, R/o {address}"
    if phone:
        result += f", Ph.{phone}"
    return result


def format_accused_list(accused: list, notice_date: str = "") -> str:
    """Format accused persons list for charge sheet."""
    if not accused:
        return "[ ACCUSED DETAILS ]"
    
    lines = []
    for i, acc in enumerate(accused):
        serial = acc.get('serial', f'A{i+1}')
        name = acc.get('name', '[ ]')
        father = acc.get('father_name', '[ ]')
        age = acc.get('age', '[ ]')
        caste = acc.get('caste', '[ ]')
        occupation = acc.get('occupation', '[ ]')
        address = acc.get('address', '[ ]')
        phone = acc.get('phone', '')
        
        line = f"{serial}. {name} S/o {father}, Age: {age} years, Caste: {caste}, Occ: {occupation}, R/o {address}"
        if phone:
            line += f" Ph. {phone}"
        lines.append(line)
    
    if notice_date:
        lines.append(f"\na) Served notice U/s 35(3) BNSS on {notice_date}")
    
    return "\n".join(lines)


def format_accused_for_cd(accused: list) -> str:
    """Format accused persons for Case Diary."""
    if not accused:
        return "[ ]"
    
    lines = []
    for i, acc in enumerate(accused):
        serial = acc.get('serial', f'A{i+1}')
        name = acc.get('name', '[ ]')
        father = acc.get('father_name', '[ ]')
        address = acc.get('address', '[ ]')
        lines.append(f"{serial}. {name} S/o {father}, R/o {address}")
    
    return "\n".join(lines)


def format_witnesses_list(witnesses: list) -> str:
    """Format witnesses list for charge sheet."""
    if not witnesses:
        return "[ WITNESS DETAILS ]"
    
    lines = []
    for i, wit in enumerate(witnesses):
        serial = wit.get('serial', f'LW-{i+1}')
        name = wit.get('name', '[ ]')
        father = wit.get('father_name', '[ ]')
        age = wit.get('age', '[ ]')
        caste = wit.get('caste', '[ ]')
        occupation = wit.get('occupation', '[ ]')
        address = wit.get('address', '[ ]')
        phone = wit.get('phone', '')
        role = wit.get('role', '')
        
        line = f"{serial}. Sri. {name} S/o {father}, Age: {age} Yrs., Caste: {caste}, Occ: {occupation}, R/o {address}"
        if phone:
            line += f" - {phone}"
        if role:
            line += f" ({role})"
        lines.append(line)
    
    return "\n\n".join(lines)


def format_witnesses_for_cd(witnesses: list) -> str:
    """Format witnesses for Case Diary."""
    if not witnesses:
        return "---"
    
    lines = []
    for i, wit in enumerate(witnesses):
        serial = wit.get('serial', f'LW-{i+1}')
        name = wit.get('name', '[ ]')
        lines.append(f"{serial}. {name}")
    
    return ", ".join(lines)
