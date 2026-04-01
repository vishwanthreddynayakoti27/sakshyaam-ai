"""
Template Service - DOCX Template-Based Document Generation
============================================================
Uses docxtpl (Jinja2 templating for Word) to generate documents from templates.

Templates:
  - chargesheet_template.docx: 18-column charge sheet format
  - casediary_template.docx: Case Diary Part-I format
  - remand_template.docx: Remand Case Diary format

Template Tags (Jinja2 syntax):
  - {{ fir_number }}
  - {{ complainant_name }}
  - {% for accused in accused_list %}...{% endfor %}
  - {{ io_name }}

NO layout generation via code - templates are pre-formatted.
"""
import os
import io
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional
from docxtpl import DocxTemplate
from .aggregator_service import UnifiedSchema

logger = logging.getLogger(__name__)


class TemplateService:
    """
    Template-based DOCX generation using docxtpl.
    NO programmatic layout - uses pre-formatted templates.
    """
    
    TEMPLATES_DIR = Path("/app/backend/templates")
    
    # Template filenames
    CHARGESHEET_TEMPLATE = "chargesheet_template.docx"
    CASEDIARY_TEMPLATE = "casediary_template.docx"
    REMAND_TEMPLATE = "remand_template.docx"
    
    def __init__(self, templates_dir: Optional[Path] = None):
        """
        Initialize template service.
        
        Args:
            templates_dir: Path to templates directory (optional)
        """
        self.templates_dir = templates_dir or self.TEMPLATES_DIR
        self.templates_dir.mkdir(parents=True, exist_ok=True)
        
        # Ensure templates exist (create placeholders if not)
        self._ensure_templates()
    
    def _ensure_templates(self):
        """Ensure all required templates exist."""
        for template_name in [self.CHARGESHEET_TEMPLATE, self.CASEDIARY_TEMPLATE, self.REMAND_TEMPLATE]:
            template_path = self.templates_dir / template_name
            if not template_path.exists():
                logger.warning(f"Template {template_name} not found. Creating placeholder.")
                self._create_placeholder_template(template_path, template_name)
    
    def _create_placeholder_template(self, path: Path, template_type: str):
        """Create a basic placeholder template if not exists."""
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
        
        if "chargesheet" in template_type.lower():
            self._create_chargesheet_template(doc)
        elif "casediary" in template_type.lower():
            self._create_casediary_template(doc)
        elif "remand" in template_type.lower():
            self._create_remand_template(doc)
        
        doc.save(str(path))
        logger.info(f"Created placeholder template: {path}")
    
    def _create_chargesheet_template(self, doc):
        """Create charge sheet template with Jinja2 tags."""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        # Title
        title = doc.add_paragraph()
        title_run = title.add_run("C H A R G E – S H E E T")
        title_run.bold = True
        title_run.font.size = Pt(14)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph()
        subtitle.add_run("(UNDER SECTION 193 BNSS.)")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        court = doc.add_paragraph()
        court_run = court.add_run("IN THE COURT OF ADDL. JUDICIAL FIRST CLASS MAGISTRATE AT {{ police_station|upper }}")
        court_run.bold = True
        court.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Create 18-row table with Jinja2 tags
        table = doc.add_table(rows=18, cols=3)
        self._add_borders_to_table(table)
        
        rows_data = [
            ("01", "Dist / PS / FIR", "Dist.:- {{ district }} Police Station: {{ police_station }} FIR. No: {{ fir_number }} Dated: {{ fir_date }}"),
            ("02", "Final Report/Charge Sheet No.", "{{ chargesheet_number }}"),
            ("03", "Date", "{{ current_date }}"),
            ("04", "Act/Sections.", "{{ sections }}"),
            ("05", "Type of Final Report", "Charge sheet"),
            ("06", "F.R. Un occurred", "---"),
            ("07", "Original/supplementary", "Original"),
            ("08", "Names of I.O.", "Sri. {{ io_name }}, {{ io_rank }} PS {{ police_station }}"),
            ("09", "Name of complainant", "{{ complainant_formatted }}"),
            ("10", "Property Recovered/Seized", "{{ property_recovered }}"),
            ("11", "Particulars of accused", "{{ accused_formatted }}"),
            ("12", "Sureties/Convictions/Absconding", "---"),
            ("13", "Accused not charge sheeted", "---"),
            ("14", "Witnesses to be examined", "{{ witnesses_formatted }}"),
            ("15", "Result of Lab Analysis", "{{ fsl_result }}"),
            ("16", "Brief facts of the case", "{{ brief_facts }}"),
            ("17", "Notice to complainant", "---"),
            ("18", "Dispatched on", "{{ current_date }}"),
        ]
        
        for i, (num, label, value) in enumerate(rows_data):
            row = table.rows[i]
            row.cells[0].text = num
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[1].text = label
            row.cells[2].text = value
        
        # Prayer
        doc.add_paragraph()
        prayer = doc.add_paragraph()
        prayer.add_run("PRAYER: ").bold = True
        prayer.add_run("Therefore, the Hon'ble Court is prayed that the accused persons mentioned in column No. 11 may be tried and dealt with suitably as per law.")
        
        # Signature
        doc.add_paragraph()
        doc.add_paragraph()
        sig = doc.add_paragraph()
        sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sig.add_run("Signature of Investigation Officer\n{{ io_name }}\n{{ io_rank }}\nPS {{ police_station }}")
    
    def _create_casediary_template(self, doc):
        """Create case diary template with Jinja2 tags."""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Header
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header.add_run("CASE DIARY Part - I")
        header_run.bold = True
        header_run.font.size = Pt(14)
        
        # Case Info
        info = doc.add_paragraph()
        info.add_run("Police Station: ").bold = True
        info.add_run("{{ police_station }}    ")
        info.add_run("Dist: ").bold = True
        info.add_run("{{ district }}\n")
        info.add_run("F.I.R. No.: ").bold = True
        info.add_run("{{ fir_number }}    ")
        info.add_run("CD Dt: ").bold = True
        info.add_run("{{ current_date }}\n")
        info.add_run("Offence u/s ").bold = True
        info.add_run("{{ sections }}")
        
        doc.add_paragraph()
        
        # 8-point table
        table = doc.add_table(rows=8, cols=3)
        self._add_borders_to_table(table)
        
        points = [
            ("1.", "Date and time of report:", "{{ fir_date }} at {{ incident_time }}"),
            ("2.", "Name of the Complainant/Informant:", "{{ complainant_formatted }}"),
            ("3.", "Name and address of accused:", "{{ accused_cd_formatted }}"),
            ("4.", "Property Lost:", "{{ property_lost }}"),
            ("5.", "Property recovered:", "{{ property_recovered }}"),
            ("6.", "Date of Last Case Diary:", "First CD"),
            ("7.", "Name and address of deceased:", "---"),
            ("8.", "Name and address of witnesses examined:", "{{ witnesses_cd_formatted }}"),
        ]
        
        for i, (num, label, value) in enumerate(points):
            row = table.rows[i]
            row.cells[0].text = num
            row.cells[1].text = label
            row.cells[2].text = value
        
        doc.add_paragraph()
        
        # Investigation Narrative
        narrative = doc.add_paragraph()
        narrative.add_run("INVESTIGATION NARRATIVE:").bold = True
        
        doc.add_paragraph("On this day I resumed further investigation into this case.\n\n{{ brief_facts }}\n\nClosed the C.D. for the day.\nFurther progress follows through my next CD.")
        
        # Signature
        doc.add_paragraph()
        sig = doc.add_paragraph()
        sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sig.add_run("({{ io_name }})\n{{ io_rank }}, PS {{ police_station }}\nCopy submitted to the SDPO {{ district }}, Through CI of Police {{ police_station }} f.f.i.")
    
    def _create_remand_template(self, doc):
        """Create remand case diary template with Jinja2 tags."""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Header
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header.add_run("REMAND CASE DIARY")
        header_run.bold = True
        header_run.font.size = Pt(14)
        
        subheader = doc.add_paragraph()
        subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subheader.add_run("Part-I")
        
        # Case Info
        info = doc.add_paragraph()
        info.add_run("Police Station: ").bold = True
        info.add_run("{{ police_station }}    ")
        info.add_run("Dist: ").bold = True
        info.add_run("{{ district }}\n")
        info.add_run("Crime No.: ").bold = True
        info.add_run("{{ fir_number }}\n")
        info.add_run("U/S: ").bold = True
        info.add_run("{{ sections }}\n")
        info.add_run("Remand CD Dt: ").bold = True
        info.add_run("{{ current_date }}")
        
        doc.add_paragraph()
        
        prev = doc.add_paragraph()
        prev.add_run("Previous case diary: ").bold = True
        prev.add_run("This is the first Remand Case Diary.")
        
        deceased = doc.add_paragraph()
        deceased.add_run("Name of the deceased: ").bold = True
        deceased.add_run("---")
        
        wit = doc.add_paragraph()
        wit.add_run("Name of the witnesses examined: ").bold = True
        wit.add_run("{{ witnesses_cd_formatted }}")
        
        doc.add_paragraph()
        
        court = doc.add_paragraph()
        court.add_run("Honoured Sir,").bold = True
        
        doc.add_paragraph()
        
        facts = doc.add_paragraph()
        facts.add_run("Brief facts of the case:\n\n").bold = True
        facts.add_run("{{ brief_facts }}")
        
        doc.add_paragraph()
        
        # Reasons for arrest
        reasons_header = doc.add_paragraph()
        reasons_header.add_run("REASONS FOR ARREST:").bold = True
        reasons_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        reasons = doc.add_paragraph()
        reasons.add_run("""{{ remand_narrative }}""")
        
        doc.add_paragraph()
        
        # Prayer
        prayer_header = doc.add_paragraph()
        prayer_header.add_run("HENCE THE REMAND REPORT:").bold = True
        prayer_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        prayer = doc.add_paragraph()
        prayer.add_run("It is therefore prayed that the Hon'ble Court may kindly remand the accused:\n\n{{ accused_cd_formatted }}\n\nto judicial custody for a period of 15 days to enable the investigating officer to complete the investigation.")
        
        doc.add_paragraph()
        
        encl = doc.add_paragraph()
        encl.add_run("Encl:").bold = True
        encl.add_run("\n1. Remand application\n2. Case diary copies\n3. Section 35(3) BNSS notice copies")
        
        # Signature
        doc.add_paragraph()
        sig = doc.add_paragraph()
        sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sig.add_run("({{ io_name }})\n{{ io_rank }}\nPS {{ police_station }}")
    
    def _add_borders_to_table(self, table):
        """Add borders to all cells in table."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for edge in ['top', 'left', 'bottom', 'right']:
                    element = OxmlElement(f'w:{edge}')
                    element.set(qn('w:val'), 'single')
                    element.set(qn('w:sz'), '4')
                    element.set(qn('w:color'), '000000')
                    tcBorders.append(element)
                tcPr.append(tcBorders)
    
    def prepare_context(self, schema: UnifiedSchema, 
                       case_info: Optional[Dict[str, str]] = None) -> Dict[str, Any]:
        """
        Prepare template context from unified schema.
        
        Args:
            schema: UnifiedSchema with all extracted data
            case_info: Optional overrides for case info
            
        Returns:
            Dict with all template variables
        """
        case_info = case_info or {}
        
        # Format complainant
        comp = schema.complainant
        complainant_formatted = self._format_person(comp) if comp.name else "[ ]"
        
        # Format accused list
        accused_formatted = self._format_accused_list(schema.accused)
        accused_cd_formatted = self._format_accused_cd(schema.accused)
        
        # Format witnesses
        witnesses_formatted = self._format_witnesses_list(schema.witnesses)
        witnesses_cd_formatted = self._format_witnesses_cd(schema.witnesses)
        
        # Get notice date from section 35(3) dates
        notice_date = schema.notices.section_35_3_dates[0] if schema.notices.section_35_3_dates else ""
        
        context = {
            # FIR details
            "fir_number": schema.fir.number or case_info.get('fir_number', ''),
            "fir_date": schema.fir.date or case_info.get('fir_date', ''),
            "police_station": schema.fir.police_station or case_info.get('police_station', ''),
            "district": schema.fir.district or case_info.get('district', ''),
            "sections": ", ".join(schema.fir.sections) or case_info.get('sections', ''),
            
            # Chargesheet number (generate if not provided)
            "chargesheet_number": case_info.get('chargesheet_number', f"          /{datetime.now().year}"),
            
            # Current date
            "current_date": datetime.now().strftime("%d.%m.%Y"),
            
            # IO details
            "io_name": schema.io_details.get('name', case_info.get('io_name', '')),
            "io_rank": schema.io_details.get('rank', case_info.get('io_rank', 'Sub Inspector of Police')),
            
            # Complainant
            "complainant_formatted": complainant_formatted,
            "complainant_name": comp.name if comp else '',
            
            # Accused
            "accused_formatted": accused_formatted,
            "accused_cd_formatted": accused_cd_formatted,
            "accused_list": [self._person_to_dict(a) for a in schema.accused],
            "accused_count": len(schema.accused),
            
            # Witnesses
            "witnesses_formatted": witnesses_formatted,
            "witnesses_cd_formatted": witnesses_cd_formatted,
            "witness_list": [self._person_to_dict(w) for w in schema.witnesses],
            "witness_count": len(schema.witnesses),
            
            # Incident
            "incident_date": schema.incident.date,
            "incident_time": schema.incident.time,
            "incident_place": schema.incident.place,
            
            # Property
            "property_lost": schema.property.lost or "---",
            "property_recovered": schema.property.recovered or "---",
            
            # Facts (AI-generated or raw)
            "brief_facts": schema.facts.ai_generated or schema.facts.raw[:3000] or "[ Brief facts to be generated ]",
            "remand_narrative": schema.facts.remand_narrative or self._default_remand_narrative(schema),
            
            # FSL
            "fsl_result": "---",
            
            # Notice date
            "notice_date": notice_date,
            
            # Medical
            "medical_findings": schema.medical.findings or "---",
        }
        
        return context
    
    def _format_person(self, person) -> str:
        """Format person details string."""
        if not person or not person.name:
            return "[ ]"
        
        parts = [f"Sri. {person.name}"]
        
        if person.father_name:
            parts.append(f"S/o {person.father_name}")
        
        if person.age:
            parts.append(f"Age: {person.age} years")
        
        if person.caste:
            parts.append(f"Caste: {person.caste}")
        
        if person.occupation:
            parts.append(f"Occ: {person.occupation}")
        
        if person.address:
            parts.append(f"R/o {person.address}")
        
        if person.phone:
            parts.append(f"Ph. {person.phone}")
        
        return ", ".join(parts)
    
    def _format_accused_list(self, accused_list) -> str:
        """Format accused list for charge sheet."""
        if not accused_list:
            return "[ ACCUSED DETAILS ]"
        
        lines = []
        for acc in accused_list:
            line = f"{acc.serial}. {self._format_person(acc)}"
            lines.append(line)
        
        return "\n".join(lines)
    
    def _format_accused_cd(self, accused_list) -> str:
        """Format accused for case diary (shorter format)."""
        if not accused_list:
            return "[ ]"
        
        lines = []
        for acc in accused_list:
            parts = [f"{acc.serial}. {acc.name}"]
            if acc.father_name:
                parts.append(f"S/o {acc.father_name}")
            if acc.address:
                parts.append(f"R/o {acc.address}")
            lines.append(", ".join(parts))
        
        return "\n".join(lines)
    
    def _format_witnesses_list(self, witnesses) -> str:
        """Format witnesses list for charge sheet."""
        if not witnesses:
            return "[ WITNESS DETAILS ]"
        
        lines = []
        for wit in witnesses:
            parts = [f"{wit.serial}. Sri. {wit.name}"]
            if wit.father_name:
                parts.append(f"S/o {wit.father_name}")
            if wit.age:
                parts.append(f"Age: {wit.age} Yrs.")
            if wit.caste:
                parts.append(f"Caste: {wit.caste}")
            if wit.occupation:
                parts.append(f"Occ: {wit.occupation}")
            if wit.address:
                parts.append(f"R/o {wit.address}")
            if wit.phone:
                parts.append(f"- {wit.phone}")
            if wit.role:
                parts.append(f"({wit.role})")
            
            lines.append(", ".join(parts))
        
        return "\n\n".join(lines)
    
    def _format_witnesses_cd(self, witnesses) -> str:
        """Format witnesses for case diary (short format)."""
        if not witnesses:
            return "---"
        
        return ", ".join([f"{w.serial}. {w.name}" for w in witnesses])
    
    def _person_to_dict(self, person) -> Dict[str, Any]:
        """Convert PersonRecord to dict for template loops."""
        return {
            "serial": person.serial,
            "name": person.name,
            "father_name": person.father_name,
            "age": person.age,
            "caste": person.caste,
            "occupation": person.occupation,
            "address": person.address,
            "phone": person.phone,
            "role": person.role
        }
    
    def _default_remand_narrative(self, schema: UnifiedSchema) -> str:
        """Generate default remand narrative if AI not available."""
        sections = ", ".join(schema.fir.sections) if schema.fir.sections else "the relevant sections"
        
        return f"""
1. The accused has committed a cognizable offence punishable under sections {sections} of BNS.

2. There is a reasonable suspicion that the accused has committed the said offence based on the evidence collected during investigation.

3. The arrest is necessary to prevent the accused from:
   a) Committing any further offence
   b) Tampering with evidence
   c) Influencing witnesses
   d) Absconding

4. The investigation is still ongoing and the accused's custody is required for further investigation.
"""
    
    def generate_chargesheet(self, schema: UnifiedSchema, 
                            case_info: Optional[Dict[str, str]] = None) -> bytes:
        """
        Generate Charge Sheet DOCX from template.
        
        Args:
            schema: UnifiedSchema with extracted data
            case_info: Optional case info overrides
            
        Returns:
            DOCX file bytes
        """
        template_path = self.templates_dir / self.CHARGESHEET_TEMPLATE
        
        doc = DocxTemplate(str(template_path))
        context = self.prepare_context(schema, case_info)
        
        doc.render(context)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        logger.info(f"Generated Charge Sheet DOCX ({len(schema.accused)} accused, {len(schema.witnesses)} witnesses)")
        return buffer.getvalue()
    
    def generate_casediary(self, schema: UnifiedSchema,
                          case_info: Optional[Dict[str, str]] = None) -> bytes:
        """
        Generate Case Diary Part-I DOCX from template.
        
        Args:
            schema: UnifiedSchema with extracted data
            case_info: Optional case info overrides
            
        Returns:
            DOCX file bytes
        """
        template_path = self.templates_dir / self.CASEDIARY_TEMPLATE
        
        doc = DocxTemplate(str(template_path))
        context = self.prepare_context(schema, case_info)
        
        doc.render(context)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        logger.info("Generated Case Diary DOCX")
        return buffer.getvalue()
    
    def generate_remand(self, schema: UnifiedSchema,
                       case_info: Optional[Dict[str, str]] = None) -> bytes:
        """
        Generate Remand Case Diary DOCX from template.
        
        Args:
            schema: UnifiedSchema with extracted data
            case_info: Optional case info overrides
            
        Returns:
            DOCX file bytes
        """
        template_path = self.templates_dir / self.REMAND_TEMPLATE
        
        doc = DocxTemplate(str(template_path))
        context = self.prepare_context(schema, case_info)
        
        doc.render(context)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        logger.info("Generated Remand Case Diary DOCX")
        return buffer.getvalue()
    
    def generate_all(self, schema: UnifiedSchema,
                    case_info: Optional[Dict[str, str]] = None) -> Dict[str, bytes]:
        """
        Generate all three documents.
        
        Returns:
            Dict with 'chargesheet', 'casediary', 'remand' keys containing DOCX bytes
        """
        return {
            'chargesheet': self.generate_chargesheet(schema, case_info),
            'casediary': self.generate_casediary(schema, case_info),
            'remand': self.generate_remand(schema, case_info)
        }
