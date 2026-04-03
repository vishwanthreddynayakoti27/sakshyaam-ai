"""
OCR Service - Unified OCR Pipeline
====================================
Primary: Azure AI Document Intelligence (high accuracy for tables)
Fallback: Google Vision API (Telugu/English support)
Emergency: Tesseract OCR (local, no API)

Extracts text from: PDF, DOCX, DOC, JPG, PNG, JPEG, WEBP, GIF
"""
import io
import os
import tempfile
import subprocess
import logging
from pathlib import Path
from typing import Optional, List, Tuple
from dataclasses import dataclass
from dotenv import load_dotenv

logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()


@dataclass
class OCRResult:
    """Result of OCR/text extraction."""
    text: str
    source_file: str
    file_type: str
    success: bool
    error: Optional[str] = None
    char_count: int = 0
    language_detected: Optional[str] = None
    ocr_engine: str = "unknown"
    confidence: float = 0.0
    tables_extracted: int = 0


class OCRService:
    """
    Multi-format text extraction service.
    
    Engine Priority:
    1. Azure Document Intelligence (best for tables and forms)
    2. Google Vision API (good for handwriting and Telugu)
    3. Tesseract OCR (local fallback)
    """
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.jpg', '.jpeg', '.png', '.gif', '.webp'}
    
    def __init__(self, prefer_azure: bool = True):
        """
        Initialize OCR service.
        
        Args:
            prefer_azure: Whether to prefer Azure over Google Vision
        """
        self.prefer_azure = prefer_azure
        
        # Initialize Azure
        self.azure_available = self._check_azure()
        
        # Initialize Google Vision
        self.vision_client = None
        self.translate_client = None
        self._init_google_clients()
        
        # Check Tesseract
        self.tesseract_available = self._check_tesseract()
        
        logger.info(f"OCR Service initialized - Azure: {self.azure_available}, "
                   f"Vision: {self.vision_client is not None}, "
                   f"Tesseract: {self.tesseract_available}")
    
    def _check_azure(self) -> bool:
        """Check if Azure Document Intelligence is configured."""
        endpoint = os.environ.get("AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT", "")
        key = os.environ.get("AZURE_DOCUMENT_INTELLIGENCE_KEY", "")
        return bool(endpoint and key)
    
    def _init_google_clients(self):
        """Initialize Google Cloud clients."""
        vision_creds = os.environ.get("GOOGLE_VISION_CREDENTIALS", 
                                      "/app/backend/credentials/google_vision_4.json")
        translate_creds = os.environ.get("GOOGLE_TRANSLATE_CREDENTIALS",
                                        "/app/backend/credentials/google_vision_3.json")
        
        # Vision API
        if os.path.exists(vision_creds):
            try:
                from google.cloud import vision
                from google.oauth2 import service_account
                
                credentials = service_account.Credentials.from_service_account_file(vision_creds)
                self.vision_client = vision.ImageAnnotatorClient(credentials=credentials)
                logger.info("Google Vision API initialized")
            except Exception as e:
                logger.warning(f"Failed to initialize Google Vision: {e}")
        
        # Translate API
        if os.path.exists(translate_creds):
            try:
                from google.cloud import translate_v2 as translate
                from google.oauth2 import service_account
                
                credentials = service_account.Credentials.from_service_account_file(translate_creds)
                self.translate_client = translate.Client(credentials=credentials)
                logger.info("Google Translate API initialized")
            except Exception as e:
                logger.warning(f"Failed to initialize Google Translate: {e}")
    
    def _check_tesseract(self) -> bool:
        """Check if Tesseract is available."""
        try:
            import pytesseract
            pytesseract.get_tesseract_version()
            return True
        except Exception:
            return False
    
    def extract_text(self, file_path: Path, contents: Optional[bytes] = None) -> OCRResult:
        """
        Extract text from a file.
        
        Args:
            file_path: Path to the file
            contents: Optional file contents
            
        Returns:
            OCRResult with extracted text
        """
        ext = file_path.suffix.lower()
        filename = file_path.name
        
        if ext not in self.SUPPORTED_EXTENSIONS:
            return OCRResult(
                text="",
                source_file=filename,
                file_type=ext,
                success=False,
                error=f"Unsupported format: {ext}"
            )
        
        # Load contents
        if contents is None:
            try:
                with open(file_path, 'rb') as f:
                    contents = f.read()
            except Exception as e:
                return OCRResult(
                    text="",
                    source_file=filename,
                    file_type=ext,
                    success=False,
                    error=f"Failed to read: {e}"
                )
        
        try:
            if ext == '.pdf':
                return self._extract_pdf(contents, filename)
            elif ext == '.docx':
                return self._extract_docx(contents, filename)
            elif ext == '.doc':
                return self._extract_doc(contents, filename)
            elif ext in {'.jpg', '.jpeg', '.png', '.gif', '.webp'}:
                return self._extract_image(contents, filename)
            else:
                return OCRResult(
                    text="",
                    source_file=filename,
                    file_type=ext,
                    success=False,
                    error="Unknown format"
                )
        except Exception as e:
            logger.error(f"Extraction failed for {filename}: {e}")
            return OCRResult(
                text="",
                source_file=filename,
                file_type=ext,
                success=False,
                error=str(e)
            )
    
    def _extract_pdf(self, contents: bytes, filename: str) -> OCRResult:
        """Extract from PDF using Azure or fallback."""
        text_parts = []
        engine_used = "unknown"
        tables_count = 0
        confidence = 0.0
        
        # Try PyPDF2 first for text-based PDFs
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(contents))
            
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text and len(page_text.strip()) > 50:
                    text_parts.append(page_text)
            
            if len("\n".join(text_parts).strip()) > 100:
                return OCRResult(
                    text="\n\n".join(text_parts),
                    source_file=filename,
                    file_type=".pdf",
                    success=True,
                    char_count=len("\n".join(text_parts)),
                    ocr_engine="pypdf2",
                    confidence=0.85
                )
        except Exception as e:
            logger.debug(f"PyPDF2 failed: {e}")
        
        # Try Azure Document Intelligence
        if self.prefer_azure and self.azure_available:
            try:
                from services.document_intelligence_service import DocumentIntelligenceService
                import asyncio
                
                service = DocumentIntelligenceService()
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                
                try:
                    result = loop.run_until_complete(
                        service.process_document(contents, filename, "chargesheet", preprocess=False)
                    )
                    
                    if result.success:
                        return OCRResult(
                            text=result.full_text,
                            source_file=filename,
                            file_type=".pdf",
                            success=True,
                            char_count=len(result.full_text),
                            ocr_engine="azure_document_intelligence",
                            confidence=result.overall_confidence,
                            tables_extracted=len(result.tables)
                        )
                finally:
                    loop.close()
            except Exception as e:
                logger.warning(f"Azure extraction failed: {e}")
        
        # Fall back to image-based OCR
        try:
            from pdf2image import convert_from_bytes
            
            images = convert_from_bytes(contents, dpi=300, fmt='PNG')
            
            for i, image in enumerate(images):
                img_bytes = io.BytesIO()
                image.save(img_bytes, format='PNG')
                img_bytes.seek(0)
                
                page_result = self._extract_image(img_bytes.getvalue(), f"{filename}_page{i+1}")
                
                if page_result.success and page_result.text.strip():
                    text_parts.append(f"--- Page {i+1} ---\n{page_result.text}")
                    engine_used = page_result.ocr_engine
                    confidence = max(confidence, page_result.confidence)
            
            return OCRResult(
                text="\n\n".join(text_parts),
                source_file=filename,
                file_type=".pdf",
                success=bool(text_parts),
                char_count=len("\n".join(text_parts)),
                ocr_engine=engine_used,
                confidence=confidence
            )
        except Exception as e:
            return OCRResult(
                text="",
                source_file=filename,
                file_type=".pdf",
                success=False,
                error=f"PDF extraction failed: {e}"
            )
    
    def _extract_docx(self, contents: bytes, filename: str) -> OCRResult:
        """Extract from DOCX."""
        from docx import Document
        
        doc = Document(io.BytesIO(contents))
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        text = "\n".join(text_parts)
        
        return OCRResult(
            text=text,
            source_file=filename,
            file_type=".docx",
            success=True,
            char_count=len(text),
            ocr_engine="python-docx",
            confidence=0.95
        )
    
    def _extract_doc(self, contents: bytes, filename: str) -> OCRResult:
        """Extract from legacy DOC."""
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
            tmp.write(contents)
            tmp_path = tmp.name
        
        try:
            result = subprocess.run(
                ['antiword', tmp_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            
            if result.returncode == 0:
                return OCRResult(
                    text=result.stdout,
                    source_file=filename,
                    file_type=".doc",
                    success=True,
                    char_count=len(result.stdout),
                    ocr_engine="antiword",
                    confidence=0.90
                )
            else:
                return OCRResult(
                    text="",
                    source_file=filename,
                    file_type=".doc",
                    success=False,
                    error=f"antiword failed: {result.stderr}"
                )
        finally:
            os.unlink(tmp_path)
    
    def _extract_image(self, contents: bytes, filename: str) -> OCRResult:
        """
        Extract from image using available OCR engines.
        Priority: Azure > Google Vision > Tesseract
        """
        # Try Azure first
        if self.prefer_azure and self.azure_available:
            try:
                from services.document_intelligence_service import DocumentIntelligenceService
                import asyncio
                
                service = DocumentIntelligenceService()
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                
                try:
                    result = loop.run_until_complete(
                        service.process_document(contents, filename, "chargesheet")
                    )
                    
                    if result.success and result.full_text.strip():
                        return OCRResult(
                            text=result.full_text,
                            source_file=filename,
                            file_type=Path(filename).suffix.lower(),
                            success=True,
                            char_count=len(result.full_text),
                            ocr_engine="azure_document_intelligence",
                            confidence=result.overall_confidence,
                            tables_extracted=len(result.tables)
                        )
                finally:
                    loop.close()
            except Exception as e:
                logger.debug(f"Azure OCR failed: {e}")
        
        # Try Google Vision
        if self.vision_client:
            try:
                text, confidence = self._google_vision_ocr(contents)
                if text and len(text.strip()) > 10:
                    return OCRResult(
                        text=text,
                        source_file=filename,
                        file_type=Path(filename).suffix.lower(),
                        success=True,
                        char_count=len(text),
                        ocr_engine="google_vision",
                        confidence=confidence
                    )
            except Exception as e:
                logger.debug(f"Google Vision failed: {e}")
        
        # Fallback to Tesseract
        if self.tesseract_available:
            try:
                text = self._tesseract_ocr(contents)
                return OCRResult(
                    text=text,
                    source_file=filename,
                    file_type=Path(filename).suffix.lower(),
                    success=bool(text.strip()),
                    char_count=len(text),
                    ocr_engine="tesseract",
                    confidence=0.70
                )
            except Exception as e:
                logger.debug(f"Tesseract failed: {e}")
        
        return OCRResult(
            text="",
            source_file=filename,
            file_type=Path(filename).suffix.lower(),
            success=False,
            error="All OCR engines failed"
        )
    
    def _google_vision_ocr(self, contents: bytes) -> Tuple[str, float]:
        """OCR using Google Vision."""
        from google.cloud import vision
        
        image = vision.Image(content=contents)
        response = self.vision_client.text_detection(image=image)
        
        if response.error.message:
            raise Exception(response.error.message)
        
        if not response.text_annotations:
            return "", 0.0
        
        text = response.text_annotations[0].description
        
        # Calculate confidence from word confidences
        confidence = 0.85  # Default for Vision API
        
        # Translate if needed
        if self.translate_client:
            try:
                detection = self.translate_client.detect_language(text[:500])
                if detection.get('language') not in ['en', 'und']:
                    result = self.translate_client.translate(text, target_language='en')
                    text = result['translatedText']
            except Exception:
                pass
        
        return text, confidence
    
    def _tesseract_ocr(self, contents: bytes) -> str:
        """OCR using Tesseract."""
        import pytesseract
        from PIL import Image
        
        image = Image.open(io.BytesIO(contents))
        
        if image.mode in ('RGBA', 'P'):
            image = image.convert('RGB')
        
        try:
            text = pytesseract.image_to_string(image, lang='eng+tel+hin')
        except Exception:
            text = pytesseract.image_to_string(image, lang='eng')
        
        return text.strip()
    
    def batch_extract(self, file_paths: List[Path]) -> List[OCRResult]:
        """Extract from multiple files."""
        results = []
        for file_path in file_paths:
            result = self.extract_text(file_path)
            results.append(result)
            logger.info(f"Extracted {result.char_count} chars from {result.source_file} "
                       f"using {result.ocr_engine}")
        return results
    
    def get_supported_languages(self) -> List[str]:
        """Get supported languages."""
        return ['eng', 'tel', 'hin']
