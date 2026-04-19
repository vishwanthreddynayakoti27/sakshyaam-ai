"""
Station-format charge sheet DOCX renderer.

Produces a Word document matching the real Makthal-style 18-column charge sheet
layout (verified against FIR 57/2026 by Y. Bhagya Lakshmi Reddy, SI).

The renderer takes the structured output of `intelligent_charge_sheet.py`
and produces a deterministic, print-ready DOCX. Missing fields render as
visible blank placeholders so officers can fill them in on the printed form.
"""
from __future__ import annotations

from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


# Visible placeholder shown for any empty/missing field so officers can write it in
# manually on the printed form. Keep it subtle but noticeable.
BLANK = "__________"


def _blank_if_empty(val: Any, width: str = BLANK) -> str:
    """Return a visible blank placeholder if value is empty/None, else the value."""
    if val is None:
        return width
    s = str(val).strip()
    return s if s else width


def _set_cell_shading(cell, fill_color: str = "D9D9D9"):
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    shd = parse_xml(rf'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    cell._tc.get_or_add_tcPr().append(shd)


def _add_kv_row(table, num: str, field: str, value: str):
    row = table.add_row()
    cells = row.cells
    cells[0].text = num
    cells[1].text = field
    cells[2].text = value or ""
    for c in cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
                r.font.name = "Times New Roman"


def _fmt_person(p: Dict[str, Any], prefix: str = "") -> str:
    """Format a person dict into the station-style name block.
    Missing fields render as visible blank placeholders so officers
    can fill them in by hand on the printed form."""
    sal = (p.get("salutation") or "Sri.").strip()
    name = _blank_if_empty(p.get("name"), "__________________")
    father = _blank_if_empty(p.get("father_name"), "__________________")
    age = _blank_if_empty(p.get("age"), "________")
    caste = _blank_if_empty(p.get("caste"), "__________")
    occ = _blank_if_empty(p.get("occupation"), "__________")
    addr = _blank_if_empty(p.get("address"), "____________________________")
    phone = _blank_if_empty(p.get("phone"), "__________")
    role = (p.get("role") or "").strip()

    core = f"{prefix}{sal} {name}".strip()
    core += f" S/o {father}"
    bits = [core + ","]
    sub = [
        f"Age: {age}",
        f"Caste: {caste}",
        f"Occ: {occ}",
        f"R/o {addr}",
        f"Ph. {phone}",
    ]
    bits.append(", ".join(sub) + ".")
    if role:
        bits.append(f"  ({role})")
    return " ".join(bits)


def render_charge_sheet_docx(data: Dict[str, Any]) -> bytes:
    """
    Render the intelligent charge sheet structured data into DOCX bytes.
    """
    doc = Document()

    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # Narrower margins to match station format
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # === HEADER ===
    h1 = doc.add_paragraph()
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h1.add_run("C H A R G E   –   S H E E T")
    run.bold = True
    run.font.size = Pt(14)

    h2 = doc.add_paragraph()
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h2.add_run("(UNDER SECTION 193 BNSS.)")
    run.font.size = Pt(10)

    h3 = doc.add_paragraph()
    h3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h3.add_run(_blank_if_empty(data.get("court"), "IN THE COURT OF JUDICIAL FIRST CLASS MAGISTRATE AT __________"))
    run.bold = True
    run.font.size = Pt(11)

    # === FIR HEADER LINE ===
    fir_line = doc.add_paragraph()
    run = fir_line.add_run(
        f"Dist.:- {_blank_if_empty(data.get('district'))}    "
        f"Police Station: {_blank_if_empty(data.get('police_station'))}    "
        f"FIR. No: {_blank_if_empty(data.get('fir_number'))}    "
        f"Dated: {_blank_if_empty(data.get('fir_date'))}"
    )
    run.bold = True
    run.font.size = Pt(10)

    # === MAIN KEY-VALUE TABLE ===
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(0.4)
    table.columns[1].width = Inches(2.8)
    table.columns[2].width = Inches(4.0)

    hdr = table.rows[0].cells
    hdr[0].text = "No."
    hdr[1].text = "Field"
    hdr[2].text = "Value"
    for c in hdr:
        _set_cell_shading(c, "BFBFBF")
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True

    _add_kv_row(table, "1", "Final Report / Charge Sheet No.", _blank_if_empty(data.get("chargesheet_no"), "______/______"))
    _add_kv_row(table, "2", "Date", _blank_if_empty(data.get("chargesheet_date")))
    _add_kv_row(table, "3", "Act / Sections", _blank_if_empty(data.get("sections"), "______________________________"))
    _add_kv_row(table, "4", "Type of Final Report", _blank_if_empty(data.get("chargesheet_type"), "Charge sheet"))
    _add_kv_row(table, "5", "F.R. Un-occurred (False / Mistake of fact / Mistake of Law / Non-cognizable / Civil Nature)", "---")
    _add_kv_row(table, "6", "If Charge Sheet (Original / Supplementary)", _blank_if_empty(data.get("chargesheet_type"), "Original"))

    io = data.get("io") or {}
    io_name = (io.get("name") or "").strip()
    io_rank = (io.get("rank") or "").strip()
    io_station = (io.get("station") or "").strip()
    io_sal = io.get("salutation") or "Sri."
    io_line = (
        f"{io_sal} {_blank_if_empty(io_name, '__________________')}, "
        f"{_blank_if_empty(io_rank, 'SI of Police')} "
        f"{_blank_if_empty(io_station, 'PS ________')} "
        f"(IO & Filed Charge sheet)"
    )
    _add_kv_row(table, "7", "Names of the Investigating Officers", io_line)

    _add_kv_row(table, "8", "Name of the Complainant / Informant with Father's / Husband's name", _fmt_person(data.get("complainant") or {}))
    _add_kv_row(table, "9", "Detail of properties / Articles / Documents Recovered / Seized during investigation and relied upon.", _blank_if_empty(data.get("property_recovered"), "---"))

    # Accused cell (spans one row, contains A1/A2... list)
    accused_rows: List[str] = []
    for a in data.get("accused") or []:
        accused_rows.append(_fmt_person(a, prefix=f"{a.get('serial','A?')}. "))
    if not accused_rows:
        accused_rows = ["A1. __________________________________________________________________"]
    _add_kv_row(table, "10", "Particulars of Charge-Sheeted Persons", "\n".join(accused_rows))

    # Sub-fields 10a / b / c / d
    notice_dates = [a.get("section_35_3_notice_date", "") for a in (data.get("accused") or []) if a.get("section_35_3_notice_date")]
    arrest_txt = f"Served a notice U/s 35(3) BNSS to accused on {', '.join(set(notice_dates))}" if notice_dates else "--"
    _add_kv_row(table, "10a", "Date of arrest, release, Forwarded to Court", arrest_txt)
    _add_kv_row(table, "10b", "Particulars of sureties if Released on bail", "--")
    _add_kv_row(table, "10c", "Previous convictions if any", "--")
    _add_kv_row(table, "10d", "Particulars of accused Persons absconding", "--")
    _add_kv_row(table, "11", "Particulars of accused persons not charge sheeted", "--")

    # === WITNESS LIST ===
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("12. List of Witnesses")
    run.bold = True
    run.font.size = Pt(11)

    wt = doc.add_table(rows=1, cols=3)
    wt.style = "Table Grid"
    wt.autofit = False
    wt.columns[0].width = Inches(0.7)
    wt.columns[1].width = Inches(4.5)
    wt.columns[2].width = Inches(2.0)
    wh = wt.rows[0].cells
    wh[0].text = "LW No."
    wh[1].text = "Name, Parentage, Age, Caste, Occupation, Address, Phone"
    wh[2].text = "Role / Purpose"
    for c in wh:
        _set_cell_shading(c, "BFBFBF")
        for para in c.paragraphs:
            for r in para.runs:
                r.bold = True

    for w in data.get("witnesses") or []:
        row = wt.add_row().cells
        row[0].text = _blank_if_empty(w.get("serial"), "LW-__")
        row[1].text = _fmt_person(w).rstrip()
        row[2].text = _blank_if_empty(w.get("role"), "__________________")

    # If no witnesses at all, add one blank row so officer can fill in by hand
    if not (data.get("witnesses") or []):
        row = wt.add_row().cells
        row[0].text = "LW-__"
        row[1].text = "_______________________________________________________________________"
        row[2].text = "__________________"

    # === BRIEF FACTS ===
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("13. Brief Facts of the Case:")
    run.bold = True
    run.font.size = Pt(11)

    brief = (data.get("brief_facts") or "").strip()
    if not brief:
        bp = doc.add_paragraph()
        run = bp.add_run("_" * 80)
        run.font.size = Pt(11)
        for _ in range(6):
            bp = doc.add_paragraph()
            bp.add_run("_" * 90).font.size = Pt(11)
    else:
        for para_text in [t.strip() for t in brief.split("\n\n") if t.strip()]:
            bp = doc.add_paragraph()
            bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = bp.add_run(para_text)
            run.font.size = Pt(11)

    # === PRAYER ===
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(data.get("prayer") or "")
    run.font.size = Pt(11)

    # === ANNEXURES ===
    doc.add_paragraph()
    ann = doc.add_paragraph()
    run = ann.add_run(
        f"17. Is ack. Copy of notice to complainant enclosed: {data.get('notice_ack_enclosed') or '--'}\n"
        f"18. Dispatched on: {data.get('chargesheet_date') or ''}"
    )
    run.font.size = Pt(10)

    # === SIGNATURE BLOCK ===
    doc.add_paragraph()
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = sig.add_run("Signature of the Investigation Officer\nSubmitting charge-sheet.\n")
    run.bold = True
    run.font.size = Pt(10)

    sig2 = doc.add_paragraph()
    sig2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = sig2.add_run(
        f"{_blank_if_empty(io.get('name'), '__________________')},\n"
        f"{_blank_if_empty(io.get('rank'), 'SI of Police')}, of {_blank_if_empty(io.get('station'), 'PS __________')}"
    )
    run.font.size = Pt(10)

    # Serialize to bytes
    import io as _io
    out = _io.BytesIO()
    doc.save(out)
    return out.getvalue()
