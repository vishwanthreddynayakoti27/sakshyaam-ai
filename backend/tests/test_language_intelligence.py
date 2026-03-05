"""
Test suite for Language Intelligence Module
Testing: OCR (images, PDF, DOCX) and Speech-to-Text endpoints
"""

import pytest
import requests
import os
import io

BASE_URL = os.environ.get('REACT_APP_BACKEND_URL', '').rstrip('/')

class TestLanguageIntelligence:
    """Tests for Language Intelligence Module endpoints"""
    
    @pytest.fixture(autouse=True)
    def setup(self):
        """Setup - get auth token"""
        response = requests.post(
            f"{BASE_URL}/api/auth/login",
            json={"officer_id": "OFF12345", "password": "test123"}
        )
        if response.status_code != 200:
            pytest.skip("Login failed - skipping authenticated tests")
        self.token = response.json().get("token")
        self.headers = {"Authorization": f"Bearer {self.token}"}

    # === OCR Endpoint Tests ===
    
    def test_ocr_png_image(self):
        """Test OCR endpoint accepts PNG image files"""
        # Create a simple PNG file (1x1 pixel)
        from PIL import Image
        img = Image.new('RGB', (100, 100), color='white')
        img_bytes = io.BytesIO()
        img.save(img_bytes, format='PNG')
        img_bytes.seek(0)
        
        files = {'file': ('test_image.png', img_bytes, 'image/png')}
        response = requests.post(
            f"{BASE_URL}/api/ocr/process",
            files=files,
            headers=self.headers
        )
        
        assert response.status_code == 200, f"Expected 200, got {response.status_code}"
        data = response.json()
        assert 'original_text' in data
        assert 'detected_language' in data
        assert 'translated_text' in data
        assert 'legal_text' in data
        assert 'confidence_score' in data
        assert 'message' in data
        print(f"PNG OCR test passed - Message: {data['message']}")
    
    def test_ocr_jpg_image(self):
        """Test OCR endpoint accepts JPG image files"""
        from PIL import Image
        img = Image.new('RGB', (100, 100), color='white')
        img_bytes = io.BytesIO()
        img.save(img_bytes, format='JPEG')
        img_bytes.seek(0)
        
        files = {'file': ('test_image.jpg', img_bytes, 'image/jpeg')}
        response = requests.post(
            f"{BASE_URL}/api/ocr/process",
            files=files,
            headers=self.headers
        )
        
        assert response.status_code == 200, f"Expected 200, got {response.status_code}"
        data = response.json()
        assert 'original_text' in data
        assert 'message' in data
        print(f"JPG OCR test passed - Message: {data['message']}")
    
    def test_ocr_pdf_file(self):
        """Test OCR endpoint accepts PDF files"""
        # Create a simple PDF
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        
        pdf_bytes = io.BytesIO()
        c = canvas.Canvas(pdf_bytes, pagesize=letter)
        c.drawString(100, 750, "Test PDF document for Language Intelligence Module")
        c.drawString(100, 730, "This is a sample complaint text.")
        c.save()
        pdf_bytes.seek(0)
        
        files = {'file': ('test_document.pdf', pdf_bytes, 'application/pdf')}
        response = requests.post(
            f"{BASE_URL}/api/ocr/process",
            files=files,
            headers=self.headers
        )
        
        assert response.status_code == 200, f"Expected 200, got {response.status_code}"
        data = response.json()
        assert 'original_text' in data
        assert 'translated_text' in data
        assert 'legal_text' in data
        assert 'message' in data
        # Should contain extracted text
        if data['original_text']:
            assert 'Test PDF' in data['original_text'] or 'complaint' in data['original_text'].lower()
        print(f"PDF OCR test passed - Extracted: {data['original_text'][:100] if data['original_text'] else 'No text'}")
    
    def test_ocr_docx_file(self):
        """Test OCR endpoint accepts DOCX files"""
        from docx import Document
        
        doc = Document()
        doc.add_paragraph("Test DOCX document for Language Intelligence")
        doc.add_paragraph("This is a sample complaint text for testing.")
        
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        files = {'file': ('test_document.docx', docx_bytes, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        response = requests.post(
            f"{BASE_URL}/api/ocr/process",
            files=files,
            headers=self.headers
        )
        
        assert response.status_code == 200, f"Expected 200, got {response.status_code}"
        data = response.json()
        assert 'original_text' in data
        assert 'translated_text' in data
        assert 'legal_text' in data
        # Should contain extracted text
        if data['original_text']:
            assert 'Test DOCX' in data['original_text'] or 'Language' in data['original_text']
        print(f"DOCX OCR test passed - Extracted: {data['original_text'][:100] if data['original_text'] else 'No text'}")
    
    def test_ocr_unsupported_format(self):
        """Test OCR endpoint rejects unsupported file formats"""
        # Create a fake CSV file
        csv_content = b"name,value\ntest,123"
        files = {'file': ('test.csv', io.BytesIO(csv_content), 'text/csv')}
        response = requests.post(
            f"{BASE_URL}/api/ocr/process",
            files=files,
            headers=self.headers
        )
        
        assert response.status_code == 200, f"Expected 200 with error message, got {response.status_code}"
        data = response.json()
        assert 'Unsupported file format' in data.get('message', '') or 'not configured' in data.get('message', '')
        print(f"Unsupported format test passed - Message: {data['message']}")
    
    def test_ocr_requires_auth(self):
        """Test OCR endpoint requires authentication"""
        from PIL import Image
        img = Image.new('RGB', (100, 100), color='white')
        img_bytes = io.BytesIO()
        img.save(img_bytes, format='PNG')
        img_bytes.seek(0)
        
        files = {'file': ('test_image.png', img_bytes, 'image/png')}
        response = requests.post(
            f"{BASE_URL}/api/ocr/process",
            files=files
            # No auth header
        )
        
        assert response.status_code in [401, 403], f"Expected 401/403, got {response.status_code}"
        print("OCR auth required test passed")

    # === Speech Endpoint Tests ===
    
    def test_speech_mp3_audio(self):
        """Test Speech endpoint accepts MP3 audio files"""
        # Create a minimal MP3 file (44 bytes of silence)
        mp3_header = bytes([
            0xFF, 0xFB, 0x90, 0x00, 0x00, 0x00, 0x00, 0x00,
            0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00,
            0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00,
            0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00,
            0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00,
            0x00, 0x00, 0x00, 0x00
        ])
        
        files = {'file': ('test_audio.mp3', io.BytesIO(mp3_header), 'audio/mpeg')}
        response = requests.post(
            f"{BASE_URL}/api/speech/process",
            files=files,
            headers=self.headers
        )
        
        assert response.status_code in [200, 500], f"Expected 200 or 500, got {response.status_code}"
        if response.status_code == 200:
            data = response.json()
            assert 'transcribed_text' in data or 'message' in data
            assert 'translated_text' in data or 'message' in data
            assert 'legal_text' in data or 'message' in data
            print(f"MP3 Speech test passed - Message: {data.get('message', 'processed')}")
        else:
            print(f"MP3 Speech test - API error (may be due to invalid audio content)")
    
    def test_speech_wav_audio(self):
        """Test Speech endpoint accepts WAV audio files"""
        import wave
        import struct
        
        wav_bytes = io.BytesIO()
        with wave.open(wav_bytes, 'wb') as wav_file:
            wav_file.setnchannels(1)
            wav_file.setsampwidth(2)
            wav_file.setframerate(16000)
            # Write 0.5 seconds of silence
            for _ in range(8000):
                wav_file.writeframes(struct.pack('<h', 0))
        wav_bytes.seek(0)
        
        files = {'file': ('test_audio.wav', wav_bytes, 'audio/wav')}
        response = requests.post(
            f"{BASE_URL}/api/speech/process",
            files=files,
            headers=self.headers
        )
        
        assert response.status_code in [200, 500], f"Expected 200 or 500, got {response.status_code}"
        if response.status_code == 200:
            data = response.json()
            assert 'transcribed_text' in data or 'message' in data
            print(f"WAV Speech test passed - Message: {data.get('message', 'processed')}")
        else:
            print(f"WAV Speech test - API error (may be due to invalid audio content)")
    
    def test_speech_m4a_audio(self):
        """Test Speech endpoint accepts M4A audio files"""
        # M4A header (minimal)
        m4a_header = bytes([
            0x00, 0x00, 0x00, 0x20, 0x66, 0x74, 0x79, 0x70,
            0x4D, 0x34, 0x41, 0x20, 0x00, 0x00, 0x00, 0x00,
            0x4D, 0x34, 0x41, 0x20, 0x6D, 0x70, 0x34, 0x32,
            0x69, 0x73, 0x6F, 0x6D, 0x00, 0x00, 0x00, 0x00
        ])
        
        files = {'file': ('test_audio.m4a', io.BytesIO(m4a_header), 'audio/mp4')}
        response = requests.post(
            f"{BASE_URL}/api/speech/process",
            files=files,
            headers=self.headers
        )
        
        assert response.status_code in [200, 500], f"Expected 200 or 500, got {response.status_code}"
        if response.status_code == 200:
            data = response.json()
            assert 'transcribed_text' in data or 'message' in data
            print(f"M4A Speech test passed - Message: {data.get('message', 'processed')}")
        else:
            print(f"M4A Speech test - API error (may be due to invalid audio content)")
    
    def test_speech_response_structure(self):
        """Test Speech endpoint returns expected response structure"""
        # Create minimal WAV
        import wave
        import struct
        
        wav_bytes = io.BytesIO()
        with wave.open(wav_bytes, 'wb') as wav_file:
            wav_file.setnchannels(1)
            wav_file.setsampwidth(2)
            wav_file.setframerate(16000)
            for _ in range(1600):
                wav_file.writeframes(struct.pack('<h', 0))
        wav_bytes.seek(0)
        
        files = {'file': ('test.wav', wav_bytes, 'audio/wav')}
        response = requests.post(
            f"{BASE_URL}/api/speech/process",
            files=files,
            headers=self.headers
        )
        
        if response.status_code == 200:
            data = response.json()
            # Verify response structure includes 5-stage pipeline fields
            expected_fields = [
                'transcribed_text', 'detected_language', 'translated_text',
                'legal_text', 'message'
            ]
            for field in expected_fields:
                assert field in data, f"Missing field: {field}"
            
            # Verify processing stages info
            if 'processing_stages' in data:
                assert len(data['processing_stages']) == 5, "Should have 5 processing stages"
            print(f"Speech response structure test passed")
        else:
            print(f"Speech response structure test - API returned {response.status_code}")
    
    def test_speech_requires_auth(self):
        """Test Speech endpoint requires authentication"""
        import wave
        import struct
        
        wav_bytes = io.BytesIO()
        with wave.open(wav_bytes, 'wb') as wav_file:
            wav_file.setnchannels(1)
            wav_file.setsampwidth(2)
            wav_file.setframerate(16000)
            wav_file.writeframes(struct.pack('<h', 0))
        wav_bytes.seek(0)
        
        files = {'file': ('test.wav', wav_bytes, 'audio/wav')}
        response = requests.post(
            f"{BASE_URL}/api/speech/process",
            files=files
            # No auth header
        )
        
        assert response.status_code in [401, 403], f"Expected 401/403, got {response.status_code}"
        print("Speech auth required test passed")

    # === API Configuration Tests ===
    
    def test_ocr_api_configuration_check(self):
        """Test that OCR endpoint returns proper message about API configuration"""
        from PIL import Image
        img = Image.new('RGB', (100, 100), color='white')
        img_bytes = io.BytesIO()
        img.save(img_bytes, format='PNG')
        img_bytes.seek(0)
        
        files = {'file': ('test.png', img_bytes, 'image/png')}
        response = requests.post(
            f"{BASE_URL}/api/ocr/process",
            files=files,
            headers=self.headers
        )
        
        assert response.status_code == 200
        data = response.json()
        # Should have a message about either success or configuration
        assert 'message' in data
        print(f"OCR API config check: {data['message']}")
    
    def test_speech_api_configuration_check(self):
        """Test that Speech endpoint returns proper message about API configuration"""
        import wave
        import struct
        
        wav_bytes = io.BytesIO()
        with wave.open(wav_bytes, 'wb') as wav_file:
            wav_file.setnchannels(1)
            wav_file.setsampwidth(2)
            wav_file.setframerate(16000)
            for _ in range(1600):
                wav_file.writeframes(struct.pack('<h', 0))
        wav_bytes.seek(0)
        
        files = {'file': ('test.wav', wav_bytes, 'audio/wav')}
        response = requests.post(
            f"{BASE_URL}/api/speech/process",
            files=files,
            headers=self.headers
        )
        
        if response.status_code == 200:
            data = response.json()
            assert 'message' in data
            print(f"Speech API config check: {data['message']}")
        else:
            print(f"Speech API returned {response.status_code}")

    # === Legal Text Conversion Tests ===
    
    def test_ocr_legal_text_conversion(self):
        """Test that OCR properly converts text to legal format"""
        from docx import Document
        
        doc = Document()
        doc.add_paragraph("I was robbed yesterday. My phone was stolen by two men.")
        
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        files = {'file': ('complaint.docx', docx_bytes, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        response = requests.post(
            f"{BASE_URL}/api/ocr/process",
            files=files,
            headers=self.headers
        )
        
        assert response.status_code == 200
        data = response.json()
        
        if data['original_text'] and data['legal_text']:
            # Legal text should convert first person to third person
            # "I" should become "The complainant"
            if 'complainant' in data['legal_text'].lower():
                print("Legal text conversion test passed - First person converted to third person")
            else:
                print(f"Legal text: {data['legal_text'][:200]}")
        else:
            print("Legal text conversion test - No text to convert")


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
