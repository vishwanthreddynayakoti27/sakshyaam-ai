"""
Validate the new FIXED-LAYOUT renderers (matching the user's real Telangana
police samples — 57-26 Chargesheet.pdf, 57-26 CD 1.pdf, 236 remand.pdf).

Tests:
  1. Charge Sheet renders all signature/title elements + 18-row table layout
  2. Sparse data → blanks render as `_____` (no AI invention)
  3. Case Diary Part-I renders header strip + numbered fields + closing line
  4. Remand renders "REMAND CASE DIARY" + "Honoured Sir," + prayer clause
  5. Aadhaar auto-extraction picks up Aadhaar fields from staged file OCR text
  6. Layout is IDENTICAL between two completely different cases
"""
import sys
sys.path.insert(0, "/app/backend")

import io
from docx import Document

from services.fixed_layout_renderer import (
    render_charge_sheet,
    render_case_diary_part1,
    render_remand_report,
    extract_aadhaar_from_files,
    BLANK,
)


def docx_text(b: bytes) -> str:
    """All visible text from a DOCX (paragraphs + table cells)."""
    doc = Document(io.BytesIO(b))
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def docx_headings(b: bytes):
    """Bold paragraph headings (used for structural-equality test)."""
    doc = Document(io.BytesIO(b))
    out = []
    for p in doc.paragraphs:
        if p.text and any(r.bold for r in p.runs if r.text):
            out.append(p.text)
    return out


# --- Test 1: full data → values plug in correctly ---
full_case = {
    "police_station": "Makthal",
    "district": "Narayanpet",
    "fir_number": "57/2026",
    "fir_date": "22.02.2026",
    "today_date": "26.03.2026",
    "sections": "118(2), 115(2), 352 R/w 3(5) BNS",
    "court_name": "Judicial First Class Magistrate",
    "court_place": "Makthal",
    "io": {"name": "Y. Bhagya Lakshmi Reddy", "designation": "SI of Police", "phone": "9876543210"},
    "complainant": {
        "salutation": "Sri.",
        "name": "Chandapuram Manikanta",
        "father": "Sri. Chandrashekar",
        "age": "22",
        "caste": "Mudiraj",
        "occupation": "Business",
        "address": "Nethajinagar, Makthal Mandal, Narayanpet",
        "phone": "9441016205",
    },
    "accused": [
        {"name": "Pujari Nandakishor", "father": "Pujari Subhash", "age": "36",
         "caste": "Yadav", "occupation": "Agriculture",
         "address": "H.No. 2-72 Maganoor Village & Mandal, Narayanpet District",
         "phone": "9959282848", "aadhaar_number": "1234 5678 9012"},
    ],
    "witnesses": [
        {"name": "Witness 1", "father": "X", "age": "30", "caste": "A", "occupation": "B",
         "address": "Makthal", "phone": "1", "type": "Eye-witness"},
        {"name": "Witness 2", "father": "Y", "age": "40", "caste": "C", "occupation": "D",
         "address": "Makthal", "phone": "2", "type": "Hearsay"},
    ],
    "brief_facts": "On 22.02.2026 the accused beat the complainant near the bus stand.",
    "arrest_release": "Served a notice u/s 35(3) BNSS to accused on 23.2.2026",
    "dispatch_date": "26.03.2026",
}

cs1 = render_charge_sheet(full_case)
assert cs1[:2] == b"PK", "Not a valid DOCX"
text1 = docx_text(cs1)
# Title (verbatim from sample: spaced-out "C H A R G E – S H E E T")
assert "C H A R G E – S H E E T" in text1
assert "(UNDER SECTION 193 BNSS.)" in text1
assert "JUDICIAL FIRST CLASS MAGISTRATE" in text1.upper()
assert "MAKTHAL" in text1.upper()
# 18 mandatory section labels (must appear verbatim — sample-matched)
mandatory_labels = [
    "Charge Sheet No.",                                    # 02
    "Date of Charge",                                      # 03
    "Act and Section of Law",                              # 04
    "Type of the final report",                            # 05
    "If final report is un-occurred",                      # 06
    "If charge sheet is original or supplementary.",       # 07
    "Name and rank of the I.O (s)",                        # 08
    "Name and Address of the complainant or informant",    # 09
    "Details of property seized during the course of investigation.",  # 10
    "Particulars of accused persons charge sheeted",       # 11
    "Particulars of the accused persons not charge sheeted",  # 12
    "Particulars of witnesses to be examined",             # 13 (heading paragraph)
    "If F.R. is false, indicate action taken",             # 14
    "Result of Laboratory Analysis",                       # 15
    "Brief facts of the case",                             # 16
    "Is ack. copy of notice to complainant is enclosed",   # 17
    "Dispatched on",                                       # 18
]
for lbl in mandatory_labels:
    assert lbl in text1, f"Missing mandatory section label: {lbl!r}"

# Header strip
assert "Narayanpet" in text1
assert "57/2026" in text1
assert "22.02.2026" in text1
# IO
assert "Bhagya Lakshmi Reddy" in text1
assert "Buddamolla" not in text1  # sanity — ours has different complainant
assert "Chandapuram Manikanta" in text1
# Accused
assert "Pujari Nandakishor" in text1
assert "A1." in text1
# Witness
assert "LW-1" in text1
# Closing
assert "Hence the charge sheet." in text1
assert "Submitting chargesheet" in text1
print("[OK] 1. Charge Sheet renders authentic 18-section Telangana layout (verbatim)")


# --- Test 2: sparse data → blanks render as `_____` ---
sparse = {"police_station": "Test PS", "fir_number": "1/2026"}
cs2 = render_charge_sheet(sparse)
text2 = docx_text(cs2)
assert "Test PS" in text2
assert "1/2026" in text2
assert text2.count(BLANK) >= 5, f"Expected ≥5 BLANK placeholders, got {text2.count(BLANK)}"
# Required fixed-text elements still present even with sparse data
for must in ["C H A R G E – S H E E T", "Particulars of accused persons charge sheeted",
             "Particulars of witnesses to be examined", "Hence the charge sheet.",
             "Submitting chargesheet", "Charge Sheet No.", "Date of Charge",
             "Act and Section of Law", "Type of the final report",
             "Result of Laboratory Analysis", "Brief facts of the case",
             "Is ack. copy of notice to complainant is enclosed", "Dispatched on"]:
    assert must in text2, f"Missing fixed element: {must!r}"
print(f"[OK] 2. Sparse data: {text2.count(BLANK)} BLANK placeholders, all fixed text intact")


# --- Test 3: case diary ---
cd = render_case_diary_part1(full_case)
text_cd = docx_text(cd)
assert "CASE DIARY" in text_cd
assert "(Part-I)" in text_cd
assert "Section 193(8) BNSS" in text_cd or "193(8)" in text_cd
assert "F.I.R No." in text_cd
assert "Date and time of report" in text_cd
assert "Property Lost" in text_cd
assert "Property recovered" in text_cd
assert "witnesses examined" in text_cd
assert "Closed the CD for the day" in text_cd
assert "Copy submitted to the SDPO" in text_cd
print("[OK] 3. Case Diary Part-I matches authentic Telangana CD format")


# --- Test 4: remand ---
rr = render_remand_report({**full_case, "remand_type": "judicial",
                          "occurrence_dtp": "22.02.2026 at 18:00 hours, Makthal Bus Stand",
                          "action_taken_datetime": "22.02.2026 at 19:30 hours"})
text_rr = docx_text(rr)
assert "REMAND CASE DIARY" in text_rr
assert "Part-I" in text_rr
assert "JUDICIAL MAGISTRATE OF FIRST CLASS" in text_rr.upper()
assert "Honoured Sir," in text_rr
assert "Name of the Investigating Officer" in text_rr
assert "Date and place of occurrence" in text_rr
assert "Reasons for arrest" in text_rr
assert "Hence the remand report" in text_rr
assert "produced before the Hon'ble Court" in text_rr
assert "judicial remand custody" in text_rr
assert "Encl:" in text_rr
print("[OK] 4. Remand Case Diary matches authentic letter-format")


# --- Test 5: Aadhaar auto-extraction ---
sample_files = [
    {"filename": "fir_copy.pdf", "ocr_text": "FIR Number 57/2026 dated 22.02.2026"},
    {"filename": "aadhaar_card.jpg", "ocr_text": (
        "Government of India\n"
        "Unique Identification Authority of India\n"
        "Pujari Nandakishor\n"
        "DOB: 12/05/1989\n"
        "Male\n"
        "S/O Pujari Subhash, H.No. 2-72 Maganoor Village & Mandal\n"
        "1234 5678 9012\n"
    )},
]
aad = extract_aadhaar_from_files(sample_files)
assert aad["aadhaar_number"] == "123456789012"
assert aad["aadhaar_dob"] == "12/05/1989"
assert aad["aadhaar_gender"] == "Male"
assert "Pujari" in aad["aadhaar_address"]
print(f"[OK] 5. Aadhaar auto-extraction: number={aad['aadhaar_number']}, dob={aad['aadhaar_dob']}, "
      f"gender={aad['aadhaar_gender']}")


# --- Test 6: identical layout across very different cases ---
import re as _re

def _structural(headings):
    """Drop case-specific lines so we compare structure only."""
    out = []
    for h in headings:
        if not h:
            continue
        # Drop lines that contain personal data / dates / IDs
        if any(s in h for s in ["FIR", "PS ", "P.S.:", "Dist:", "Date:", "Dated:",
                                "SDPO", "Y. Bhagya"]):
            continue
        # Drop dynamic per-accused / per-witness labels
        if _re.match(r"^(A|LW-?|MO)\d+\.?$", h.strip()):
            continue
        out.append(h.strip())
    return out

case_a = {"police_station": "PS A", "fir_number": "1/2026", "complainant": {"name": "Alice"}}
case_b = {"police_station": "PS B", "fir_number": "999/2026",
          "accused": [{"name": "Bob"}, {"name": "Carol"}, {"name": "Dave"}],
          "witnesses": [{"name": f"W{i}"} for i in range(8)]}

for name, fn in [("charge_sheet", render_charge_sheet),
                 ("case_diary_part1", render_case_diary_part1),
                 ("remand_report", render_remand_report)]:
    da = fn(case_a)
    db = fn(case_b)
    sa = _structural(docx_headings(da))
    sb = _structural(docx_headings(db))
    assert sa == sb, (
        f"{name} structural headings differ between cases!\n"
        f"  A: {sa}\n  B: {sb}"
    )
    print(f"[OK] 6.{name}: identical structure across two completely different cases ({len(sa)} headings)")

print("\nALL FIXED-LAYOUT TESTS PASSED")
