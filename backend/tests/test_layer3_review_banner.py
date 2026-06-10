"""
Smoke test for the LAYER 3 Review Summary banner injected at the top of
the generated chargesheet DOCX by `fixed_layout_renderer.render_charge_sheet`.

Ensures:
  1. If no quality_review is passed, the renderer does NOT crash (skips banner).
  2. If a complete quality_review is passed, the rendered DOCX contains:
     - the status pill text (READY/REVIEW/OFFICER MUST COMPLETE)
     - the "X% complete" string
     - each bullet of items_to_verify
     - the C1..C9 audit summary line
"""
from __future__ import annotations

import io
import sys
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document  # noqa: E402

from services.fixed_layout_renderer import render_charge_sheet  # noqa: E402


def _extract_all_text(docx_bytes: bytes) -> str:
    doc = Document(io.BytesIO(docx_bytes))
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text)
    return "\n".join(parts)


BASE_CASE = {
    "court_name": "JUDICIAL FIRST CLASS MAGISTRATE",
    "court_place": "MAKTHAL",
    "police_station": "Makthal",
    "district": "Narayanpet",
    "fir_number": "100/2025",
    "fir_date": "23.04.2025",
    "io": {"name": "K Lal Singh", "designation": "HC 248"},
    "complainant": {"name": "Jingiti Aruna", "father": "Late Chinna Thayappa"},
    "accused": [{"name": "John Doe"}],
    "witnesses": [{"name": "LW-1"}],
    "sections": "115(2) BNS",
}


class TestLayer3ReviewBanner(unittest.TestCase):

    def test_renderer_without_quality_review_does_not_crash(self):
        docx_bytes = render_charge_sheet(BASE_CASE)
        self.assertGreater(len(docx_bytes), 5000, "DOCX should not be tiny")
        text = _extract_all_text(docx_bytes)
        # Banner should not appear (no QR provided)
        self.assertNotIn("DRAFT QUALITY REPORT", text)
        # But normal sections should still be present
        self.assertIn("C H A R G E", text)
        self.assertIn("100/2025", text)

    def test_renderer_renders_complete_review_banner(self):
        case = dict(BASE_CASE)
        case["quality_review"] = {
            "completion_pct": 78,
            "fixes_applied": [
                {"check": "C2", "before": "simple", "after": "grievous", "reason": "fracture → grievous"}
            ],
            "items_to_verify": [
                "complainant.phone",
                "accused[0].caste",
                "witnesses[1].address",
            ],
            "audit_checks": {
                "C1_dual_listed_person": "PASS",
                "C2_injury_gravity": "FIXED",
                "C9_injuries_wrong_person": "PASS",
            },
            "overall_status": "REVIEW_NEEDED",
        }
        docx_bytes = render_charge_sheet(case)
        text = _extract_all_text(docx_bytes)
        # Banner header present
        self.assertIn("DRAFT QUALITY REPORT", text)
        self.assertIn("REVIEW NEEDED", text)
        # Completion %
        self.assertIn("78% complete", text)
        # Items to verify
        self.assertIn("complainant.phone", text)
        self.assertIn("accused[0].caste", text)
        self.assertIn("witnesses[1].address", text)
        # Audit grid line
        self.assertIn("C1=PASS", text)
        self.assertIn("C2=FIXED", text)
        # Fix details (truncated)
        self.assertIn("fracture", text)

    def test_ready_to_file_status_renders_green_pill(self):
        case = dict(BASE_CASE)
        case["quality_review"] = {
            "completion_pct": 100,
            "fixes_applied": [],
            "items_to_verify": [],
            "audit_checks": {"C1_dual_listed_person": "PASS"},
            "overall_status": "READY_TO_FILE",
        }
        docx_bytes = render_charge_sheet(case)
        text = _extract_all_text(docx_bytes)
        self.assertIn("READY TO FILE", text)
        self.assertIn("100% complete", text)

    def test_officer_must_complete_status_renders_red_pill(self):
        case = dict(BASE_CASE)
        case["quality_review"] = {
            "completion_pct": 40,
            "fixes_applied": [],
            "items_to_verify": ["accused[0].name", "io.name"],
            "audit_checks": {"C7_lw_mentioned_but_missing": "FLAG"},
            "overall_status": "OFFICER_MUST_COMPLETE",
        }
        docx_bytes = render_charge_sheet(case)
        text = _extract_all_text(docx_bytes)
        self.assertIn("OFFICER MUST COMPLETE", text)
        self.assertIn("40% complete", text)
        self.assertIn("C7=FLAG", text)


if __name__ == "__main__":
    unittest.main()
