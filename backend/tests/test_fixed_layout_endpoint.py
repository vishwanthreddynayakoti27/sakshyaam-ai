"""
Integration tests for GET /api/staging/render-fixed/{doc_type}/{case_id}

Covers:
 1. Auth required (401 without token)
 2. Invalid doc_type → 400
 3. Missing case folder → 404
 4. Metadata missing → 404
 5. Happy path: create-case → render all 3 doc_types → 200 + DOCX bytes
 6. Filename + content-type headers correct
 7. Aadhaar auto-extraction: metadata.json with ocr_text populates A1.aadhaar_number
 8. 0 credits: profile.credits unchanged before/after
"""
import io
import json
import os
import time
import uuid
from pathlib import Path

import pytest
import requests
from docx import Document

BASE_URL = os.environ.get("REACT_APP_BACKEND_URL", "").rstrip("/")
if not BASE_URL:
    # fallback to public preview URL
    try:
        with open("/app/frontend/.env") as f:
            for line in f:
                if line.startswith("REACT_APP_BACKEND_URL="):
                    BASE_URL = line.split("=", 1)[1].strip().rstrip("/")
                    break
    except Exception:
        pass

STAGING_BASE = Path("/app/backend/staging")
OFFICER = {"officer_id": "TEST001", "password": "Test123!"}


# ------------------------ Fixtures ------------------------

@pytest.fixture(scope="module")
def api():
    s = requests.Session()
    s.headers.update({"Content-Type": "application/json"})
    return s


@pytest.fixture(scope="module")
def token(api):
    r = api.post(f"{BASE_URL}/api/auth/login", json=OFFICER, timeout=30)
    assert r.status_code == 200, f"login failed: {r.status_code} {r.text}"
    j = r.json()
    return j["token"]


@pytest.fixture(scope="module")
def auth(api, token):
    api.headers.update({"Authorization": f"Bearer {token}"})
    return api


@pytest.fixture(scope="module")
def me(auth):
    r = auth.get(f"{BASE_URL}/api/auth/profile", timeout=30)
    assert r.status_code == 200, f"profile failed: {r.status_code} {r.text}"
    return r.json()


@pytest.fixture(scope="module")
def staged_case(auth):
    """Create a case via API."""
    data = {
        "police_station": "Makthal PS",
        "district": "Narayanpet",
        "fir_number": "TEST/57/2026",
        "sections": "BNS 115(2)",
    }
    # multipart/form because endpoint uses Form(...)
    headers = {k: v for k, v in auth.headers.items() if k.lower() != "content-type"}
    r = requests.post(f"{BASE_URL}/api/staging/create-case", data=data, headers=headers, timeout=30)
    assert r.status_code == 200, f"create-case failed: {r.status_code} {r.text}"
    case_id = r.json()["case_id"]
    yield case_id


# ------------------------ Tests ------------------------

def test_unauth_returns_401():
    r = requests.get(
        f"{BASE_URL}/api/staging/render-fixed/charge_sheet/CASE-XYZ",
        timeout=30,
    )
    assert r.status_code in (401, 403), f"expected 401/403 got {r.status_code}"


def test_invalid_doc_type(auth, staged_case):
    r = auth.get(
        f"{BASE_URL}/api/staging/render-fixed/NOT_A_DOC/{staged_case}", timeout=30
    )
    assert r.status_code == 400
    assert "Invalid doc_type" in r.text


def test_missing_case_returns_404(auth):
    r = auth.get(
        f"{BASE_URL}/api/staging/render-fixed/charge_sheet/CASE-DOES-NOT-EXIST-{uuid.uuid4().hex[:6]}",
        timeout=30,
    )
    # The endpoint uses get_case_folder which CREATES the folder if missing, but metadata.json won't exist.
    # So we should get 404 either on "Case folder not found" or "Case metadata not found"
    assert r.status_code == 404, f"expected 404 got {r.status_code}: {r.text}"


@pytest.mark.parametrize("doc_type,expected_heading", [
    ("charge_sheet", "CHARGE"),
    ("case_diary_part1", "CASE DIARY"),
    ("remand_report", "REMAND REPORT"),
])
def test_render_happy_path(auth, staged_case, doc_type, expected_heading):
    r = auth.get(
        f"{BASE_URL}/api/staging/render-fixed/{doc_type}/{staged_case}", timeout=60
    )
    assert r.status_code == 200, f"{doc_type} failed: {r.status_code} {r.text[:400]}"
    ct = r.headers.get("content-type", "")
    assert "officedocument.wordprocessingml.document" in ct, f"bad content-type: {ct}"
    cd = r.headers.get("content-disposition", "")
    assert "attachment" in cd and ".docx" in cd, f"bad content-disposition: {cd}"

    # Validate DOCX bytes
    assert r.content[:2] == b"PK", "Not a valid DOCX"
    doc = Document(io.BytesIO(r.content))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                all_text += "\n" + cell.text
    assert expected_heading in all_text.upper(), f"heading not found in {doc_type}"
    # FIR number + police_station should be in rendered doc
    assert "Makthal PS" in all_text
    assert "TEST/57/2026" in all_text
    # Missing fields should render as blanks
    assert "_____" in all_text, "Blanks not present for missing fields"


def test_zero_credits_deducted(auth, me, staged_case):
    """Render all 3 docs and verify credits do not change."""
    before = auth.get(f"{BASE_URL}/api/auth/profile", timeout=30).json().get("credits", 0)
    for dt in ("charge_sheet", "case_diary_part1", "remand_report"):
        r = auth.get(
            f"{BASE_URL}/api/staging/render-fixed/{dt}/{staged_case}", timeout=60
        )
        assert r.status_code == 200
    after = auth.get(f"{BASE_URL}/api/auth/profile", timeout=30).json().get("credits", 0)
    assert before == after, f"credits changed: {before} -> {after}"


def test_aadhaar_auto_extraction(auth, me):
    """
    Seed metadata.json with a file containing Aadhaar OCR text and verify
    A1.aadhaar_number gets populated in the rendered DOCX.
    """
    officer_id = me.get("officer_id") or OFFICER["officer_id"]
    case_id = f"CASE-TESTAADHAAR-{uuid.uuid4().hex[:6].upper()}"
    folder = STAGING_BASE / officer_id / case_id
    folder.mkdir(parents=True, exist_ok=True)
    metadata = {
        "case_id": case_id,
        "officer_id": officer_id,
        "police_station": "AadTest PS",
        "district": "TestDist",
        "fir_number": "AAD/1/2026",
        "sections": "BNS 115(2)",
        "accused": [],
        "files": [
            {
                "filename": "aadhaar_card.jpg",
                "ocr_text": (
                    "Government of India\n"
                    "Unique Identification Authority of India\n"
                    "Vijay Reddy\n"
                    "DOB: 12/05/1989\n"
                    "Male\n"
                    "S/O Krishna Reddy, Hyderabad\n"
                    "1234 5678 9012\n"
                ),
            }
        ],
    }
    with open(folder / "metadata.json", "w") as f:
        json.dump(metadata, f)

    try:
        r = auth.get(
            f"{BASE_URL}/api/staging/render-fixed/charge_sheet/{case_id}", timeout=60
        )
        assert r.status_code == 200, r.text[:400]
        doc = Document(io.BytesIO(r.content))
        text = "\n".join(p.text for p in doc.paragraphs)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    text += "\n" + cell.text
        # Aadhaar number should appear (possibly spaced as "1234 5678 9012" or compact)
        assert ("1234 5678 9012" in text) or ("123456789012" in text), (
            f"Aadhaar not auto-extracted into DOCX. Text sample:\n{text[:600]}"
        )
    finally:
        # Cleanup
        try:
            (folder / "metadata.json").unlink(missing_ok=True)
            folder.rmdir()
        except Exception:
            pass
