"""
Staged Upload System for Charge Sheet Fusion
- Zero-credit file collection into Case Folder
- Single credit deduction on "Generate Triple Fusion" trigger
- Supports unlimited batch uploads
- Roll-back on failure (no credits deducted)
- ASYNC BACKGROUND PROCESSING for large batches
"""
from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from datetime import datetime, timezone
from typing import Optional, List
from pydantic import BaseModel
from pathlib import Path
import os
import jwt
import logging
import re
import shutil
import uuid
import json
import io
import tempfile
import subprocess
import asyncio

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/staging", tags=["Staged Upload"])
security = HTTPBearer()

# Staging folder base path
STAGING_BASE = Path("/app/backend/staging")
STAGING_BASE.mkdir(parents=True, exist_ok=True)

# Database connection
db = None
JWT_SECRET = os.environ['JWT_SECRET']


def set_database(database):
    global db
    db = database


async def get_current_officer(credentials: HTTPAuthorizationCredentials = Depends(security)) -> dict:
    try:
        token = credentials.credentials
        payload = jwt.decode(token, JWT_SECRET, algorithms=['HS256'])
        return payload
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")


def get_case_folder(officer_id: str, case_id: str) -> Path:
    """Get or create case staging folder."""
    folder = STAGING_BASE / officer_id / case_id
    folder.mkdir(parents=True, exist_ok=True)
    return folder


def get_file_info(file_path: Path) -> dict:
    """Get file metadata."""
    stat = file_path.stat()
    return {
        "filename": file_path.name,
        "size": stat.st_size,
        "size_kb": round(stat.st_size / 1024, 2),
        "uploaded_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
        "extension": file_path.suffix.lower()
    }


# =============================================================================
# STAGING ENDPOINTS - ZERO CREDIT COST
# =============================================================================

@router.post("/create-case")
async def create_staging_case(
    police_station: str = Form(...),
    district: str = Form(...),
    fir_number: str = Form(default=""),
    sections: str = Form(default=""),
    # 11 new V3.0 manual-input fields (filled by the police writer, never altered)
    fir_date: str = Form(default=""),
    chargesheet_no: str = Form(default=""),
    chargesheet_date: str = Form(default=""),
    report_type: str = Form(default="Charge Sheet"),
    un_occurred_reason: str = Form(default=""),
    chargesheet_type: str = Form(default="Original"),
    io_name: str = Form(default=""),
    io_rank: str = Form(default=""),
    court_name: str = Form(default=""),
    dispatch_date: str = Form(default=""),
    ack_enclosed: str = Form(default="No"),
    officer: dict = Depends(get_current_officer)
):
    """
    Create a new case folder for staging files.

    Accepts the 15-field manual-input form (4 legacy + 11 V3.0). All fields
    are persisted in metadata.json under `manual_input` and passed to the
    LLM verbatim as "CONFIRMED MANUAL INPUT — USE EXACTLY AS PROVIDED".

    COST: 0 CREDITS
    """
    case_id = f"CASE-{datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:4].upper()}"
    folder = get_case_folder(officer.get("officer_id", "unknown"), case_id)

    # Save case metadata
    metadata = {
        "case_id": case_id,
        "officer_id": officer.get("officer_id"),
        "police_station": police_station,
        "district": district,
        "fir_number": fir_number,
        "sections": sections,
        # V3.0 manual input — kept under a dedicated subobject so the
        # intelligent-charge-sheet endpoint can flag them as "Phase 1" data
        "manual_input": {
            "district": district,
            "police_station": police_station,
            "fir_number": fir_number,
            "fir_date": fir_date,
            "chargesheet_no": chargesheet_no,
            "chargesheet_date": chargesheet_date,
            "sections": sections,
            "report_type": report_type,
            "un_occurred_reason": un_occurred_reason,
            "chargesheet_type": chargesheet_type,
            "io_name": io_name,
            "io_rank": io_rank,
            "court_name": court_name,
            "dispatch_date": dispatch_date,
            "ack_enclosed": ack_enclosed,
        },
        "created_at": datetime.now(timezone.utc).isoformat(),
        "status": "staging",
        "files": []
    }

    with open(folder / "metadata.json", "w") as f:
        json.dump(metadata, f, indent=2, ensure_ascii=False)

    logger.info(
        f"Created staging case: {case_id} with full manual-input form "
        f"(0 credits, {sum(1 for v in metadata['manual_input'].values() if v)} fields populated)"
    )

    return {
        "success": True,
        "case_id": case_id,
        "message": "Case folder created with manual input locked",
        "credits_used": 0
    }


@router.post("/upload-files/{case_id}")
async def upload_files_to_staging(
    case_id: str,
    files: List[UploadFile] = File(...),
    officer: dict = Depends(get_current_officer)
):
    """
    Upload files to staging folder in batches.
    COST: 0 CREDITS - Files are saved without processing.
    Supports: PDF, DOCX, DOC, JPEG, PNG, JPG
    """
    folder = get_case_folder(officer.get("officer_id", "unknown"), case_id)
    
    if not folder.exists():
        raise HTTPException(status_code=404, detail="Case folder not found")
    
    # Load existing metadata
    metadata_path = folder / "metadata.json"
    if metadata_path.exists():
        with open(metadata_path) as f:
            metadata = json.load(f)
    else:
        metadata = {"files": []}
    
    saved_files = []
    allowed_extensions = {'.pdf', '.docx', '.doc', '.jpeg', '.jpg', '.png', '.gif', '.webp'}
    
    for file in files:
        ext = Path(file.filename).suffix.lower()
        if ext not in allowed_extensions:
            logger.warning(f"Skipped unsupported file: {file.filename}")
            continue
        
        # Generate unique filename to avoid overwrite
        unique_name = f"{uuid.uuid4().hex[:8]}_{file.filename}"
        file_path = folder / unique_name
        
        # Save file
        contents = await file.read()
        with open(file_path, "wb") as f:
            f.write(contents)
        
        file_info = {
            "original_name": file.filename,
            "saved_name": unique_name,
            "size": len(contents),
            "type": ext,
            "uploaded_at": datetime.now(timezone.utc).isoformat()
        }
        metadata["files"].append(file_info)
        saved_files.append(file_info)
    
    # Update metadata
    with open(metadata_path, "w") as f:
        json.dump(metadata, f, indent=2)
    
    logger.info(f"Staged {len(saved_files)} files for case {case_id} (0 credits)")
    
    return {
        "success": True,
        "case_id": case_id,
        "files_saved": len(saved_files),
        "total_files_in_case": len(metadata["files"]),
        "saved_files": saved_files,
        "credits_used": 0,
        "message": f"{len(saved_files)} files saved to staging. No credits deducted."
    }


@router.get("/case/{case_id}")
async def get_staged_files(
    case_id: str,
    officer: dict = Depends(get_current_officer)
):
    """
    Get list of all staged files in a case folder.
    COST: 0 CREDITS
    """
    folder = get_case_folder(officer.get("officer_id", "unknown"), case_id)
    
    if not folder.exists():
        raise HTTPException(status_code=404, detail="Case folder not found")
    
    metadata_path = folder / "metadata.json"
    if metadata_path.exists():
        with open(metadata_path) as f:
            metadata = json.load(f)
    else:
        metadata = {"files": []}
    
    return {
        "success": True,
        "case_id": case_id,
        "metadata": metadata,
        "file_count": len(metadata.get("files", [])),
        "credits_used": 0
    }


@router.delete("/case/{case_id}/file/{filename}")
async def delete_staged_file(
    case_id: str,
    filename: str,
    officer: dict = Depends(get_current_officer)
):
    """
    Delete a file from staging.
    COST: 0 CREDITS
    """
    folder = get_case_folder(officer.get("officer_id", "unknown"), case_id)
    file_path = folder / filename
    
    if file_path.exists():
        file_path.unlink()
        
        # Update metadata
        metadata_path = folder / "metadata.json"
        if metadata_path.exists():
            with open(metadata_path) as f:
                metadata = json.load(f)
            metadata["files"] = [f for f in metadata["files"] if f["saved_name"] != filename]
            with open(metadata_path, "w") as f:
                json.dump(metadata, f, indent=2)
        
        return {"success": True, "message": f"File {filename} deleted", "credits_used": 0}
    
    raise HTTPException(status_code=404, detail="File not found")


@router.get("/my-cases")
async def list_my_cases(
    officer: dict = Depends(get_current_officer)
):
    """
    List all staging cases for current officer.
    COST: 0 CREDITS
    """
    officer_folder = STAGING_BASE / officer.get("officer_id", "unknown")
    
    if not officer_folder.exists():
        return {"cases": [], "count": 0, "credits_used": 0}
    
    cases = []
    for case_folder in officer_folder.iterdir():
        if case_folder.is_dir():
            metadata_path = case_folder / "metadata.json"
            if metadata_path.exists():
                with open(metadata_path) as f:
                    metadata = json.load(f)
                cases.append({
                    "case_id": case_folder.name,
                    "fir_number": metadata.get("fir_number", ""),
                    "police_station": metadata.get("police_station", ""),
                    "file_count": len(metadata.get("files", [])),
                    "status": metadata.get("status", "staging"),
                    "created_at": metadata.get("created_at", "")
                })
    
    return {"cases": cases, "count": len(cases), "credits_used": 0}


# =============================================================================
# TRIPLE FUSION - CREDIT DEDUCTION ONLY HERE
# =============================================================================

async def extract_text_from_staged_file(file_path: Path) -> str:
    """Extract text from a staged file using OCR/parsing."""
    ext = file_path.suffix.lower()
    
    with open(file_path, "rb") as f:
        contents = f.read()
    
    # DOCX
    if ext == '.docx':
        try:
            from docx import Document
            doc = Document(io.BytesIO(contents))
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            return "\n".join(text_parts)
        except Exception as e:
            return f"[DOCX error: {e}]"
    
    # DOC (legacy)
    elif ext == '.doc':
        try:
            with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
                tmp.write(contents)
                tmp_path = tmp.name
            
            result = subprocess.run(['antiword', tmp_path], capture_output=True, text=True, timeout=30)
            os.unlink(tmp_path)
            
            if result.returncode == 0:
                return result.stdout
            return f"[DOC extraction failed: {result.stderr}]"
        except Exception as e:
            return f"[DOC error: {e}]"
    
    # PDF
    elif ext == '.pdf':
        text_parts = []
        # Phase 1 — try the embedded text layer (fast, works for digital PDFs)
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(contents))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text)
            extracted_text = "\n".join(text_parts).strip()
            if len(extracted_text) >= 80:
                return extracted_text
        except Exception as e:
            logger.warning(f"PyPDF2 phase failed for {file_path.name}: {e}")
            extracted_text = ""

        # Phase 2 — text layer was empty/sparse → render to images and OCR via Google Vision
        try:
            from pdf2image import convert_from_bytes
            pages = convert_from_bytes(contents, dpi=200, fmt="png")
            logger.info(
                f"[PDF-OCR] {file_path.name}: PyPDF2 returned "
                f"{len(extracted_text)} chars, falling back to {len(pages)}-page Vision OCR"
            )
            GOOGLE_VISION_CREDENTIALS = os.environ.get("GOOGLE_VISION_CREDENTIALS", "")
            if GOOGLE_VISION_CREDENTIALS and os.path.exists(GOOGLE_VISION_CREDENTIALS):
                from google.cloud import vision
                from google.oauth2 import service_account
                credentials = service_account.Credentials.from_service_account_file(
                    GOOGLE_VISION_CREDENTIALS
                )
                vision_client = vision.ImageAnnotatorClient(credentials=credentials)
                ocr_pages = []
                for idx, img in enumerate(pages, 1):
                    buf = io.BytesIO()
                    img.save(buf, format="PNG")
                    image = vision.Image(content=buf.getvalue())
                    response = vision_client.document_text_detection(image=image)
                    if response.error.message:
                        logger.warning(
                            f"[PDF-OCR] {file_path.name} page {idx}: "
                            f"Vision error {response.error.message}"
                        )
                        continue
                    ocr_pages.append(
                        response.full_text_annotation.text if response.full_text_annotation else ""
                    )
                combined = "\n--- PAGE BREAK ---\n".join(p for p in ocr_pages if p).strip()
                if combined:
                    return combined
            # No vision creds → fall back to whatever PyPDF2 gave us, even if short
            return extracted_text or "[Scanned PDF — OCR unavailable: configure GOOGLE_VISION_CREDENTIALS]"
        except Exception as e:
            logger.warning(f"PDF OCR fallback failed for {file_path.name}: {e}")
            return extracted_text or f"[PDF OCR error: {e}]"
    
    # Images (JPG, PNG) - Use OCR
    elif ext in ['.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp', '.tiff', '.tif']:
        GOOGLE_VISION_CREDENTIALS = os.environ.get('GOOGLE_VISION_CREDENTIALS', '')
        GOOGLE_TRANSLATE_CREDENTIALS = os.environ.get('GOOGLE_TRANSLATE_CREDENTIALS', '')
        
        if GOOGLE_VISION_CREDENTIALS and os.path.exists(GOOGLE_VISION_CREDENTIALS):
            try:
                from google.cloud import vision
                from google.oauth2 import service_account
                
                credentials = service_account.Credentials.from_service_account_file(GOOGLE_VISION_CREDENTIALS)
                vision_client = vision.ImageAnnotatorClient(credentials=credentials)
                
                image = vision.Image(content=contents)
                response = vision_client.text_detection(image=image)
                
                if response.text_annotations:
                    text = response.text_annotations[0].description
                    
                    # Translate if needed
                    if GOOGLE_TRANSLATE_CREDENTIALS and os.path.exists(GOOGLE_TRANSLATE_CREDENTIALS):
                        from google.cloud import translate_v2 as translate
                        trans_creds = service_account.Credentials.from_service_account_file(GOOGLE_TRANSLATE_CREDENTIALS)
                        translate_client = translate.Client(credentials=trans_creds)
                        translation = translate_client.translate(text, target_language='en')
                        return translation['translatedText']
                    
                    return text
                return "[No text detected in image]"
            except Exception as e:
                return f"[OCR error: {e}]"
        else:
            return f"[OCR not configured - file: {file_path.name}]"

    # Plain text / structured text formats
    elif ext in ['.txt', '.md', '.csv', '.json', '.html', '.htm', '.xml', '.log', '.rtf']:
        try:
            return contents.decode('utf-8', errors='replace').strip()
        except Exception as e:
            return f"[Text decode error: {e}]"

    # Legacy MS Word (.doc) — try antiword if installed
    elif ext == '.doc':
        try:
            result = subprocess.run(
                ['antiword', str(file_path)],
                capture_output=True, text=True, timeout=20
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except FileNotFoundError:
            pass
        except Exception as e:
            logger.warning(f"antiword failed on {file_path.name}: {e}")
        # Final fallback: read as latin-1 to capture any embedded ASCII strings
        try:
            return contents.decode('latin-1', errors='replace')
        except Exception as e:
            return f"[.doc parse error: {e}]"

    # Excel
    elif ext in ['.xlsx', '.xls']:
        try:
            from openpyxl import load_workbook
            wb = load_workbook(io.BytesIO(contents), data_only=True, read_only=True)
            chunks = []
            for ws in wb.worksheets:
                chunks.append(f"--- Sheet: {ws.title} ---")
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(c.strip() for c in cells):
                        chunks.append(" | ".join(cells))
            return "\n".join(chunks)
        except Exception as e:
            return f"[xlsx parse error: {e}]"

    # Audio formats — transcribe via OpenAI Whisper through the user's key
    elif ext in ['.mp3', '.wav', '.m4a', '.aac', '.ogg', '.flac', '.webm', '.mp4']:
        try:
            openai_key = os.environ.get('OPENAI_API_KEY')
            if not openai_key:
                return "[Audio file: OPENAI_API_KEY not configured for Whisper]"
            from openai import OpenAI
            client = OpenAI(api_key=openai_key)
            buf = io.BytesIO(contents)
            buf.name = file_path.name
            resp = client.audio.transcriptions.create(
                model="whisper-1", file=buf, response_format="verbose_json",
            )
            return (getattr(resp, "text", "") or "").strip() or "[Whisper returned empty transcript]"
        except Exception as e:
            logger.warning(f"Whisper failed on {file_path.name}: {e}")
            return f"[Whisper error: {e}]"

    # Last-resort fallback: try to decode the raw bytes as UTF-8 text
    try:
        decoded = contents.decode('utf-8', errors='strict')
        if decoded.strip():
            return decoded.strip()
    except Exception:
        pass
    return f"[Unsupported format: {ext}]"


@router.post("/generate-triple-fusion/{case_id}")
async def generate_triple_fusion(
    case_id: str,
    force: bool = False,
    officer: dict = Depends(get_current_officer)
):
    """
    Kick off Triple Fusion generation (Charge Sheet + Case Diary + Remand CD).

    Uses a DB-backed async job queue:
    - Returns immediately with status="processing" and a job_id
    - Background task persists progress to `triple_fusion_jobs` collection
    - Frontend polls GET /api/staging/job-status/{case_id}

    COST: Credits are deducted inside the background task only on success.
          Failed jobs deduct 0 credits (rollback-safe).

    Pass `?force=true` to bypass the result cache and re-run a fresh
    OpenAI pipeline (full 5-credit charge applies).
    """
    officer_id = officer.get("officer_id", "unknown")
    folder = get_case_folder(officer_id, case_id)

    if not folder.exists():
        raise HTTPException(status_code=404, detail="Case folder not found")

    metadata_path = folder / "metadata.json"
    if not metadata_path.exists():
        raise HTTPException(status_code=404, detail="Case metadata not found")

    with open(metadata_path) as f:
        metadata = json.load(f)

    if not metadata.get("files"):
        raise HTTPException(status_code=400, detail="No files staged for fusion")

    # If forcing a fresh run, purge ALL cached fusion rows + any stale extraction
    # cache entries for the files in this case BEFORE the credit check.
    if force and db is not None:
        await db.triple_fusions.delete_many({"case_id": case_id, "officer_id": officer_id})
        await db.triple_fusion_jobs.delete_many({"case_id": case_id, "officer_id": officer_id})
        # Best-effort: also evict document_cache entries whose key matches any
        # ocr_text we have on disk (used by legal_llm for entity extraction).
        try:
            from services.document_cache import generate_cache_key
            keys = []
            for f_meta in metadata.get("files", []):
                txt = (f_meta.get("ocr_text") or f_meta.get("text") or "").strip()
                if txt:
                    keys.append(generate_cache_key(txt, "entity_extraction"))
                    keys.append(generate_cache_key(txt, "translation"))
            if keys:
                deleted = await db.document_cache.delete_many({"cache_key": {"$in": keys}})
                logger.info(f"[force=true] evicted {deleted.deleted_count} document_cache entries")
        except Exception as e:
            logger.warning(f"force-purge: document_cache evict skipped: {e}")
        logger.info(f"[force=true] Purged cache for case {case_id}")

    # Credit balance pre-check (idempotent: cached fusions cost 0 and skip this)
    if db is not None:
        cached = await db.triple_fusions.find_one(
            {"case_id": case_id, "officer_id": officer_id, "status": "completed"},
            {"_id": 0, "status": 1}
        )
        if not cached:
            officer_doc = await db.officers.find_one({"officer_id": officer_id}, {"_id": 0, "credits": 1})
            current = int((officer_doc or {}).get("credits", 0) or 0)
            if current < 5:
                raise HTTPException(
                    status_code=402,
                    detail=f"Insufficient credits — Triple Fusion costs 5 credits, you have {current}. Buy more at /credits.",
                )

    # Fast path: return cached fusion if already completed (skipped if force=true above)
    if db is not None:
        existing_fusion = await db.triple_fusions.find_one(
            {"case_id": case_id, "officer_id": officer_id},
            {"_id": 0}
        )
        if existing_fusion and existing_fusion.get("status") == "completed":
            logger.info(f"Returning cached fusion for case {case_id}")
            return {
                "success": True,
                "status": "completed",
                "transaction_id": existing_fusion.get("transaction_id"),
                "case_id": case_id,
                "documents_processed": existing_fusion.get("documents_processed", 0),
                "credits_used": 0,
                "extracted_data": existing_fusion.get("extracted_data", {}),
                "documents": {
                    "charge_sheet": existing_fusion.get("charge_sheet_html", ""),
                    "case_diary": existing_fusion.get("case_diary_html", ""),
                    "remand_cd": existing_fusion.get("remand_cd_html", "")
                },
                "message": "Triple Fusion retrieved from cache (0 credits used)"
            }

    # Idempotency: if a job is already processing, return current state
    if db is not None:
        existing_job = await db.triple_fusion_jobs.find_one(
            {"case_id": case_id, "officer_id": officer_id},
            {"_id": 0}
        )
        if existing_job and existing_job.get("status") == "processing":
            return {
                "success": True,
                "status": "processing",
                "job_id": existing_job.get("job_id"),
                "case_id": case_id,
                "progress": existing_job.get("progress", 0),
                "stage": existing_job.get("stage", "queued"),
                "message": "Job already in progress. Poll /job-status/{case_id} for updates."
            }

    # Create a new job and fire off background task
    job_id = f"JOB-{datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:6]}"
    now_iso = datetime.now(timezone.utc).isoformat()

    job_doc = {
        "job_id": job_id,
        "case_id": case_id,
        "officer_id": officer_id,
        "status": "processing",
        "progress": 0,
        "stage": "queued",
        "created_at": now_iso,
        "updated_at": now_iso,
        "file_count": len(metadata.get("files", [])),
        "error": None,
    }

    if db is not None:
        await db.triple_fusion_jobs.update_one(
            {"case_id": case_id, "officer_id": officer_id},
            {"$set": job_doc},
            upsert=True
        )

    # Detach background processing (non-blocking)
    asyncio.create_task(
        _process_triple_fusion_background(
            case_id=case_id,
            officer=officer,
            metadata=metadata,
            folder=folder,
            job_id=job_id,
        )
    )

    logger.info(f"Triple Fusion job {job_id} queued for case {case_id} ({len(metadata.get('files', []))} files)")

    return {
        "success": True,
        "status": "processing",
        "job_id": job_id,
        "case_id": case_id,
        "progress": 0,
        "stage": "queued",
        "file_count": len(metadata.get("files", [])),
        "message": "Triple Fusion started. Poll /api/staging/job-status/{case_id} for updates."
    }


@router.get("/fusion/{case_id}")
async def get_fusion_result(
    case_id: str,
    officer: dict = Depends(get_current_officer)
):
    """
    Get previously generated Triple Fusion result.
    COST: 0 CREDITS
    """
    if db is not None:
        result = await db.triple_fusions.find_one(
            {"case_id": case_id, "officer_id": officer.get("officer_id")},
            {"_id": 0}
        )
        
        if result:
            return {"success": True, "fusion": result, "credits_used": 0}
    
    raise HTTPException(status_code=404, detail="Fusion result not found")


@router.get("/download/{case_id}/{doc_type}")
async def download_fusion_document(
    case_id: str,
    doc_type: str,
    officer: dict = Depends(get_current_officer)
):
    """
    Download generated Word document.
    doc_type: 'chargesheet', 'casediary', 'remand'
    """
    from fastapi.responses import FileResponse
    
    folder = get_case_folder(officer.get("officer_id", "unknown"), case_id)
    
    file_map = {
        "chargesheet": f"{case_id}_ChargeSheet.docx",
        "casediary": f"{case_id}_CaseDiary.docx",
        "remand": f"{case_id}_RemandCD.docx"
    }
    
    filename = file_map.get(doc_type)
    if not filename:
        raise HTTPException(status_code=400, detail="Invalid document type")
    
    file_path = folder / filename
    if not file_path.exists():
        # Try global staging folder
        file_path = STAGING_BASE / filename
    
    if file_path.exists():
        return FileResponse(
            path=str(file_path),
            filename=filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    raise HTTPException(status_code=404, detail="Document not found")



async def _update_job(case_id: str, officer_id: str, **fields):
    """Persist job progress/status to `triple_fusion_jobs` collection."""
    if db is None:
        return
    fields["updated_at"] = datetime.now(timezone.utc).isoformat()
    try:
        await db.triple_fusion_jobs.update_one(
            {"case_id": case_id, "officer_id": officer_id},
            {"$set": fields},
            upsert=True
        )
    except Exception as e:
        logger.error(f"Failed to update job status in DB: {e}")


async def _process_triple_fusion_background(
    case_id: str,
    officer: dict,
    metadata: dict,
    folder: Path,
    job_id: str
):
    """
    Background task to process Triple Fusion.

    Persists progress to `triple_fusion_jobs` collection so frontend can
    safely poll `/api/staging/job-status/{case_id}` regardless of which
    backend worker served the request.
    """
    from services.template_generator import (
        generate_html_table_charge_sheet,
        generate_case_diary_html,
    )
    from services.remand_generator import generate_remand_case_diary_html

    officer_id = officer.get("officer_id", "unknown")
    transaction_id = f"TXN-{datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:4]}"
    credits_to_deduct = 5

    try:
        logger.info(f"[BACKGROUND] Job {job_id} starting for case {case_id} ({len(metadata.get('files', []))} files)")
        await _update_job(case_id, officer_id, status="processing", progress=10, stage="extracting_text", job_id=job_id)

        # Collect valid file paths
        file_paths = []
        for file_info in metadata.get("files", []):
            fp = folder / file_info["saved_name"]
            if fp.exists():
                file_paths.append(fp)

        if not file_paths:
            raise RuntimeError("No valid files found in staging folder")

        # Extract text from each staged file in parallel (async but chunked for memory safety)
        extracted_texts = []
        total = len(file_paths)
        for idx, fp in enumerate(file_paths, start=1):
            try:
                txt = await extract_text_from_staged_file(fp)
                if txt:
                    extracted_texts.append(txt)
            except Exception as ex:
                logger.warning(f"Text extraction failed for {fp.name}: {ex}")
            # progress 10% -> 50% across files
            progress = 10 + int(40 * idx / max(total, 1))
            await _update_job(case_id, officer_id, progress=progress, stage=f"extracting_text ({idx}/{total})")

        combined_text = "\n\n".join(extracted_texts)[:20000]  # cap for memory

        # Build case_info dict (ALWAYS a dict - this is the fix for
        # `'str' object has no attribute 'get'` bug in generators)
        case_info = {
            "police_station": metadata.get("police_station", ""),
            "district": metadata.get("district", ""),
            "fir_number": metadata.get("fir_number", ""),
            "sections": metadata.get("sections", ""),
            "io_name": officer.get("name", ""),
            "io_rank": officer.get("rank", "Sub Inspector of Police"),
        }

        await _update_job(case_id, officer_id, progress=55, stage="parsing_entities")

        # Attempt structured parsing via enhanced_legal_parser
        extracted_data = {
            "complainant": {},
            "accused_persons": [],
            "witnesses": [],
            "fir_number": case_info["fir_number"],
            "police_station": case_info["police_station"],
            "district": case_info["district"],
            "sections": case_info["sections"].split(",") if case_info["sections"] else [],
            "act_type": "BNS",
            "io_name": case_info["io_name"],
            "io_rank": case_info["io_rank"],
            "incident_date": "",
            "incident_time": "",
            "incident_place": case_info["police_station"],
            "brief_facts": combined_text[:3000] if combined_text else "Case details to be extracted from uploaded documents.",
        }

        try:
            from services.enhanced_legal_parser import EnhancedLegalParser
            parser = EnhancedLegalParser()
            parsed = parser.parse(combined_text) if combined_text else None
            if parsed:
                extracted_data["complainant"] = parsed.get("complainant", {}) or {}
                extracted_data["accused_persons"] = parsed.get("accused", []) or []
                extracted_data["witnesses"] = parsed.get("witnesses", []) or []
        except Exception as ex:
            logger.warning(f"Parser fallback - using metadata-only extraction: {ex}")

        await _update_job(case_id, officer_id, progress=75, stage="generating_documents")

        # Generate HTML documents (pass case_info DICT, not fir_number string!)
        charge_sheet_html = generate_html_table_charge_sheet(extracted_data, case_info)
        case_diary_html = generate_case_diary_html(extracted_data, case_info)
        remand_cd_html = generate_remand_case_diary_html(extracted_data, case_info)

        await _update_job(case_id, officer_id, progress=90, stage="persisting_result")

        # Persist fusion result + deduct credits
        if db is not None:
            fusion_result = {
                "case_id": case_id,
                "transaction_id": transaction_id,
                "officer_id": officer_id,
                "police_station": metadata.get("police_station", ""),
                "district": metadata.get("district", ""),
                "fir_number": metadata.get("fir_number", ""),
                "documents_processed": len(file_paths),
                "extracted_data": extracted_data,
                "charge_sheet_html": charge_sheet_html,
                "case_diary_html": case_diary_html,
                "remand_cd_html": remand_cd_html,
                "credits_used": credits_to_deduct,
                "created_at": datetime.now(timezone.utc).isoformat(),
                "status": "completed"
            }
            await db.triple_fusions.update_one(
                {"case_id": case_id, "officer_id": officer_id},
                {"$set": fusion_result},
                upsert=True
            )
            await db.officers.update_one(
                {"officer_id": officer_id},
                {"$inc": {"credits": -credits_to_deduct}}
            )
            await db.action_logs.insert_one({
                "officer_id": officer_id,
                "action": "TRIPLE_FUSION_GENERATE",
                "credit_cost": credits_to_deduct,
                "status": "SUCCESS",
                "correlation_id": transaction_id,
                "timestamp": datetime.now(timezone.utc).isoformat(),
                "case_id": case_id,
                "files_processed": len(file_paths),
            })

        # Update staging metadata
        metadata["status"] = "completed"
        metadata["fusion_transaction_id"] = transaction_id
        with open(folder / "metadata.json", "w") as f:
            json.dump(metadata, f, indent=2)

        # Finalize job record with full result so polling can fetch it
        final_result = {
            "success": True,
            "transaction_id": transaction_id,
            "case_id": case_id,
            "documents_processed": len(file_paths),
            "credits_used": credits_to_deduct,
            "documents": {
                "charge_sheet": charge_sheet_html,
                "case_diary": case_diary_html,
                "remand_cd": remand_cd_html
            },
            "extracted_data": extracted_data,
            "message": "Triple Fusion completed successfully!"
        }
        await _update_job(
            case_id, officer_id,
            status="completed", progress=100, stage="done",
            transaction_id=transaction_id, result=final_result,
            error=None,
        )
        logger.info(f"[BACKGROUND] Job {job_id} completed for case {case_id}")

    except Exception as e:
        logger.exception(f"[BACKGROUND] Job {job_id} FAILED for case {case_id}: {e}")
        if db is not None:
            try:
                await db.action_logs.insert_one({
                    "officer_id": officer_id,
                    "action": "TRIPLE_FUSION_GENERATE",
                    "credit_cost": 0,
                    "status": "FAILED",
                    "correlation_id": f"ERR-{transaction_id}",
                    "timestamp": datetime.now(timezone.utc).isoformat(),
                    "case_id": case_id,
                    "error": str(e)
                })
            except Exception:
                pass
        await _update_job(
            case_id, officer_id,
            status="failed", progress=0, stage="failed",
            error=str(e), result=None,
        )


@router.get("/job-status/{case_id}")
async def get_job_status(
    case_id: str,
    officer: dict = Depends(get_current_officer)
):
    """
    Check the status of the Triple Fusion job for this case.
    Reads from the `triple_fusion_jobs` collection (DB-backed queue).
    """
    officer_id = officer.get("officer_id", "unknown")

    if db is None:
        raise HTTPException(status_code=500, detail="Database not initialized")

    job = await db.triple_fusion_jobs.find_one(
        {"case_id": case_id, "officer_id": officer_id},
        {"_id": 0}
    )

    if not job:
        # No job ever queued — maybe fusion was completed earlier
        existing = await db.triple_fusions.find_one(
            {"case_id": case_id, "officer_id": officer_id},
            {"_id": 0}
        )
        if existing and existing.get("status") == "completed":
            return {
                "success": True,
                "status": "completed",
                "case_id": case_id,
                "message": "Triple Fusion already completed. POST /generate-triple-fusion to view cached result.",
                "cached": True
            }
        return {
            "success": False,
            "status": "not_found",
            "case_id": case_id,
            "message": "No processing job found for this case"
        }

    status = job.get("status", "processing")

    if status == "completed":
        return {
            "success": True,
            "status": "completed",
            "case_id": case_id,
            "job_id": job.get("job_id"),
            "progress": 100,
            "stage": "done",
            **(job.get("result") or {})
        }

    if status == "failed":
        return {
            "success": False,
            "status": "failed",
            "case_id": case_id,
            "job_id": job.get("job_id"),
            "progress": job.get("progress", 0),
            "stage": job.get("stage", "failed"),
            "error": job.get("error", "Unknown error"),
            "message": f"Triple Fusion failed. NO CREDITS DEDUCTED. Error: {job.get('error', 'Unknown')}"
        }

    # processing / queued
    return {
        "success": True,
        "status": status,
        "case_id": case_id,
        "job_id": job.get("job_id"),
        "progress": job.get("progress", 0),
        "stage": job.get("stage", "queued"),
        "message": f"Processing... {job.get('progress', 0)}% ({job.get('stage', 'queued')})"
    }



# =====================================================================
# Intelligent Charge Sheet Generator (station-writer grade)
# =====================================================================
# Schema adapter: LLM JSON → fixed_layout_renderer schema
# =====================================================================
NOT_FOUND = "NOT FOUND IN DOCUMENTS"


def _nf(value, default=NOT_FOUND):
    """Return value or 'NOT FOUND IN DOCUMENTS' instead of an empty string.

    Phase-1 (manual) values that come in as empty are rare; they should
    still be flagged so the officer sees the gap.
    """
    if value is None:
        return default
    s = str(value).strip()
    if not s or s == "_____":
        return default
    return s


def _adapt_person(p: dict) -> dict:
    """Translate LLM person schema → fixed_layout_renderer schema."""
    if not isinstance(p, dict):
        return {}
    # Infer marital status from salutation when possible (Smt. → married, Kum. → unmarried)
    salutation = (p.get("salutation") or "").strip()
    marital = p.get("marital_status") or ""
    if not marital:
        if salutation.lower().startswith("smt"):
            marital = "married"
        elif salutation.lower().startswith("kum"):
            marital = "unmarried"
    # Infer gender from salutation if not provided
    gender = (p.get("gender") or "").lower()
    if not gender:
        if salutation.lower().startswith(("smt", "kum")):
            gender = "female"
        elif salutation.lower().startswith("sri"):
            gender = "male"
    return {
        "salutation": salutation,
        "name": _nf(p.get("name"), ""),
        "father": _nf(p.get("father_name") or p.get("father"), ""),
        "age": _nf(p.get("age"), ""),
        "caste": _nf(p.get("caste"), ""),
        "occupation": _nf(p.get("occupation") or p.get("occ"), ""),
        "address": _nf(p.get("address") or p.get("permanent_address"), ""),
        "phone": _nf(p.get("phone"), ""),
        "gender": gender,
        "marital_status": marital,
        "relation": p.get("relation") or "",
        "aadhaar_number": p.get("aadhaar_number") or "",
        "type": p.get("role") or p.get("type") or "",
        "role": p.get("role") or p.get("type") or "",
        "alias": p.get("alias") or "",
    }


def _adapt_llm_schema_to_fixed_layout(cs: dict) -> dict:
    """
    Translate the V3.0 intelligent_charge_sheet LLM JSON output into the
    schema expected by services.fixed_layout_renderer.render_charge_sheet
    so the AUTHENTIC 18-section Telangana layout is produced.

    "Not found" policy: blank cells render as 'NOT FOUND IN DOCUMENTS'
    (never '_____').
    """
    if not isinstance(cs, dict):
        return {}
    court = (cs.get("court") or "").strip()
    court_name, court_place = "", ""
    if " AT " in court.upper():
        idx = court.upper().rfind(" AT ")
        court_name = court[:idx].replace("IN THE COURT OF", "").strip()
        court_place = court[idx + 4:].strip().rstrip(",.").strip()
    brief = cs.get("brief_facts") or ""
    if isinstance(brief, list):
        brief_paragraphs = [str(x).strip() for x in brief if x]
    elif isinstance(brief, str):
        brief_paragraphs = [p.strip() for p in brief.split("\n\n") if p.strip()]
    else:
        brief_paragraphs = []
    prayer = (cs.get("prayer") or "").strip()
    if prayer and (not brief_paragraphs or prayer not in brief_paragraphs[-1]):
        brief_paragraphs.append(prayer)

    io_d = cs.get("io") or {}
    return {
        "court_name": court_name or "JUDICIAL FIRST CLASS MAGISTRATE",
        "court_place": court_place or _nf(cs.get("district"), "NOT FOUND IN DOCUMENTS"),
        "police_station": _nf(cs.get("police_station")),
        "district": _nf(cs.get("district")),
        "fir_number": _nf(cs.get("fir_number")),
        "fir_date": _nf(cs.get("fir_date")),
        "charge_sheet_no": _nf(cs.get("chargesheet_no")),
        "today_date": _nf(cs.get("chargesheet_date")),
        "sections": _nf(cs.get("sections")),
        "final_report_type": _nf(cs.get("report_type"), "Charge Sheet."),
        "fr_unoccurred": _nf(cs.get("un_occurred_reason"), "----"),
        "charge_sheet_kind": _nf(cs.get("chargesheet_type"), "Original."),
        "io": {
            "salutation": io_d.get("salutation") or "Sri.",
            "name": _nf(io_d.get("name"), "NOT FOUND IN DOCUMENTS"),
            "designation": _nf(io_d.get("rank") or io_d.get("designation"),
                                "Sub Inspector of Police"),
            "rank": _nf(io_d.get("rank") or io_d.get("designation"),
                         "Sub Inspector of Police"),
        },
        "complainant": _adapt_person(cs.get("complainant") or {}),
        "accused": [_adapt_person(a) for a in (cs.get("accused") or [])],
        "witnesses": [_adapt_person(w) for w in (cs.get("witnesses") or [])],
        "properties_seized": _nf(cs.get("property_recovered"), "---"),
        "ack_notice_enclosed": _nf(cs.get("notice_ack_enclosed"), "No."),
        "dispatch_date": _nf(cs.get("dispatch_date")),
        "brief_facts_paragraphs": brief_paragraphs,
        "arrest_release": _nf(cs.get("arrest_release"), "--"),
        "sureties": _nf(cs.get("sureties"), "--"),
        "previous_convictions": _nf(cs.get("previous_convictions"), "--"),
        "absconding": _nf(cs.get("absconding"), "--"),
        "accused_not_chargesheeted": _nf(cs.get("accused_not_chargesheeted"), "Nil"),
        "fr_false_action": _nf(cs.get("fr_false_action"), "--Nil--"),
        "lab_result": _nf(cs.get("lab_result"), "--Nil--"),
        # Pass-through fields used by the response (not rendered)
        "_extraction_report": cs.get("extraction_report") or {},
        "_corrections_applied": cs.get("corrections_applied") or [],
    }


# =====================================================================
# INTELLIGENT CHARGE SHEET — ASYNC background-job pattern
#
# WHY: the synchronous version was hitting the K8s ingress 60s timeout for
# 23-file cases (LLM + OCR + render runs 40-70s). The user saw "Intelligent
# generation failed" with HTTP 502 even though the backend completed
# successfully — the proxy killed the connection before the DOCX could be
# streamed back.
#
# HOW: POST returns immediately with {status: 'processing'}. The work runs
# in a detached asyncio task that saves the DOCX to disk + updates
# intelligent_chargesheets.{status} on completion. The frontend polls
# GET /intelligent-chargesheet/{case_id} until status='completed' and then
# downloads from GET /intelligent-chargesheet/{case_id}/download.
# =====================================================================
async def _process_icgs_background(*, case_id: str, officer: dict, credits_to_deduct: int):
    """Run the heavy ICGS generation in a detached task — survives K8s 60s ingress."""
    from services.intelligent_charge_sheet import generate_intelligent_charge_sheet
    from services.fixed_layout_renderer import render_charge_sheet as render_authentic_charge_sheet

    officer_id = officer.get("officer_id", "unknown")
    try:
        existing = await db.triple_fusions.find_one(
            {"case_id": case_id, "officer_id": officer_id},
            {"_id": 0}
        )
        if not existing:
            raise RuntimeError("No Triple Fusion result found. Run Generate Triple Fusion first.")

        extracted = existing.get("extracted_data") or {}
        folder = get_case_folder(officer_id, case_id)
        metadata_path = folder / "metadata.json"
        metadata: dict = {}
        if metadata_path.exists():
            with open(metadata_path) as f:
                metadata = json.load(f)

        sections_raw = extracted.get("sections") or []
        sections_text = ", ".join(sections_raw) if isinstance(sections_raw, list) else str(sections_raw)

        files_meta = metadata.get("files") or []
        uploaded_documents: List[str] = []
        documents_corpus_parts: List[str] = []
        for fmeta in files_meta:
            fname = fmeta.get("filename") or fmeta.get("saved_name") or "unknown"
            uploaded_documents.append(fname)
            text = (fmeta.get("ocr_text") or fmeta.get("text") or "").strip()
            if not text or len(text) < 40 or text.startswith("[No text in PDF]") or text.startswith("[PDF error"):
                saved = fmeta.get("saved_name")
                if saved:
                    fp = folder / saved
                    if fp.exists():
                        try:
                            text = (await extract_text_from_staged_file(fp)).strip()
                            fmeta["ocr_text"] = text
                        except Exception as e:
                            logger.warning(f"[ICGS-BG] re-OCR failed for {saved}: {e}")
            if text:
                documents_corpus_parts.append(
                    f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                    f"FILE: {fname}\n"
                    f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                    f"{text[:25000]}"
                )
        try:
            with open(metadata_path, "w") as f:
                json.dump(metadata, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.warning(f"[ICGS-BG] could not persist refreshed metadata: {e}")

        documents_corpus = "\n\n".join(documents_corpus_parts)

        raw_data = {
            "fir_number": extracted.get("fir_number") or metadata.get("fir_number", ""),
            "fir_date": extracted.get("fir_date", ""),
            "chargesheet_date": datetime.now(timezone.utc).strftime("%d.%m.%Y"),
            "police_station": extracted.get("police_station") or metadata.get("police_station", ""),
            "district": extracted.get("district") or metadata.get("district", ""),
            "sections": sections_text or metadata.get("sections", ""),
            "io": {
                "name": extracted.get("io_name") or officer.get("name", ""),
                "rank": extracted.get("io_rank") or officer.get("rank", "SI of Police"),
                "station": f"PS {metadata.get('police_station','Makthal')}",
                "salutation": "Sri.",
            },
            "complainant": extracted.get("complainant") or {},
            "accused_persons": extracted.get("accused_persons") or [],
            "witnesses": extracted.get("witnesses") or [],
            "incident_date": extracted.get("incident_date", ""),
            "incident_time": extracted.get("incident_time", ""),
            "incident_place": extracted.get("incident_place", ""),
            "medical_findings": extracted.get("medical_findings", ""),
            "section_35_3_dates": extracted.get("section_35_3_dates", ""),
            "brief_facts": extracted.get("brief_facts", ""),
            "uploaded_documents": uploaded_documents,
            "documents_corpus": documents_corpus,
        }

        # V3.0 PHASE-1 OVERRIDE (manual_input is authoritative)
        manual = metadata.get("manual_input") or {}
        if manual:
            def _fmt_date(d):
                if not d:
                    return ""
                s = str(d).strip()
                if "-" in s and len(s.split("-")[0]) == 4:
                    y, m, dd = s.split("-")
                    return f"{dd}.{m}.{y}"
                return s.replace("/", ".")

            raw_data.update({
                "district":           manual.get("district") or raw_data["district"],
                "police_station":     manual.get("police_station") or raw_data["police_station"],
                "fir_number":         manual.get("fir_number") or raw_data["fir_number"],
                "fir_date":           _fmt_date(manual.get("fir_date")) or raw_data["fir_date"],
                "chargesheet_no":     manual.get("chargesheet_no", ""),
                "chargesheet_date":   _fmt_date(manual.get("chargesheet_date")) or raw_data["chargesheet_date"],
                "sections":           manual.get("sections") or raw_data["sections"],
                "report_type":        manual.get("report_type") or "Charge Sheet.",
                "un_occurred_reason": manual.get("un_occurred_reason") or "----",
                "chargesheet_type":   manual.get("chargesheet_type") or "Original.",
                "io": {
                    **(raw_data.get("io") or {}),
                    "name":    manual.get("io_name") or (raw_data.get("io") or {}).get("name", ""),
                    "rank":    manual.get("io_rank") or (raw_data.get("io") or {}).get("rank", ""),
                    "station": manual.get("police_station") or (raw_data.get("io") or {}).get("station", ""),
                },
                "court":              manual.get("court_name") and f"IN THE COURT OF {manual['court_name']}" or "",
                "court_name":         manual.get("court_name", ""),
                "dispatch_date":      _fmt_date(manual.get("dispatch_date")),
                "notice_ack_enclosed": (manual.get("ack_enclosed") or "No") + ".",
                "manual_input_locked": True,
            })

        await db.intelligent_chargesheets.update_one(
            {"case_id": case_id, "officer_id": officer_id},
            {"$set": {
                "stage": "llm_composing",
                "progress": 35,
                "updated_at": datetime.now(timezone.utc).isoformat(),
            }},
        )

        logger.info(f"[ICGS-BG] LLM call starting for case {case_id}")
        cs_data = await generate_intelligent_charge_sheet(raw_data, session_id=f"ics-{case_id}")

        # PHASE-1 RE-LOCK
        if manual:
            cs_data["court"] = (
                f"IN THE COURT OF {manual['court_name']}"
                if manual.get("court_name") else cs_data.get("court", "")
            )
            cs_data["district"]            = manual.get("district") or cs_data.get("district", "")
            cs_data["police_station"]      = manual.get("police_station") or cs_data.get("police_station", "")
            cs_data["fir_number"]          = manual.get("fir_number") or cs_data.get("fir_number", "")
            cs_data["fir_date"]            = raw_data.get("fir_date") or cs_data.get("fir_date", "")
            cs_data["chargesheet_no"]      = manual.get("chargesheet_no") or cs_data.get("chargesheet_no", "")
            cs_data["chargesheet_date"]    = raw_data.get("chargesheet_date") or cs_data.get("chargesheet_date", "")
            cs_data["sections"]            = manual.get("sections") or cs_data.get("sections", "")
            cs_data["report_type"]         = manual.get("report_type") or cs_data.get("report_type", "")
            cs_data["chargesheet_type"]    = manual.get("chargesheet_type") or cs_data.get("chargesheet_type", "")
            cs_data["un_occurred_reason"]  = manual.get("un_occurred_reason") or cs_data.get("un_occurred_reason", "")
            io_now = cs_data.get("io") or {}
            io_now["name"]    = manual.get("io_name") or io_now.get("name", "")
            io_now["rank"]    = manual.get("io_rank") or io_now.get("rank", "")
            io_now["station"] = manual.get("police_station") or io_now.get("station", "")
            cs_data["io"] = io_now
            cs_data["dispatch_date"]       = raw_data.get("dispatch_date") or cs_data.get("dispatch_date", "")
            cs_data["notice_ack_enclosed"] = (manual.get("ack_enclosed") or "No") + "."

        await db.intelligent_chargesheets.update_one(
            {"case_id": case_id, "officer_id": officer_id},
            {"$set": {
                "stage": "rendering_docx",
                "progress": 85,
                "updated_at": datetime.now(timezone.utc).isoformat(),
            }},
        )

        adapted = _adapt_llm_schema_to_fixed_layout(cs_data)
        docx_bytes = render_authentic_charge_sheet(adapted)

        # Persist DOCX to the case folder so /download can stream it
        docx_path = folder / "intelligent_charge_sheet.docx"
        docx_path.write_bytes(docx_bytes)

        fir_safe = (cs_data.get("fir_number") or "case").replace("/", "-")
        download_filename = f"{fir_safe}_IntelligentChargeSheet.docx"

        await db.intelligent_chargesheets.update_one(
            {"case_id": case_id, "officer_id": officer_id},
            {"$set": {
                "case_id": case_id,
                "officer_id": officer_id,
                "fir_number": cs_data.get("fir_number", ""),
                "structured_data": cs_data,
                "corrections_applied": cs_data.get("corrections_applied", []),
                "model_used": cs_data.get("_model_used", ""),
                "completed_at": datetime.now(timezone.utc).isoformat(),
                "credits_used": credits_to_deduct,
                "docx_filename": download_filename,
                "docx_relative_path": "intelligent_charge_sheet.docx",
                "status": "completed",
                "stage": "completed",
                "progress": 100,
                "error": None,
            }},
            upsert=True,
        )
        await db.officers.update_one(
            {"officer_id": officer_id},
            {"$inc": {"credits": -credits_to_deduct}}
        )
        await db.action_logs.insert_one({
            "officer_id": officer_id,
            "action": "INTELLIGENT_CHARGESHEET_GENERATE",
            "credit_cost": credits_to_deduct,
            "status": "SUCCESS",
            "correlation_id": case_id,
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "model_used": cs_data.get("_model_used", ""),
        })
        logger.info(f"[ICGS-BG] case {case_id}: completed → {download_filename}")
    except Exception as e:
        logger.exception(f"[ICGS-BG] case {case_id} FAILED: {e}")
        try:
            await db.intelligent_chargesheets.update_one(
                {"case_id": case_id, "officer_id": officer_id},
                {"$set": {
                    "status": "failed",
                    "stage": "failed",
                    "error": str(e),
                    "failed_at": datetime.now(timezone.utc).isoformat(),
                }},
                upsert=True,
            )
            await db.action_logs.insert_one({
                "officer_id": officer_id,
                "action": "INTELLIGENT_CHARGESHEET_GENERATE",
                "credit_cost": 0,
                "status": "FAILED",
                "correlation_id": case_id,
                "timestamp": datetime.now(timezone.utc).isoformat(),
                "error": str(e),
            })
        except Exception:
            pass


@router.post("/generate-intelligent-charge-sheet/{case_id}")
async def generate_intelligent_charge_sheet_endpoint(
    case_id: str,
    officer: dict = Depends(get_current_officer),
):
    """
    Kick off Intelligent Charge Sheet generation as a BACKGROUND job.
    Returns immediately (under 1s) with {status: 'processing'}. The frontend
    polls GET /staging/intelligent-chargesheet/{case_id} until
    status='completed' and then GETs /download to fetch the DOCX.
    COST: 3 credits (deducted only on success in the background task).
    """
    officer_id = officer.get("officer_id", "unknown")
    credits_to_deduct = 3

    if db is None:
        raise HTTPException(status_code=500, detail="Database not initialized")

    # Credit balance pre-check
    officer_doc = await db.officers.find_one({"officer_id": officer_id}, {"_id": 0, "credits": 1})
    current = int((officer_doc or {}).get("credits", 0) or 0)
    if current < credits_to_deduct:
        raise HTTPException(
            status_code=402,
            detail=f"Insufficient credits — Intelligent Charge Sheet costs {credits_to_deduct} credits, you have {current}. Buy more at /credits.",
        )

    # Idempotency: if a previous job is still running, return its status
    existing = await db.intelligent_chargesheets.find_one(
        {"case_id": case_id, "officer_id": officer_id},
        {"_id": 0}
    )
    if existing and existing.get("status") == "processing":
        return {
            "success": True,
            "status": "processing",
            "case_id": case_id,
            "stage": existing.get("stage", "queued"),
            "progress": existing.get("progress", 0),
            "message": "Job already in progress. Poll GET /staging/intelligent-chargesheet/{case_id}.",
        }

    # Pre-check that a triple_fusion exists so we fail fast instead of in background
    tf = await db.triple_fusions.find_one(
        {"case_id": case_id, "officer_id": officer_id}, {"_id": 0, "extracted_data": 1}
    )
    if not tf:
        raise HTTPException(
            status_code=400,
            detail="No Triple Fusion result found. Run Generate Triple Fusion first."
        )

    # Mark processing + spawn background task
    started_at = datetime.now(timezone.utc).isoformat()
    await db.intelligent_chargesheets.update_one(
        {"case_id": case_id, "officer_id": officer_id},
        {"$set": {
            "case_id": case_id, "officer_id": officer_id,
            "status": "processing",
            "stage": "queued",
            "progress": 5,
            "started_at": started_at,
            "updated_at": started_at,
            "error": None,
        }},
        upsert=True,
    )
    asyncio.create_task(
        _process_icgs_background(
            case_id=case_id, officer=officer,
            credits_to_deduct=credits_to_deduct,
        )
    )
    logger.info(f"[ICGS] Started background job for case {case_id} (officer {officer_id})")
    return {
        "success": True,
        "status": "processing",
        "case_id": case_id,
        "stage": "queued",
        "progress": 5,
        "started_at": started_at,
        "message": "Intelligent Charge Sheet started. Poll GET /staging/intelligent-chargesheet/{case_id} every 4-5s; when status='completed' GET /download to fetch the DOCX.",
    }


@router.get("/intelligent-chargesheet/{case_id}/download")
async def download_intelligent_chargesheet(
    case_id: str,
    officer: dict = Depends(get_current_officer),
):
    """Stream the DOCX produced by the background Intelligent Charge Sheet job."""
    from fastapi.responses import FileResponse
    if db is None:
        raise HTTPException(status_code=500, detail="Database not initialized")
    officer_id = officer.get("officer_id", "unknown")
    doc = await db.intelligent_chargesheets.find_one(
        {"case_id": case_id, "officer_id": officer_id},
        {"_id": 0, "status": 1, "docx_relative_path": 1, "docx_filename": 1, "fir_number": 1}
    )
    if not doc:
        raise HTTPException(status_code=404, detail="No Intelligent Charge Sheet generated for this case yet")
    status = doc.get("status") or "completed"  # legacy rows have no status field — assume completed
    if status == "processing":
        raise HTTPException(
            status_code=409,
            detail=f"Still generating (stage={doc.get('stage','?')}). Poll /staging/intelligent-chargesheet/{{case_id}}."
        )
    if status == "failed":
        raise HTTPException(status_code=500, detail=f"Generation failed: {doc.get('error','unknown')}")
    folder = get_case_folder(officer_id, case_id)
    rel = doc.get("docx_relative_path") or "intelligent_charge_sheet.docx"
    fp = folder / rel
    if not fp.exists():
        raise HTTPException(status_code=410, detail="DOCX no longer available — re-run Generate")
    fir_safe = (doc.get("fir_number") or "case").replace("/", "-")
    filename = doc.get("docx_filename") or f"{fir_safe}_IntelligentChargeSheet.docx"
    return FileResponse(
        path=str(fp),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/intelligent-chargesheet/{case_id}")
async def get_intelligent_chargesheet_metadata(
    case_id: str,
    officer: dict = Depends(get_current_officer),
):
    """Return the corrections applied + structured data for the last intelligent charge sheet."""
    if db is None:
        raise HTTPException(status_code=500, detail="Database not initialized")
    doc = await db.intelligent_chargesheets.find_one(
        {"case_id": case_id, "officer_id": officer.get("officer_id", "unknown")},
        {"_id": 0}
    )
    if not doc:
        return {"success": False, "message": "No intelligent charge sheet generated for this case yet"}
    return {"success": True, **doc}


def _build_cctns_autofill(cs: dict) -> dict:
    """
    Build a FLAT JSON block matching the CCTNS Charge Sheet form fields so
    the police writer can copy/paste straight into the CCTNS portal.

    CCTNS form (Crime and Criminal Tracking Network & Systems) expects a
    fixed shape — this helper flattens the V3.0 chargesheet structured_data
    into that shape with verbatim values (never re-derived, never AI'd).
    """
    if not isinstance(cs, dict):
        cs = {}

    def _safe(v, default=""):
        if v is None:
            return default
        s = str(v).strip()
        return s if s and s != "NOT FOUND IN DOCUMENTS" else default

    def _person_flat(p, prefix):
        p = p or {}
        return {
            f"{prefix}_salutation":   _safe(p.get("salutation")),
            f"{prefix}_name":         _safe(p.get("name")),
            f"{prefix}_relation":     _safe(p.get("relation")),
            f"{prefix}_father_name":  _safe(p.get("father_name") or p.get("father")),
            f"{prefix}_gender":       _safe(p.get("gender")),
            f"{prefix}_marital":      _safe(p.get("marital_status")),
            f"{prefix}_age":          _safe(p.get("age")),
            f"{prefix}_caste":        _safe(p.get("caste")),
            f"{prefix}_occupation":   _safe(p.get("occupation") or p.get("occ")),
            f"{prefix}_address":      _safe(p.get("address") or p.get("permanent_address")),
            f"{prefix}_phone":        _safe(p.get("phone")),
            f"{prefix}_aadhaar":      _safe(p.get("aadhaar_number")),
        }

    io_d = cs.get("io") or {}
    sections_raw = _safe(cs.get("sections"))
    sections_list = [s.strip() for s in re.split(r"[,;]", sections_raw) if s.strip()] if sections_raw else []

    accused = cs.get("accused") or []
    witnesses = cs.get("witnesses") or []

    flat: dict = {
        # Case header
        "fir_number":            _safe(cs.get("fir_number")),
        "fir_date":              _safe(cs.get("fir_date")),
        "chargesheet_no":        _safe(cs.get("chargesheet_no")),
        "chargesheet_date":      _safe(cs.get("chargesheet_date")),
        "district":              _safe(cs.get("district")),
        "police_station":        _safe(cs.get("police_station")),
        "state":                 "Telangana",
        "court_name":            _safe(cs.get("court")),
        # Offence
        "sections":              sections_raw,
        "sections_list":         sections_list,
        "report_type":           _safe(cs.get("report_type"), "Charge Sheet"),
        "chargesheet_type":      _safe(cs.get("chargesheet_type"), "Original"),
        "fr_unoccurred_reason":  _safe(cs.get("un_occurred_reason"), "----"),
        # Investigating Officer
        "io_salutation":         _safe(io_d.get("salutation"), "Sri."),
        "io_name":               _safe(io_d.get("name")),
        "io_rank":               _safe(io_d.get("rank") or io_d.get("designation")),
        "io_station":            _safe(io_d.get("station")),
        # Counts
        "total_accused":         len(accused),
        "total_witnesses":       len(witnesses),
        # Property
        "property_seized":       _safe(cs.get("property_recovered"), "---"),
        # Notice / arrest
        "notice_35_3":           _safe(cs.get("arrest_release"), "--"),
        "sureties":              _safe(cs.get("sureties"), "--"),
        "previous_convictions":  _safe(cs.get("previous_convictions"), "--"),
        "absconding":            _safe(cs.get("absconding"), "--"),
        "accused_not_chargesheeted": _safe(cs.get("accused_not_chargesheeted"), "Nil"),
        # FR / Lab / Brief facts
        "fr_false_action":       _safe(cs.get("fr_false_action"), "--Nil--"),
        "lab_result":            _safe(cs.get("lab_result"), "--Nil--"),
        "brief_facts":           _safe(cs.get("brief_facts"))[:8000],
        "prayer":                _safe(cs.get("prayer")),
        "notice_ack_enclosed":   _safe(cs.get("notice_ack_enclosed"), "No."),
        "dispatch_date":         _safe(cs.get("dispatch_date")),
    }
    # Flat complainant block
    flat.update(_person_flat(cs.get("complainant"), "complainant"))
    # Flat A1..AN blocks
    for i, a in enumerate(accused, 1):
        flat.update(_person_flat(a, f"a{i}"))
        flat[f"a{i}_alias"] = (a or {}).get("alias", "")
        flat[f"a{i}_notice_35_3_date"] = (a or {}).get("section_35_3_notice_date", "")
    # Flat LW-1..LW-N blocks (with role)
    for i, w in enumerate(witnesses, 1):
        flat.update(_person_flat(w, f"lw{i}"))
        flat[f"lw{i}_role"] = (w or {}).get("role") or (w or {}).get("type") or ""

    # Generated-at metadata
    flat["_cctns_schema_version"] = "1.0"
    flat["_generated_at"] = datetime.now(timezone.utc).isoformat()
    return flat


@router.get("/cctns-autofill/{case_id}")
async def get_cctns_autofill(
    case_id: str,
    officer: dict = Depends(get_current_officer),
):
    """
    Return a FLAT JSON object mapping the case's intelligent charge sheet
    onto the CCTNS (Crime & Criminal Tracking Network) form fields. The
    police writer can copy-paste the values into the CCTNS portal without
    re-deriving anything. 0 credits — pure data view.
    """
    if db is None:
        raise HTTPException(status_code=500, detail="Database not initialized")
    officer_id = officer.get("officer_id", "unknown")
    ics = await db.intelligent_chargesheets.find_one(
        {"case_id": case_id, "officer_id": officer_id}, {"_id": 0}
    )
    if not ics or not ics.get("structured_data"):
        return {
            "success": False,
            "message": "Generate the Intelligent Charge Sheet first — CCTNS autofill is built from its structured data.",
        }
    return {
        "success": True,
        "case_id": case_id,
        "fir_number": ics.get("fir_number", ""),
        "cctns_autofill": _build_cctns_autofill(ics["structured_data"]),
    }


# =====================================================================
# REGENERATE WITH CORRECTIONS — Section G of V3.0 spec
# Reads the last-generated chargesheet JSON, applies user-supplied
# corrections via the LLM cascade rules, and emits a fresh DOCX.
# Cost: 0 credits (it's a re-render, the user already paid for v1).
# =====================================================================
class _CorrectionItem(BaseModel):
    field: str   # e.g., "Field 08 IO Name"
    instruction: str  # e.g., "Change IO name from K Lal to K. Lal Singh"


class _RegenerateRequest(BaseModel):
    corrections: List[_CorrectionItem]


@router.post("/regenerate-charge-sheet/{case_id}")
async def regenerate_charge_sheet(
    case_id: str,
    body: _RegenerateRequest,
    officer: dict = Depends(get_current_officer),
):
    """
    Re-run the intelligent charge sheet with user-supplied corrections.
    The LLM receives the previously-generated JSON + a CORRECTIONS block
    listing the cascade rules from the V3.0 spec, then regenerates a
    complete fresh chargesheet with all dependent fields updated.

    Cost: 0 credits (re-render of an already-paid run).
    """
    officer_id = officer.get("officer_id", "unknown")
    folder = get_case_folder(officer_id, case_id)
    if not folder.exists():
        raise HTTPException(status_code=404, detail="Case folder not found")
    if not body.corrections:
        raise HTTPException(status_code=400, detail="Provide at least one correction")
    if db is None:
        raise HTTPException(status_code=500, detail="Database not initialized")

    # Load the previously-generated chargesheet JSON (sole source of truth for v2)
    prev = await db.intelligent_chargesheets.find_one(
        {"case_id": case_id, "officer_id": officer_id}, {"_id": 0}
    )
    if not prev:
        raise HTTPException(
            status_code=400,
            detail="No previous intelligent charge sheet found. Run Generate Station-Format Charge Sheet first.",
        )

    prev_payload = prev.get("structured_data") or {}

    # Re-derive the documents corpus so the LLM still has the ground truth
    # available alongside the previous payload + corrections.
    metadata_path = folder / "metadata.json"
    if not metadata_path.exists():
        raise HTTPException(status_code=404, detail="Case metadata missing")
    with open(metadata_path) as f:
        metadata = json.load(f)

    files_meta = metadata.get("files") or []
    uploaded_documents = [
        f.get("filename") or f.get("saved_name") or "?" for f in files_meta
    ]
    corpus_parts = []
    for fmeta in files_meta:
        text = (fmeta.get("ocr_text") or fmeta.get("text") or "").strip()
        if not text:
            saved = fmeta.get("saved_name")
            if saved:
                fp = folder / saved
                if fp.exists():
                    try:
                        text = (await extract_text_from_staged_file(fp)).strip()
                    except Exception:
                        text = ""
        if text:
            corpus_parts.append(
                f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                f"FILE: {fmeta.get('filename') or fmeta.get('saved_name')}\n"
                f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                f"{text[:25000]}"
            )

    # Build a raw_data payload that includes BOTH the previous JSON and the
    # user corrections — the LLM's prompt builder will append the
    # corrections + cascade rules block automatically.
    # Regenerate is FAST-PATH: the LLM has the previous_payload as ground truth,
    # so we trim the documents_corpus to keep the call under the 60s ingress.
    manual = metadata.get("manual_input") or {}
    trimmed_corpus = "\n\n".join(corpus_parts)
    if trimmed_corpus:
        trimmed_corpus = trimmed_corpus[:8000]
    raw_data = {
        **prev_payload,
        "uploaded_documents": uploaded_documents,
        "documents_corpus": trimmed_corpus,
        "corrections": [c.dict() for c in body.corrections],
        "previous_payload": prev_payload,
    }
    # Re-apply manual lock (so corrections don't drift to use stale defaults)
    if manual:
        raw_data["manual_input_locked"] = True

    try:
        logger.info(
            f"[ICGS-REGEN] case {case_id}: applying {len(body.corrections)} "
            f"correction(s) via {os.environ.get('OPENAI_DEFAULT_MODEL', 'gpt-4o')}"
        )
        from services.intelligent_charge_sheet import generate_intelligent_charge_sheet
        from services.fixed_layout_renderer import render_charge_sheet as render_authentic_charge_sheet
        from fastapi.responses import Response
        cs_data = await generate_intelligent_charge_sheet(
            raw_data, session_id=f"ics-regen-{case_id}-{uuid.uuid4().hex[:6]}"
        )

        # V3.0 PHASE-1 RE-LOCK (same as the original endpoint)
        if manual:
            def _fmt_date(d):
                if not d:
                    return ""
                s = str(d).strip()
                if "-" in s and len(s.split("-")[0]) == 4:
                    y, m, dd = s.split("-")
                    return f"{dd}.{m}.{y}"
                return s.replace("/", ".")
            cs_data["court"] = (
                f"IN THE COURT OF {manual['court_name']}" if manual.get("court_name")
                else cs_data.get("court", "")
            )
            cs_data["district"]         = manual.get("district") or cs_data.get("district", "")
            cs_data["police_station"]   = manual.get("police_station") or cs_data.get("police_station", "")
            # When the user explicitly corrects the FIR / sections / IO, the LLM's
            # corrections_applied will reflect the new value. Otherwise re-lock to
            # the manual input.
            user_corrected_fields = " ".join(
                (c.field + " " + c.instruction).lower() for c in body.corrections
            )
            if "fir" not in user_corrected_fields:
                cs_data["fir_number"] = manual.get("fir_number") or cs_data.get("fir_number", "")
                cs_data["fir_date"]   = _fmt_date(manual.get("fir_date")) or cs_data.get("fir_date", "")
            if "section" not in user_corrected_fields and "sections" not in user_corrected_fields:
                cs_data["sections"]   = manual.get("sections") or cs_data.get("sections", "")
            if "io" not in user_corrected_fields:
                io_now = cs_data.get("io") or {}
                io_now["name"]    = manual.get("io_name") or io_now.get("name", "")
                io_now["rank"]    = manual.get("io_rank") or io_now.get("rank", "")
                io_now["station"] = manual.get("police_station") or io_now.get("station", "")
                cs_data["io"] = io_now
            if "dispatch" not in user_corrected_fields:
                cs_data["dispatch_date"] = _fmt_date(manual.get("dispatch_date")) or cs_data.get("dispatch_date", "")
            if "ack" not in user_corrected_fields:
                cs_data["notice_ack_enclosed"] = (manual.get("ack_enclosed") or "No") + "."

        adapted = _adapt_llm_schema_to_fixed_layout(cs_data)
        docx_bytes = render_authentic_charge_sheet(adapted)

        # Persist the new version (overwrites the previous one — corrections are cumulative)
        await db.intelligent_chargesheets.update_one(
            {"case_id": case_id, "officer_id": officer_id},
            {"$set": {
                "case_id": case_id, "officer_id": officer_id,
                "fir_number": cs_data.get("fir_number", ""),
                "structured_data": cs_data,
                "corrections_applied": cs_data.get("corrections_applied", []),
                "regeneration_count": (prev.get("regeneration_count", 0) or 0) + 1,
                "last_corrections": [c.dict() for c in body.corrections],
                "model_used": cs_data.get("_model_used", ""),
                "regenerated_at": datetime.now(timezone.utc).isoformat(),
            }},
            upsert=True,
        )
        await db.action_logs.insert_one({
            "officer_id": officer_id,
            "action": "INTELLIGENT_CHARGESHEET_REGENERATE",
            "credit_cost": 0,
            "status": "SUCCESS",
            "correlation_id": case_id,
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "metadata": {"correction_count": len(body.corrections)},
        })

        fir_safe = (cs_data.get("fir_number") or "case").replace("/", "-")
        rev = (prev.get("regeneration_count", 0) or 0) + 1
        filename = f"{fir_safe}_IntelligentChargeSheet_rev{rev}.docx"
        import json as _json
        cascade_header = _json.dumps({
            "corrections_applied": cs_data.get("corrections_applied", []),
            "regeneration_count": rev,
            "user_corrections": [c.dict() for c in body.corrections],
        })[:6000]
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
                "X-Corrections-Count": str(len(cs_data.get("corrections_applied", []))),
                "X-Regeneration-Count": str(rev),
                "X-Cascade-Report": cascade_header,
                "Access-Control-Expose-Headers": (
                    "X-Corrections-Count, X-Regeneration-Count, X-Cascade-Report"
                ),
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.exception(f"[ICGS-REGEN] failed for case {case_id}: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Regenerate failed: {str(e)}"
        )


# =====================================================================
# V3.0 helpers — shared raw-data builder for Case Diary & Remand Report
# =====================================================================
async def _assemble_subdoc_raw_data(
    *, case_id: str, officer_id: str, ics_payload: dict,
) -> dict:
    """
    Build the common raw_data dict that both the case-diary V3.0 and
    remand V3.0 LLM calls consume:

        - manual_input fields (Phase 1, authoritative)
        - ics_structured_data (the already-corrected charge sheet JSON)
        - documents_corpus    (full OCR'd text — sole Phase-2 truth)

    Mirrors the assembly logic in generate_intelligent_charge_sheet_endpoint
    so behaviour and edge cases stay consistent.
    """
    folder = get_case_folder(officer_id, case_id)
    metadata_path = folder / "metadata.json"
    metadata: dict = {}
    if metadata_path.exists():
        with open(metadata_path) as f:
            metadata = json.load(f)

    files_meta = metadata.get("files") or []
    uploaded_documents: List[str] = []
    corpus_parts: List[str] = []
    for fmeta in files_meta:
        fname = fmeta.get("filename") or fmeta.get("saved_name") or "unknown"
        uploaded_documents.append(fname)
        text = (fmeta.get("ocr_text") or fmeta.get("text") or "").strip()
        if not text or len(text) < 40 or text.startswith("[No text in PDF]") or text.startswith("[PDF error"):
            saved = fmeta.get("saved_name")
            if saved:
                fp = folder / saved
                if fp.exists():
                    try:
                        text = (await extract_text_from_staged_file(fp)).strip()
                        fmeta["ocr_text"] = text
                    except Exception as e:
                        logger.warning(f"[V3-SUBDOC] re-OCR failed for {saved}: {e}")
        if text:
            corpus_parts.append(
                f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                f"FILE: {fname}\n"
                f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                f"{text[:25000]}"
            )
    # Persist freshened OCR so subsequent calls skip the re-OCR work
    try:
        with open(metadata_path, "w") as f:
            json.dump(metadata, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

    documents_corpus = "\n\n".join(corpus_parts)

    def _fmt_date(d):
        if not d:
            return ""
        s = str(d).strip()
        if "-" in s and len(s.split("-")[0]) == 4:
            y, m, dd = s.split("-")
            return f"{dd}.{m}.{y}"
        return s.replace("/", ".")

    manual = metadata.get("manual_input") or {}
    raw = {
        "district":           manual.get("district") or metadata.get("district", ""),
        "police_station":     manual.get("police_station") or metadata.get("police_station", ""),
        "fir_number":         manual.get("fir_number") or metadata.get("fir_number", ""),
        "fir_date":           _fmt_date(manual.get("fir_date")) or "",
        "sections":           manual.get("sections") or metadata.get("sections", ""),
        "court_name":         manual.get("court_name") or "",
        "chargesheet_date":   _fmt_date(manual.get("chargesheet_date")) or "",
        "report_type":        manual.get("report_type") or "Charge Sheet.",
        "io": {
            "name":    manual.get("io_name") or (ics_payload.get("io") or {}).get("name", ""),
            "rank":    manual.get("io_rank") or (ics_payload.get("io") or {}).get("rank", ""),
            "salutation": (ics_payload.get("io") or {}).get("salutation") or "Sri.",
            "station": manual.get("police_station") or "",
        },
        "ics_structured_data": ics_payload,
        "uploaded_documents":  uploaded_documents,
        "documents_corpus":    documents_corpus,
    }
    return raw


def _adapt_case_diary_for_fixed_layout(cd: dict) -> dict:
    """Translate intelligent_case_diary V3.0 JSON → fixed_layout_renderer schema."""
    if not isinstance(cd, dict):
        return {}
    io_d = cd.get("io") or {}
    return {
        "police_station":    _nf(cd.get("police_station"), ""),
        "district":          _nf(cd.get("district"), ""),
        "fir_number":        _nf(cd.get("fir_number"), ""),
        "fir_date":          _nf(cd.get("fir_date"), ""),
        "sections":          _nf(cd.get("sections"), ""),
        "occurrence_dtp":    _nf(cd.get("occurrence_dtp"), ""),
        "cd_date":           _nf(cd.get("cd_date") or cd.get("fir_date"), ""),
        "report_datetime":   _nf(cd.get("report_datetime"), ""),
        "last_cd_date":      _nf(cd.get("last_cd_date"), "First CD"),
        "property_lost":     _nf(cd.get("property_lost"), "Nil"),
        "property_recovered":_nf(cd.get("property_recovered"), "Nil"),
        "deceased":          _nf(cd.get("deceased"), "Nil"),
        "circle":            _nf(cd.get("circle"), ""),
        "place":             _nf(cd.get("police_station"), ""),
        "io": {
            "salutation":  io_d.get("salutation") or "Sri.",
            "name":        _nf(io_d.get("name"), ""),
            "designation": _nf(io_d.get("rank") or io_d.get("designation"),
                                "Sub Inspector of Police"),
            "rank":        _nf(io_d.get("rank") or io_d.get("designation"),
                                "Sub Inspector of Police"),
        },
        "complainant":         _adapt_person(cd.get("complainant") or {}),
        "accused":             [_adapt_person(a) for a in (cd.get("accused") or [])],
        "witnesses_examined":  [_adapt_person(w) for w in (cd.get("witnesses_examined") or cd.get("witnesses") or [])],
        "witnesses":           [_adapt_person(w) for w in (cd.get("witnesses_examined") or cd.get("witnesses") or [])],
        "brief_facts":         _nf(cd.get("brief_facts"), ""),
        "investigation_steps": [str(s).strip() for s in (cd.get("investigation_steps") or []) if s],
    }


def _adapt_remand_for_fixed_layout(rr: dict) -> dict:
    """Translate intelligent_remand_report V3.0 JSON → fixed_layout_renderer schema."""
    if not isinstance(rr, dict):
        return {}
    io_d = rr.get("io") or {}
    court_place = rr.get("court_place") or ""
    if not court_place:
        court = (rr.get("court_name") or "").upper()
        if " AT " in court:
            court_place = court.split(" AT ", 1)[1].strip().rstrip(",.")
    return {
        "police_station":        _nf(rr.get("police_station"), ""),
        "district":              _nf(rr.get("district"), ""),
        "fir_number":            _nf(rr.get("fir_number"), ""),
        "fir_date":               _nf(rr.get("fir_date"), ""),
        "sections":              _nf(rr.get("sections"), ""),
        "occurrence_dtp":        _nf(rr.get("occurrence_dtp"), ""),
        "action_taken_datetime": _nf(rr.get("action_taken_datetime"), _nf(rr.get("fir_date"), "")),
        "court_place":           court_place or _nf(rr.get("district"), ""),
        "property_lost":         _nf(rr.get("property_lost"), "Nil"),
        "property_recovered":    _nf(rr.get("property_recovered"), "Nil"),
        "deceased":              _nf(rr.get("deceased"), "Nil"),
        "remand_type":           _nf(rr.get("remand_type"), "judicial"),
        "io": {
            "salutation":  io_d.get("salutation") or "Sri.",
            "name":        _nf(io_d.get("name"), ""),
            "designation": _nf(io_d.get("rank") or io_d.get("designation"),
                                "Sub Inspector of Police"),
            "rank":        _nf(io_d.get("rank") or io_d.get("designation"),
                                "Sub Inspector of Police"),
        },
        "complainant":       _adapt_person(rr.get("complainant") or {}),
        "accused":           [_adapt_person(a) for a in (rr.get("accused") or [])],
        "witnesses":         [_adapt_person(w) for w in (rr.get("witnesses") or [])],
        "brief_facts":       _nf(rr.get("brief_facts"), ""),
        "investigation_done":_nf(rr.get("investigation_done"), ""),
        "grounds_of_arrest": _nf(rr.get("grounds_of_arrest"), ""),
        "enclosures":        [str(e).strip() for e in (rr.get("enclosures") or []) if e],
        "escort":            _nf(rr.get("escort"), ""),
    }


# =====================================================================
# V3.0 — Intelligent Case Diary Part-I (Master IO persona)
# =====================================================================
@router.post("/generate-intelligent-case-diary/{case_id}")
async def generate_intelligent_case_diary_endpoint(
    case_id: str,
    officer: dict = Depends(get_current_officer),
):
    """
    Generate Case Diary Part-I in the V3.0 Master IO style. Reuses the
    already-corrected ICGS structured data, the documents_corpus, and the
    15-field manual input. Output is rendered via the deterministic
    `fixed_layout_renderer.render_case_diary_part1` template.
    COST: 2 credits (deducted on success).
    """
    from fastapi.responses import Response
    from services.intelligent_case_diary import generate_intelligent_case_diary
    from services.fixed_layout_renderer import render_case_diary_part1

    officer_id = officer.get("officer_id", "unknown")
    credits_to_deduct = 2

    if db is None:
        raise HTTPException(status_code=500, detail="Database not initialized")

    officer_doc = await db.officers.find_one({"officer_id": officer_id}, {"_id": 0, "credits": 1})
    current = int((officer_doc or {}).get("credits", 0) or 0)
    if current < credits_to_deduct:
        raise HTTPException(
            status_code=402,
            detail=f"Insufficient credits — Intelligent Case Diary costs {credits_to_deduct} credits, you have {current}. Buy more at /credits.",
        )

    ics = await db.intelligent_chargesheets.find_one(
        {"case_id": case_id, "officer_id": officer_id},
        {"_id": 0}
    )
    if not ics or not ics.get("structured_data"):
        raise HTTPException(
            status_code=400,
            detail="No intelligent charge sheet found. Run 'Generate Station-Format Charge Sheet' first."
        )

    raw_data = await _assemble_subdoc_raw_data(
        case_id=case_id, officer_id=officer_id,
        ics_payload=ics["structured_data"],
    )

    try:
        logger.info(f"[ICD] Generating Case Diary Part-I for case {case_id}")
        cd_data = await generate_intelligent_case_diary(
            raw_data, session_id=f"icd-{case_id}",
        )
        adapted = _adapt_case_diary_for_fixed_layout(cd_data)
        docx_bytes = render_case_diary_part1(adapted)

        await db.intelligent_case_diaries.update_one(
            {"case_id": case_id, "officer_id": officer_id},
            {"$set": {
                "case_id": case_id,
                "officer_id": officer_id,
                "fir_number": cd_data.get("fir_number", ""),
                "structured_data": cd_data,
                "steps_count": len(cd_data.get("investigation_steps", [])),
                "model_used": cd_data.get("_model_used", ""),
                "generated_at": datetime.now(timezone.utc).isoformat(),
                "credits_used": credits_to_deduct,
                "regeneration_count": 0,
            }},
            upsert=True,
        )
        await db.officers.update_one(
            {"officer_id": officer_id},
            {"$inc": {"credits": -credits_to_deduct}}
        )
        await db.action_logs.insert_one({
            "officer_id": officer_id,
            "action": "INTELLIGENT_CASE_DIARY_GENERATE",
            "credit_cost": credits_to_deduct,
            "status": "SUCCESS",
            "correlation_id": case_id,
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "model_used": cd_data.get("_model_used", ""),
        })

        fir_safe = (cd_data.get("fir_number") or "case").replace("/", "-")
        filename = f"{fir_safe}_IntelligentCaseDiary.docx"
        report = cd_data.get("extraction_report") or {}
        import json as _json
        report_header = _json.dumps({
            "manual_input_fields_used": report.get("manual_input_fields_used", 9),
            "investigation_steps_count": len(cd_data.get("investigation_steps", [])),
            "total_accused": len(cd_data.get("accused") or []),
            "total_witnesses": len(cd_data.get("witnesses_examined") or cd_data.get("witnesses") or []),
            "not_found_fields": report.get("not_found_fields", []),
            "confidence": report.get("confidence", "High"),
            "confidence_reason": report.get("confidence_reason", ""),
        })[:6000]
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
                "X-Steps-Count": str(len(cd_data.get("investigation_steps", []))),
                "X-Model-Used": cd_data.get("_model_used", ""),
                "X-Extraction-Report": report_header,
                "Access-Control-Expose-Headers": (
                    "X-Steps-Count, X-Model-Used, X-Extraction-Report"
                ),
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.exception(f"[ICD] Case Diary FAILED for case {case_id}: {e}")
        if db is not None:
            await db.action_logs.insert_one({
                "officer_id": officer_id,
                "action": "INTELLIGENT_CASE_DIARY_GENERATE",
                "credit_cost": 0,
                "status": "FAILED",
                "correlation_id": case_id,
                "timestamp": datetime.now(timezone.utc).isoformat(),
                "error": str(e),
            })
        raise HTTPException(
            status_code=500,
            detail=f"Case Diary generation failed. NO CREDITS DEDUCTED. Error: {str(e)}"
        )


@router.get("/intelligent-case-diary/{case_id}")
async def get_intelligent_case_diary_metadata(
    case_id: str,
    officer: dict = Depends(get_current_officer),
):
    """Return the corrections applied + structured data for the last intelligent case diary."""
    if db is None:
        raise HTTPException(status_code=500, detail="Database not initialized")
    doc = await db.intelligent_case_diaries.find_one(
        {"case_id": case_id, "officer_id": officer.get("officer_id", "unknown")},
        {"_id": 0}
    )
    if not doc:
        return {"success": False, "message": "No intelligent case diary generated for this case yet"}
    return {"success": True, **doc}


@router.post("/regenerate-case-diary/{case_id}")
async def regenerate_case_diary(
    case_id: str,
    body: _RegenerateRequest,
    officer: dict = Depends(get_current_officer),
):
    """
    Re-run the Intelligent Case Diary with user-supplied corrections.
    Cost: 0 credits (re-render of an already-paid run).
    """
    from fastapi.responses import Response
    from services.intelligent_case_diary import generate_intelligent_case_diary
    from services.fixed_layout_renderer import render_case_diary_part1

    officer_id = officer.get("officer_id", "unknown")
    if not body.corrections:
        raise HTTPException(status_code=400, detail="Provide at least one correction")
    if db is None:
        raise HTTPException(status_code=500, detail="Database not initialized")

    prev = await db.intelligent_case_diaries.find_one(
        {"case_id": case_id, "officer_id": officer_id}, {"_id": 0}
    )
    if not prev:
        raise HTTPException(
            status_code=400,
            detail="No previous intelligent case diary found. Generate one first.",
        )
    ics = await db.intelligent_chargesheets.find_one(
        {"case_id": case_id, "officer_id": officer_id}, {"_id": 0}
    )
    if not ics or not ics.get("structured_data"):
        raise HTTPException(status_code=400, detail="ICGS payload missing")

    raw_data = await _assemble_subdoc_raw_data(
        case_id=case_id, officer_id=officer_id,
        ics_payload=ics["structured_data"],
    )
    raw_data["corrections"] = [c.dict() for c in body.corrections]
    raw_data["previous_payload"] = prev.get("structured_data") or {}
    # Regenerate is FAST-PATH: the LLM has the previous_payload as ground truth;
    # trimming the documents_corpus keeps the call under the 60s K8s ingress limit
    # while still allowing the LLM to reference source documents when needed.
    if raw_data.get("documents_corpus"):
        raw_data["documents_corpus"] = raw_data["documents_corpus"][:8000]

    try:
        logger.info(
            f"[ICD-REGEN] case {case_id}: applying {len(body.corrections)} correction(s)"
        )
        cd_data = await generate_intelligent_case_diary(
            raw_data, session_id=f"icd-regen-{case_id}-{uuid.uuid4().hex[:6]}",
        )
        adapted = _adapt_case_diary_for_fixed_layout(cd_data)
        docx_bytes = render_case_diary_part1(adapted)

        rev = (prev.get("regeneration_count", 0) or 0) + 1
        await db.intelligent_case_diaries.update_one(
            {"case_id": case_id, "officer_id": officer_id},
            {"$set": {
                "case_id": case_id, "officer_id": officer_id,
                "fir_number": cd_data.get("fir_number", ""),
                "structured_data": cd_data,
                "corrections_applied": cd_data.get("corrections_applied", []),
                "regeneration_count": rev,
                "last_corrections": [c.dict() for c in body.corrections],
                "model_used": cd_data.get("_model_used", ""),
                "regenerated_at": datetime.now(timezone.utc).isoformat(),
            }},
            upsert=True,
        )
        await db.action_logs.insert_one({
            "officer_id": officer_id,
            "action": "INTELLIGENT_CASE_DIARY_REGENERATE",
            "credit_cost": 0,
            "status": "SUCCESS",
            "correlation_id": case_id,
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "metadata": {"correction_count": len(body.corrections)},
        })

        fir_safe = (cd_data.get("fir_number") or "case").replace("/", "-")
        filename = f"{fir_safe}_IntelligentCaseDiary_rev{rev}.docx"
        import json as _json
        cascade_header = _json.dumps({
            "corrections_applied": cd_data.get("corrections_applied", []),
            "regeneration_count": rev,
            "user_corrections": [c.dict() for c in body.corrections],
        })[:6000]
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
                "X-Corrections-Count": str(len(cd_data.get("corrections_applied", []))),
                "X-Regeneration-Count": str(rev),
                "X-Cascade-Report": cascade_header,
                "Access-Control-Expose-Headers": (
                    "X-Corrections-Count, X-Regeneration-Count, X-Cascade-Report"
                ),
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.exception(f"[ICD-REGEN] failed for case {case_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Regenerate failed: {str(e)}")


# =====================================================================
# V3.0 — Intelligent Remand Report (Master IO persona)
# =====================================================================
@router.post("/generate-intelligent-remand-report/{case_id}")
async def generate_intelligent_remand_report_endpoint(
    case_id: str,
    officer: dict = Depends(get_current_officer),
):
    """
    Generate Remand Case Diary Part-I letter in V3.0 Master IO style.
    Reuses the ICGS structured data + documents_corpus + manual input.
    Output rendered via `fixed_layout_renderer.render_remand_report`.
    COST: 2 credits (deducted on success).
    """
    from fastapi.responses import Response
    from services.intelligent_remand_report import generate_intelligent_remand_report
    from services.fixed_layout_renderer import render_remand_report

    officer_id = officer.get("officer_id", "unknown")
    credits_to_deduct = 2

    if db is None:
        raise HTTPException(status_code=500, detail="Database not initialized")

    officer_doc = await db.officers.find_one({"officer_id": officer_id}, {"_id": 0, "credits": 1})
    current = int((officer_doc or {}).get("credits", 0) or 0)
    if current < credits_to_deduct:
        raise HTTPException(
            status_code=402,
            detail=f"Insufficient credits — Intelligent Remand Report costs {credits_to_deduct} credits, you have {current}. Buy more at /credits.",
        )

    ics = await db.intelligent_chargesheets.find_one(
        {"case_id": case_id, "officer_id": officer_id},
        {"_id": 0}
    )
    if not ics or not ics.get("structured_data"):
        raise HTTPException(
            status_code=400,
            detail="No intelligent charge sheet found. Run 'Generate Station-Format Charge Sheet' first."
        )

    raw_data = await _assemble_subdoc_raw_data(
        case_id=case_id, officer_id=officer_id,
        ics_payload=ics["structured_data"],
    )

    try:
        logger.info(f"[IRR] Generating Remand Report for case {case_id}")
        rr_data = await generate_intelligent_remand_report(
            raw_data, session_id=f"irr-{case_id}",
        )
        adapted = _adapt_remand_for_fixed_layout(rr_data)
        docx_bytes = render_remand_report(adapted)

        await db.intelligent_remand_reports.update_one(
            {"case_id": case_id, "officer_id": officer_id},
            {"$set": {
                "case_id": case_id,
                "officer_id": officer_id,
                "fir_number": rr_data.get("fir_number", ""),
                "structured_data": rr_data,
                "model_used": rr_data.get("_model_used", ""),
                "generated_at": datetime.now(timezone.utc).isoformat(),
                "credits_used": credits_to_deduct,
                "regeneration_count": 0,
            }},
            upsert=True,
        )
        await db.officers.update_one(
            {"officer_id": officer_id},
            {"$inc": {"credits": -credits_to_deduct}}
        )
        await db.action_logs.insert_one({
            "officer_id": officer_id,
            "action": "INTELLIGENT_REMAND_GENERATE",
            "credit_cost": credits_to_deduct,
            "status": "SUCCESS",
            "correlation_id": case_id,
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "model_used": rr_data.get("_model_used", ""),
        })

        fir_safe = (rr_data.get("fir_number") or "case").replace("/", "-")
        filename = f"{fir_safe}_IntelligentRemandReport.docx"
        report = rr_data.get("extraction_report") or {}
        import json as _json
        report_header = _json.dumps({
            "manual_input_fields_used": report.get("manual_input_fields_used", 9),
            "total_accused": len(rr_data.get("accused") or []),
            "total_witnesses": len(rr_data.get("witnesses") or []),
            "not_found_fields": report.get("not_found_fields", []),
            "confidence": report.get("confidence", "High"),
            "confidence_reason": report.get("confidence_reason", ""),
        })[:6000]
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
                "X-Model-Used": rr_data.get("_model_used", ""),
                "X-Extraction-Report": report_header,
                "Access-Control-Expose-Headers": (
                    "X-Model-Used, X-Extraction-Report"
                ),
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.exception(f"[IRR] Remand Report FAILED for case {case_id}: {e}")
        if db is not None:
            await db.action_logs.insert_one({
                "officer_id": officer_id,
                "action": "INTELLIGENT_REMAND_GENERATE",
                "credit_cost": 0,
                "status": "FAILED",
                "correlation_id": case_id,
                "timestamp": datetime.now(timezone.utc).isoformat(),
                "error": str(e),
            })
        raise HTTPException(
            status_code=500,
            detail=f"Remand Report generation failed. NO CREDITS DEDUCTED. Error: {str(e)}"
        )


@router.get("/intelligent-remand-report/{case_id}")
async def get_intelligent_remand_report_metadata(
    case_id: str,
    officer: dict = Depends(get_current_officer),
):
    """Return the corrections applied + structured data for the last intelligent remand report."""
    if db is None:
        raise HTTPException(status_code=500, detail="Database not initialized")
    doc = await db.intelligent_remand_reports.find_one(
        {"case_id": case_id, "officer_id": officer.get("officer_id", "unknown")},
        {"_id": 0}
    )
    if not doc:
        return {"success": False, "message": "No intelligent remand report generated for this case yet"}
    return {"success": True, **doc}


@router.post("/regenerate-remand-report/{case_id}")
async def regenerate_remand_report(
    case_id: str,
    body: _RegenerateRequest,
    officer: dict = Depends(get_current_officer),
):
    """
    Re-run the Intelligent Remand Report with user-supplied corrections.
    Cost: 0 credits (re-render of an already-paid run).
    """
    from fastapi.responses import Response
    from services.intelligent_remand_report import generate_intelligent_remand_report
    from services.fixed_layout_renderer import render_remand_report

    officer_id = officer.get("officer_id", "unknown")
    if not body.corrections:
        raise HTTPException(status_code=400, detail="Provide at least one correction")
    if db is None:
        raise HTTPException(status_code=500, detail="Database not initialized")

    prev = await db.intelligent_remand_reports.find_one(
        {"case_id": case_id, "officer_id": officer_id}, {"_id": 0}
    )
    if not prev:
        raise HTTPException(
            status_code=400,
            detail="No previous intelligent remand report found. Generate one first.",
        )
    ics = await db.intelligent_chargesheets.find_one(
        {"case_id": case_id, "officer_id": officer_id}, {"_id": 0}
    )
    if not ics or not ics.get("structured_data"):
        raise HTTPException(status_code=400, detail="ICGS payload missing")

    raw_data = await _assemble_subdoc_raw_data(
        case_id=case_id, officer_id=officer_id,
        ics_payload=ics["structured_data"],
    )
    raw_data["corrections"] = [c.dict() for c in body.corrections]
    raw_data["previous_payload"] = prev.get("structured_data") or {}
    # Regenerate is FAST-PATH: trim documents_corpus to keep under K8s 60s ingress.
    if raw_data.get("documents_corpus"):
        raw_data["documents_corpus"] = raw_data["documents_corpus"][:8000]

    try:
        logger.info(
            f"[IRR-REGEN] case {case_id}: applying {len(body.corrections)} correction(s)"
        )
        rr_data = await generate_intelligent_remand_report(
            raw_data, session_id=f"irr-regen-{case_id}-{uuid.uuid4().hex[:6]}",
        )
        adapted = _adapt_remand_for_fixed_layout(rr_data)
        docx_bytes = render_remand_report(adapted)

        rev = (prev.get("regeneration_count", 0) or 0) + 1
        await db.intelligent_remand_reports.update_one(
            {"case_id": case_id, "officer_id": officer_id},
            {"$set": {
                "case_id": case_id, "officer_id": officer_id,
                "fir_number": rr_data.get("fir_number", ""),
                "structured_data": rr_data,
                "corrections_applied": rr_data.get("corrections_applied", []),
                "regeneration_count": rev,
                "last_corrections": [c.dict() for c in body.corrections],
                "model_used": rr_data.get("_model_used", ""),
                "regenerated_at": datetime.now(timezone.utc).isoformat(),
            }},
            upsert=True,
        )
        await db.action_logs.insert_one({
            "officer_id": officer_id,
            "action": "INTELLIGENT_REMAND_REGENERATE",
            "credit_cost": 0,
            "status": "SUCCESS",
            "correlation_id": case_id,
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "metadata": {"correction_count": len(body.corrections)},
        })

        fir_safe = (rr_data.get("fir_number") or "case").replace("/", "-")
        filename = f"{fir_safe}_IntelligentRemandReport_rev{rev}.docx"
        import json as _json
        cascade_header = _json.dumps({
            "corrections_applied": rr_data.get("corrections_applied", []),
            "regeneration_count": rev,
            "user_corrections": [c.dict() for c in body.corrections],
        })[:6000]
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
                "X-Corrections-Count": str(len(rr_data.get("corrections_applied", []))),
                "X-Regeneration-Count": str(rev),
                "X-Cascade-Report": cascade_header,
                "Access-Control-Expose-Headers": (
                    "X-Corrections-Count, X-Regeneration-Count, X-Cascade-Report"
                ),
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.exception(f"[IRR-REGEN] failed for case {case_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Regenerate failed: {str(e)}")



# ============================================================
# FIXED-LAYOUT DOCUMENT GENERATION
# (Charge Sheet, Case Diary Part-I, Remand Report)
# Layout never changes. Missing fields render as blank "_____".
# Aadhaar is auto-extracted from already-staged files when present.
# ============================================================

from services.fixed_layout_renderer import render_fixed_doc, extract_aadhaar_from_files
from fastapi.responses import StreamingResponse


def _build_case_data_from_metadata(metadata: dict, officer: dict, doc_type: str) -> dict:
    """
    Walk the staged metadata + per-file extracted_data and build a normalised
    `case_data` dict the renderer understands. We stay strictly mechanical here
    — no AI, no inference. Whatever the OCR/parser produced is what we plug in.
    """
    files = metadata.get("files", []) or []
    case = {
        "police_station": metadata.get("police_station") or metadata.get("ps_name") or "",
        "district": metadata.get("district") or "",
        "fir_number": metadata.get("fir_number") or metadata.get("fir_no") or "",
        "fir_date": metadata.get("fir_date") or "",
        "sections": metadata.get("sections") or "",
        "place": metadata.get("place") or metadata.get("police_station") or "",
        "today_date": datetime.now().strftime("%d-%m-%Y"),
        "io": {
            "name": officer.get("name", ""),
            "designation": officer.get("rank") or officer.get("role", ""),
            "phone": officer.get("phone", ""),
        },
        "complainant": metadata.get("complainant") or {},
        "accused": metadata.get("accused") or [],
        "witnesses": metadata.get("witnesses") or [],
        "material_objects": metadata.get("material_objects") or [],
        "brief_facts": metadata.get("brief_facts") or "",
        # Case-diary-specific
        "diary_no": metadata.get("diary_no", ""),
        "diary_date": metadata.get("diary_date", ""),
        "places_visited": metadata.get("places_visited", ""),
        "distance_travelled": metadata.get("distance_travelled", ""),
        "time_departure": metadata.get("time_departure", ""),
        "time_arrival": metadata.get("time_arrival", ""),
        "witnesses_examined": metadata.get("witnesses_examined", []),
        "search_seizure": metadata.get("search_seizure", ""),
        "material_seized": metadata.get("material_seized", ""),
        "actions_today": metadata.get("actions_today", ""),
        "findings": metadata.get("findings", ""),
        "next_steps": metadata.get("next_steps", ""),
        # Remand-specific
        "court_name": metadata.get("court_name", ""),
        "court_address": metadata.get("court_address", ""),
        "arrest_datetime": metadata.get("arrest_datetime", ""),
        "production_datetime": metadata.get("production_datetime", ""),
        "grounds_of_arrest": metadata.get("grounds_of_arrest", ""),
        "investigation_done": metadata.get("investigation_done", ""),
        "reasons_for_remand": metadata.get("reasons_for_remand", ""),
        "further_investigation": metadata.get("further_investigation", ""),
        "remand_type": metadata.get("remand_type", ""),
        "remand_duration": metadata.get("remand_duration", ""),
        "remand_from": metadata.get("remand_from", ""),
        "remand_to": metadata.get("remand_to", ""),
    }

    # Auto-extract Aadhaar fields from any staged file with OCR text and
    # plug them into A1 if A1 has no aadhaar number yet.
    aad = extract_aadhaar_from_files(files)
    if aad.get("aadhaar_number"):
        if not case["accused"]:
            case["accused"] = [{}]
        a1 = case["accused"][0] if isinstance(case["accused"][0], dict) else {}
        if not a1.get("aadhaar_number"):
            a1["aadhaar_number"] = aad["aadhaar_number"]
        if not a1.get("name") and aad.get("aadhaar_name"):
            a1["name"] = aad["aadhaar_name"]
        if not a1.get("permanent_address") and aad.get("aadhaar_address"):
            a1["permanent_address"] = aad["aadhaar_address"]
            # Also set the generic 'address' so the renderer's R/o block uses it
            if not a1.get("address"):
                a1["address"] = aad["aadhaar_address"]
        if not a1.get("gender") and aad.get("aadhaar_gender"):
            a1["gender"] = aad["aadhaar_gender"]
        # DOB is metadata, not the literal age — keep it in a separate field
        if not a1.get("dob") and aad.get("aadhaar_dob"):
            a1["dob"] = aad["aadhaar_dob"]
        case["accused"][0] = a1

    return case


@router.get("/render-fixed/{doc_type}/{case_id}")
async def render_fixed_layout_doc(
    doc_type: str,
    case_id: str,
    officer: dict = Depends(get_current_officer)
):
    """
    Generate one of the 3 fixed-layout documents from the staged case files.
    Layout is hard-coded; values are taken from metadata.json + auto-extracted
    Aadhaar fields. Missing fields render as `_____` blanks for the officer
    to fill in Word.

    doc_type ∈ {charge_sheet, case_diary_part1, remand_report}
    """
    if doc_type not in ("charge_sheet", "case_diary_part1", "remand_report"):
        raise HTTPException(status_code=400, detail="Invalid doc_type")

    folder = get_case_folder(officer.get("officer_id", "unknown"), case_id)
    if not folder.exists():
        raise HTTPException(status_code=404, detail="Case folder not found")

    metadata_path = folder / "metadata.json"
    if not metadata_path.exists():
        raise HTTPException(status_code=404, detail="Case metadata not found — upload files first")

    with open(metadata_path) as f:
        metadata = json.load(f)

    try:
        case_data = _build_case_data_from_metadata(metadata, officer, doc_type)
        docx_bytes, filename = render_fixed_doc(doc_type, case_data)
    except Exception as e:
        logger.error(f"Fixed-layout render failed for {doc_type}/{case_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Rendering failed: {str(e)}")

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ────────────────────────────────────────────────────────────────
# Convenience alias endpoints (1-to-1 mapping for clearer URLs)
# ────────────────────────────────────────────────────────────────
@router.get("/generate-fixed-charge-sheet/{case_id}")
async def generate_fixed_charge_sheet(
    case_id: str,
    officer: dict = Depends(get_current_officer),
):
    return await render_fixed_layout_doc("charge_sheet", case_id, officer)


@router.get("/generate-fixed-case-diary/{case_id}")
async def generate_fixed_case_diary(
    case_id: str,
    officer: dict = Depends(get_current_officer),
):
    return await render_fixed_layout_doc("case_diary_part1", case_id, officer)


@router.get("/generate-fixed-remand/{case_id}")
async def generate_fixed_remand(
    case_id: str,
    officer: dict = Depends(get_current_officer),
):
    return await render_fixed_layout_doc("remand_report", case_id, officer)
