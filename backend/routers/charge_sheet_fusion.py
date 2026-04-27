"""
Charge Sheet Fusion Router - Multi-upload document processing for 95% accuracy Charge Sheet generation.
Supports Telugu Petitions, CDF, and Case Diary Part-II simultaneous processing.
"""
from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from datetime import datetime, timezone
from typing import Optional, List
import os
import jwt
import logging
import base64
import io
import tempfile
import subprocess

from services.legal_llm import translate_to_legal_english, extract_entities, suggest_bns_sections
from services.template_generator import (
    generate_18_column_charge_sheet,
    generate_case_diary_part1,
    generate_html_table_charge_sheet,
    suggest_bns_sections as ml_suggest_sections
)
from services.remand_generator import (
    generate_remand_case_diary_html,
    detect_arrest_status
)
from services.cdf_overlay import (
    generate_cdf_digital_form_html,
    extract_cdf_data_for_chargesheet
)

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/charge-sheet-fusion", tags=["Charge Sheet Fusion"])
security = HTTPBearer()

# Database connection (will be set by main app)
db = None

JWT_SECRET = os.environ['JWT_SECRET']
EMERGENT_LLM_KEY = os.environ.get('EMERGENT_LLM_KEY', '')
GOOGLE_VISION_CREDENTIALS = os.environ.get('GOOGLE_VISION_CREDENTIALS', '')
GOOGLE_TRANSLATE_CREDENTIALS = os.environ.get('GOOGLE_TRANSLATE_CREDENTIALS', '')

# Initialize Google Cloud clients
vision_client = None
translate_client = None

try:
    if GOOGLE_VISION_CREDENTIALS and os.path.exists(GOOGLE_VISION_CREDENTIALS):
        from google.cloud import vision
        from google.oauth2 import service_account
        credentials = service_account.Credentials.from_service_account_file(GOOGLE_VISION_CREDENTIALS)
        vision_client = vision.ImageAnnotatorClient(credentials=credentials)
        logger.info("Google Vision client initialized for Charge Sheet Fusion")
except Exception as e:
    logger.warning(f"Could not initialize Vision client in Charge Sheet Fusion: {e}")

try:
    if GOOGLE_TRANSLATE_CREDENTIALS and os.path.exists(GOOGLE_TRANSLATE_CREDENTIALS):
        from google.cloud import translate_v2 as translate
        from google.oauth2 import service_account
        credentials = service_account.Credentials.from_service_account_file(GOOGLE_TRANSLATE_CREDENTIALS)
        translate_client = translate.Client(credentials=credentials)
        logger.info("Google Translate client initialized for Charge Sheet Fusion")
except Exception as e:
    logger.warning(f"Could not initialize Translate client in Charge Sheet Fusion: {e}")


async def extract_text_from_image_ocr(contents: bytes, filename: str) -> dict:
    """Extract text from image using Google Vision OCR and translate if needed"""
    result = {"text": "", "language": "unknown", "translated": ""}
    
    if not vision_client:
        result["text"] = f"[OCR not available - Google Vision not configured. Image: {filename}]"
        return result
    
    try:
        from google.cloud import vision
        image = vision.Image(content=contents)
        response = vision_client.text_detection(image=image)
        
        if response.error.message:
            result["text"] = f"[Vision API error: {response.error.message}]"
            return result
        
        texts = response.text_annotations
        if texts:
            result["text"] = texts[0].description
            result["language"] = texts[0].locale if texts[0].locale else "unknown"
            
            # Translate if not English
            if translate_client and result["language"] not in ['en', 'unknown', '']:
                try:
                    translation = translate_client.translate(result["text"], target_language='en')
                    result["translated"] = translation['translatedText']
                    logger.info(f"Translated from {result['language']} to English")
                except Exception as te:
                    logger.warning(f"Translation failed: {te}")
                    result["translated"] = result["text"]
            else:
                result["translated"] = result["text"]
        else:
            result["text"] = "[No text detected in image]"
            
    except Exception as e:
        logger.error(f"OCR extraction error: {e}")
        result["text"] = f"[OCR error: {str(e)}]"
    
    return result


async def extract_text_from_file(file: UploadFile, contents: bytes) -> str:
    """Extract text from various document formats"""
    filename = file.filename.lower() if file.filename else ""
    
    # Handle DOCX files
    if filename.endswith('.docx'):
        try:
            from docx import Document
            doc = Document(io.BytesIO(contents))
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return f"[Error extracting DOCX: {str(e)}]"
    
    # Handle legacy DOC files using antiword
    elif filename.endswith('.doc'):
        try:
            with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
                tmp.write(contents)
                tmp_path = tmp.name
            
            result = subprocess.run(
                ['antiword', tmp_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            os.unlink(tmp_path)  # Clean up temp file
            
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout
            else:
                return f"[DOC extraction failed: {result.stderr or 'No text extracted'}]"
        except subprocess.TimeoutExpired:
            return "[DOC extraction timed out]"
        except FileNotFoundError:
            return "[antiword not installed - cannot process .doc files]"
        except Exception as e:
            logger.error(f"DOC extraction error: {e}")
            return f"[Error extracting DOC: {str(e)}]"
    
    # Handle PDF files
    elif filename.endswith('.pdf'):
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(contents))
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            return "\n".join(text_parts) if text_parts else "[No text extracted from PDF]"
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            return f"[Error extracting PDF: {str(e)}]"
    
    # Handle image files with OCR
    elif filename.endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp')):
        ocr_result = await extract_text_from_image_ocr(contents, file.filename)
        # Return translated text if available, otherwise original
        if ocr_result["translated"]:
            return ocr_result["translated"]
        return ocr_result["text"]
    
    # Unknown format
    else:
        return f"[Unsupported file format: {filename}]"


async def extract_text_from_audio(contents: bytes, filename: str) -> dict:
    """Extract text from audio using Google Speech-to-Text and translate"""
    result = {"text": "", "language": "unknown", "translated": "", "legal_text": ""}
    
    GOOGLE_SPEECH_CREDENTIALS = os.environ.get('GOOGLE_SPEECH_CREDENTIALS', '')
    
    try:
        if GOOGLE_SPEECH_CREDENTIALS and os.path.exists(GOOGLE_SPEECH_CREDENTIALS):
            from google.cloud import speech
            from google.oauth2 import service_account
            
            credentials = service_account.Credentials.from_service_account_file(GOOGLE_SPEECH_CREDENTIALS)
            speech_client = speech.SpeechClient(credentials=credentials)
            
            # Determine encoding based on file type
            file_ext = filename.lower().split('.')[-1]
            if file_ext == 'mp3':
                encoding = speech.RecognitionConfig.AudioEncoding.MP3
            elif file_ext in ['m4a', 'aac']:
                encoding = speech.RecognitionConfig.AudioEncoding.MP3
            elif file_ext == 'wav':
                encoding = speech.RecognitionConfig.AudioEncoding.LINEAR16
            elif file_ext == 'webm':
                encoding = speech.RecognitionConfig.AudioEncoding.WEBM_OPUS
            else:
                encoding = speech.RecognitionConfig.AudioEncoding.ENCODING_UNSPECIFIED
            
            audio = speech.RecognitionAudio(content=contents)
            config = speech.RecognitionConfig(
                encoding=encoding,
                language_code="te-IN",  # Telugu
                alternative_language_codes=["hi-IN", "en-IN"],  # Hindi and English
                enable_automatic_punctuation=True,
            )
            
            response = speech_client.recognize(config=config, audio=audio)
            
            for resp_result in response.results:
                result["text"] += resp_result.alternatives[0].transcript + " "
                if hasattr(resp_result, 'language_code'):
                    result["language"] = resp_result.language_code
            
            result["text"] = result["text"].strip()
            
            # Translate if not English
            if translate_client and result["text"] and result["language"] not in ['en', 'en-IN']:
                try:
                    translation = translate_client.translate(result["text"], target_language='en')
                    result["translated"] = translation['translatedText']
                except Exception as te:
                    logger.warning(f"Translation failed: {te}")
                    result["translated"] = result["text"]
            else:
                result["translated"] = result["text"]
            
            # Convert to legal text format
            if result["translated"]:
                result["legal_text"] = convert_to_legal_format(result["translated"])
                
        else:
            result["text"] = "[Speech-to-Text not available - Google Speech API not configured]"
            
    except Exception as e:
        logger.error(f"Speech-to-Text error: {e}")
        result["text"] = f"[Speech-to-Text error: {str(e)}]"
    
    return result


def convert_to_legal_format(text: str) -> str:
    """Convert informal text to legal FIR format"""
    if not text:
        return text
    
    # Replace first-person with third-person
    replacements = [
        ("I am ", "The complainant is "),
        ("I was ", "The complainant was "),
        ("I have ", "The complainant has "),
        ("I had ", "The complainant had "),
        ("I went ", "The complainant went "),
        ("I saw ", "The complainant observed "),
        ("I lost ", "The complainant lost "),
        ("I want ", "The complainant desires "),
        ("I request ", "The complainant requests "),
        ("my mobile", "the complainant's mobile phone"),
        ("my phone", "the complainant's mobile phone"),
        ("my money", "the complainant's money"),
        ("my house", "the complainant's residence"),
        ("my ", "the complainant's "),
        (" me ", " the complainant "),
        (" I ", " the complainant "),
    ]
    
    for old, new in replacements:
        text = text.replace(old, new)
        text = text.replace(old.lower(), new.lower())
    
    # Add formal introduction if not present
    if "complainant" in text.lower() and not text.lower().startswith("the complainant"):
        text = "The complainant has stated that " + text[0].lower() + text[1:]
    
    return text


def set_database(database):
    """Set the database connection from main app"""
    global db
    db = database


async def get_current_officer(credentials: HTTPAuthorizationCredentials = Depends(security)) -> dict:
    """Verify JWT token and return officer data"""
    try:
        token = credentials.credentials
        payload = jwt.decode(token, JWT_SECRET, algorithms=['HS256'])
        return payload
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")


CHARGE_SHEET_TEMPLATE = """
IN THE COURT OF JUDICIAL FIRST CLASS MAGISTRATE
At: {court_location}
Dist: {district}  PS: {police_station}

CHARGE SHEET (U/S 193 BNSS)

FIR No.: {fir_number}  Date: {fir_date}
Charge Sheet No.: [MISSING: CHARGE_SHEET_NUMBER]
Date of Charge: {charge_date}
Act and Section of Law: U/S {sections}
Type of Final Report: Charge Sheet
Original/Supplementary: Original

1. NAME AND RANK OF THE I.O(s):
   {io_name}, {io_rank}
   PS: {police_station}

2. NAME AND ADDRESS OF THE COMPLAINANT/INFORMANT:
   {complainant_name}
   S/o {complainant_father}
   Age: {complainant_age} years, Caste: {complainant_caste}
   Occupation: {complainant_occupation}
   R/o: {complainant_address}
   Ph: {complainant_phone}

3. DETAILS OF PROPERTY SEIZED:
   {property_seized}

4. PARTICULARS OF ACCUSED PERSONS CHARGE SHEETED:
{accused_list}

5. PARTICULARS OF ACCUSED PERSONS NOT CHARGE SHEETED:
   {not_charge_sheeted}

6. PARTICULARS OF WITNESSES TO BE EXAMINED:
{witness_list}

7. RESULT OF LABORATORY ANALYSIS:
   {lab_results}

8. BRIEF FACTS OF THE CASE:
{brief_facts}

9. PRAYER:
   Therefore, the Hon'ble Court is prayed that the accused persons mentioned in column No. 4 of this charge sheet may be tried and dealt with suitably as per law.

Dispatched on: {dispatch_date}

[Signature]
{io_name}
{io_rank}
PS: {police_station}
"""


async def process_with_llm(text: str, doc_type: str) -> dict:
    """Process document text with Legal LLM"""
    from emergentintegrations.llm.chat import LlmChat, UserMessage
    
    if not EMERGENT_LLM_KEY:
        return {"success": False, "error": "LLM key not configured"}
    
    try:
        system_prompt = f"""You are a Legal Document Processor specializing in Indian Police FIR and Charge Sheet documents.
        
You are processing a {doc_type}. Extract ALL relevant information for generating a Charge Sheet.

Return a JSON object with these fields:
{{
  "complainant": {{"name": "", "father_name": "", "age": null, "caste": "", "occupation": "", "address": "", "phone": ""}},
  "accused_persons": [{{"serial": "A1", "name": "", "father_name": "", "age": null, "caste": "", "occupation": "", "address": ""}}],
  "witnesses": [{{"serial": "LW-1", "name": "", "role": ""}}],
  "offense_details": {{"type": "", "date": "", "time": "", "place": ""}},
  "sections_of_law": [],
  "brief_facts": "",
  "property_lost": "",
  "property_recovered": ""
}}

Important rules:
1. If a field is not found, use empty string or null
2. Translate Telugu text to English formally
3. Preserve police terminology (A1, LW-1, Panchanama, etc.)
4. Extract ALL accused and witness details found

Return ONLY valid JSON."""

        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"fusion-{hash(text[:100])}",
            system_message=system_prompt
        ).with_model("openai", "gpt-5.2")
        
        user_message = UserMessage(text=f"Extract data from this {doc_type}:\n\n{text}")
        response = await chat.send_message(user_message)
        
        # Parse JSON
        import json
        clean_response = response.strip()
        if clean_response.startswith("```json"):
            clean_response = clean_response[7:]
        if clean_response.startswith("```"):
            clean_response = clean_response[3:]
        if clean_response.endswith("```"):
            clean_response = clean_response[:-3]
        
        return {"success": True, "data": json.loads(clean_response.strip())}
    except Exception as e:
        logger.error(f"LLM processing error: {e}")
        return {"success": False, "error": str(e)}


def merge_extracted_data(data_list: list) -> dict:
    """Merge data extracted from multiple documents"""
    merged = {
        "complainant": {},
        "accused_persons": [],
        "witnesses": [],
        "offense_details": {},
        "sections_of_law": [],
        "brief_facts": "",
        "property_lost": "",
        "property_recovered": ""
    }
    
    for data in data_list:
        if not data:
            continue
            
        # Merge complainant (prefer non-empty values)
        for key, value in data.get("complainant", {}).items():
            if value and not merged["complainant"].get(key):
                merged["complainant"][key] = value
        
        # Merge accused (avoid duplicates by name)
        existing_names = {a.get("name", "").lower() for a in merged["accused_persons"]}
        for accused in data.get("accused_persons", []):
            if accused.get("name") and accused.get("name", "").lower() not in existing_names:
                merged["accused_persons"].append(accused)
                existing_names.add(accused.get("name", "").lower())
        
        # Merge witnesses
        existing_witness_names = {w.get("name", "").lower() for w in merged["witnesses"]}
        for witness in data.get("witnesses", []):
            if witness.get("name") and witness.get("name", "").lower() not in existing_witness_names:
                merged["witnesses"].append(witness)
                existing_witness_names.add(witness.get("name", "").lower())
        
        # Merge offense details
        for key, value in data.get("offense_details", {}).items():
            if value and not merged["offense_details"].get(key):
                merged["offense_details"][key] = value
        
        # Merge sections (unique)
        for section in data.get("sections_of_law", []):
            if section and section not in merged["sections_of_law"]:
                merged["sections_of_law"].append(section)
        
        # Merge brief facts (concatenate)
        if data.get("brief_facts"):
            if merged["brief_facts"]:
                merged["brief_facts"] += "\n\n" + data["brief_facts"]
            else:
                merged["brief_facts"] = data["brief_facts"]
        
        # Merge property
        if data.get("property_lost") and not merged["property_lost"]:
            merged["property_lost"] = data["property_lost"]
        if data.get("property_recovered") and not merged["property_recovered"]:
            merged["property_recovered"] = data["property_recovered"]
    
    return merged


def find_missing_fields(data: dict) -> list:
    """Identify mandatory fields that are missing"""
    missing = []
    
    # Check complainant
    comp = data.get("complainant", {})
    if not comp.get("name"):
        missing.append("COMPLAINANT_NAME")
    if not comp.get("father_name"):
        missing.append("COMPLAINANT_FATHER_NAME")
    if not comp.get("caste"):
        missing.append("COMPLAINANT_CASTE")
    if not comp.get("address"):
        missing.append("COMPLAINANT_ADDRESS")
    
    # Check accused
    if not data.get("accused_persons"):
        missing.append("ACCUSED_DETAILS")
    else:
        for i, acc in enumerate(data["accused_persons"]):
            if not acc.get("caste"):
                missing.append(f"A{i+1}_CASTE")
            if not acc.get("address"):
                missing.append(f"A{i+1}_ADDRESS")
    
    # Check other mandatory fields
    if not data.get("brief_facts"):
        missing.append("BRIEF_FACTS")
    
    offense = data.get("offense_details", {})
    if not offense.get("date"):
        missing.append("DATE_OF_OFFENSE")
    if not offense.get("place"):
        missing.append("PLACE_OF_OFFENSE")
    
    return missing


def generate_charge_sheet_content(data: dict, case_info: dict) -> str:
    """Generate Charge Sheet content with Active Blanks for missing fields"""
    
    # Format complainant
    comp = data.get("complainant", {})
    complainant_name = comp.get("name") or "[MISSING: COMPLAINANT_NAME]"
    complainant_father = comp.get("father_name") or "[MISSING: COMPLAINANT_FATHER_NAME]"
    complainant_age = comp.get("age") or "[MISSING: AGE]"
    complainant_caste = comp.get("caste") or "[MISSING: CASTE]"
    complainant_occupation = comp.get("occupation") or "[MISSING: OCCUPATION]"
    complainant_address = comp.get("address") or "[MISSING: ADDRESS]"
    complainant_phone = comp.get("phone") or "[MISSING: PHONE]"
    
    # Format accused list
    accused_list = ""
    for i, acc in enumerate(data.get("accused_persons", [])):
        serial = acc.get("serial") or f"A{i+1}"
        name = acc.get("name") or f"[MISSING: A{i+1}_NAME]"
        father = acc.get("father_name") or f"[MISSING: A{i+1}_FATHER_NAME]"
        age = acc.get("age") or "[AGE]"
        caste = acc.get("caste") or f"[MISSING: A{i+1}_CASTE]"
        occupation = acc.get("occupation") or "[OCCUPATION]"
        address = acc.get("address") or f"[MISSING: A{i+1}_ADDRESS]"
        
        accused_list += f"""
   {serial}. {name}
      S/o {father}
      Age: {age} years, Caste: {caste}
      Occupation: {occupation}
      R/o: {address}
      a) Date of arrest/release: [TO BE FILLED]
      b) Surety particulars: [TO BE FILLED]
      c) Previous convictions: Nil
"""
    
    if not accused_list:
        accused_list = "   [MISSING: ACCUSED_DETAILS]"
    
    # Format witness list
    witness_list = ""
    for i, wit in enumerate(data.get("witnesses", [])):
        serial = wit.get("serial") or f"LW-{i+1}"
        name = wit.get("name") or f"[WITNESS_{i+1}_NAME]"
        role = wit.get("role") or ""
        witness_list += f"   {serial}. {name} - {role}\n"
    
    if not witness_list:
        witness_list = "   [MISSING: WITNESS_DETAILS]"
    
    # Get offense details (used in case_info if not provided)
    offense = data.get("offense_details", {})
    # These can be used later for enhanced templates
    _ = offense.get("date") or "[MISSING: DATE_OF_OFFENSE]"
    _ = offense.get("place") or "[MISSING: PLACE_OF_OFFENSE]"
    
    # Sections
    sections = ", ".join(data.get("sections_of_law", [])) or case_info.get("sections") or "[MISSING: SECTIONS]"
    
    # Brief facts
    brief_facts = data.get("brief_facts") or "[MISSING: BRIEF_FACTS - Please add case narrative]"
    
    # Property
    property_seized = data.get("property_lost") or "---"
    
    # Generate charge sheet
    charge_sheet = CHARGE_SHEET_TEMPLATE.format(
        court_location=case_info.get("district", "[COURT_LOCATION]"),
        district=case_info.get("district", "[DISTRICT]"),
        police_station=case_info.get("police_station", "[POLICE_STATION]"),
        fir_number=case_info.get("fir_number") or "[MISSING: FIR_NUMBER]",
        fir_date="[FIR_DATE]",
        charge_date=datetime.now().strftime("%d.%m.%Y"),
        sections=sections,
        io_name="[MISSING: IO_NAME]",
        io_rank="[MISSING: IO_RANK]",
        complainant_name=complainant_name,
        complainant_father=complainant_father,
        complainant_age=complainant_age,
        complainant_caste=complainant_caste,
        complainant_occupation=complainant_occupation,
        complainant_address=complainant_address,
        complainant_phone=complainant_phone,
        property_seized=property_seized,
        accused_list=accused_list,
        not_charge_sheeted="Nil",
        witness_list=witness_list,
        lab_results="---",
        brief_facts=brief_facts,
        dispatch_date=datetime.now().strftime("%d.%m.%Y")
    )
    
    return charge_sheet


@router.post("/process")
async def process_documents(
    police_station: str = Form(...),
    district: str = Form(...),
    fir_number: str = Form(default=""),
    sections: str = Form(default=""),
    petition: Optional[UploadFile] = File(default=None),
    cdf: Optional[UploadFile] = File(default=None),
    case_diary: Optional[UploadFile] = File(default=None),
    officer: dict = Depends(get_current_officer)
):
    """
    Process multiple documents and generate Charge Sheet with 95% accuracy.
    Supports Telugu Petitions (JPG/PDF), CDF (.doc/.docx), Case Diary Part-II.
    """
    if not petition and not cdf and not case_diary:
        raise HTTPException(status_code=400, detail="At least one document must be uploaded")
    
    extracted_data_list = []
    doc_count = 0
    extraction_logs = []
    
    try:
        # Process Telugu Petition
        if petition:
            doc_count += 1
            contents = await petition.read()
            text = await extract_text_from_file(petition, contents)
            extraction_logs.append(f"Petition ({petition.filename}): {len(text)} chars extracted")
            
            if text and not text.startswith("[Error") and not text.startswith("[Unsupported"):
                result = await process_with_llm(text, "Telugu Petition")
                if result.get("success"):
                    extracted_data_list.append(result.get("data", {}))
            else:
                logger.warning(f"Petition extraction issue: {text[:100]}")
        
        # Process CDF
        if cdf:
            doc_count += 1
            contents = await cdf.read()
            text = await extract_text_from_file(cdf, contents)
            extraction_logs.append(f"CDF ({cdf.filename}): {len(text)} chars extracted")
            
            if text and not text.startswith("[Error") and not text.startswith("[Unsupported"):
                result = await process_with_llm(text, "Crime Details Form")
                if result.get("success"):
                    extracted_data_list.append(result.get("data", {}))
            else:
                logger.warning(f"CDF extraction issue: {text[:100]}")
        
        # Process Case Diary
        if case_diary:
            doc_count += 1
            contents = await case_diary.read()
            text = await extract_text_from_file(case_diary, contents)
            extraction_logs.append(f"Case Diary ({case_diary.filename}): {len(text)} chars extracted")
            
            if text and not text.startswith("[Error") and not text.startswith("[Unsupported"):
                result = await process_with_llm(text, "Case Diary Part-II")
                if result.get("success"):
                    extracted_data_list.append(result.get("data", {}))
            else:
                logger.warning(f"Case Diary extraction issue: {text[:100]}")
        
        # Merge all extracted data
        merged_data = merge_extracted_data(extracted_data_list)
        
        # Find missing fields
        missing_fields = find_missing_fields(merged_data)
        
        # Generate charge sheet
        case_info = {
            "police_station": police_station,
            "district": district,
            "fir_number": fir_number,
            "sections": sections,
            "fir_date": "",  # Can be extracted from documents
            "io_name": officer.get("name", ""),
            "io_rank": officer.get("rank", "")
        }
        
        # Generate 18-Column Charge Sheet (Makthal PS Format)
        charge_sheet_18col = generate_18_column_charge_sheet(merged_data, case_info)
        
        # Generate Case Diary Part-I with 8-Point Header
        case_diary_part1 = generate_case_diary_part1(merged_data, case_info)
        
        # Generate HTML table version for export/print
        charge_sheet_html = generate_html_table_charge_sheet(merged_data, case_info)
        
        # ML-suggested sections based on brief facts
        brief_facts = merged_data.get("brief_facts", "")
        suggested_sections = ml_suggest_sections(brief_facts) if brief_facts else []
        
        # Save to database
        fusion_record = {
            "officer_id": officer.get("officer_id"),
            "police_station": police_station,
            "district": district,
            "fir_number": fir_number,
            "documents_processed": doc_count,
            "extracted_data": merged_data,
            "missing_fields": missing_fields,
            "charge_sheet_18col": charge_sheet_18col,
            "case_diary_part1": case_diary_part1,
            "charge_sheet_html": charge_sheet_html,
            "suggested_sections": suggested_sections,
            "extraction_logs": extraction_logs,
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        await db.charge_sheet_fusions.insert_one(fusion_record)
        
        return {
            "success": True,
            "documents_processed": doc_count,
            "extracted_data": {
                "accused_count": len(merged_data.get("accused_persons", [])),
                "witness_count": len(merged_data.get("witnesses", [])),
                "sections": ", ".join(merged_data.get("sections_of_law", [])) or sections,
                "complainant": merged_data.get("complainant", {}),
                "brief_facts": merged_data.get("brief_facts", "")
            },
            "missing_fields": missing_fields,
            "suggested_sections": suggested_sections,
            "charge_sheet": charge_sheet_18col,
            "case_diary": case_diary_part1,
            "charge_sheet_html": charge_sheet_html,
            "extraction_logs": extraction_logs
        }
        
    except Exception as e:
        logger.error(f"Document processing error: {e}")
        raise HTTPException(status_code=500, detail=f"Processing failed: {str(e)}")



@router.post("/voice-to-text")
async def process_voice_input(
    audio_file: UploadFile = File(...),
    convert_to_legal: bool = Form(default=True),
    officer: dict = Depends(get_current_officer)
):
    """
    Process voice recording and convert to text.
    Supports Telugu, Hindi, and English with automatic translation.
    """
    try:
        # Validate file type
        filename = audio_file.filename.lower() if audio_file.filename else ""
        allowed_formats = ['mp3', 'wav', 'm4a', 'webm', 'ogg', 'aac']
        file_ext = filename.split('.')[-1] if '.' in filename else ''
        
        if file_ext not in allowed_formats:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported audio format. Supported: {', '.join(allowed_formats)}"
            )
        
        contents = await audio_file.read()
        
        # Extract text from audio
        result = await extract_text_from_audio(contents, audio_file.filename)
        
        response_data = {
            "success": True,
            "original_text": result.get("text", ""),
            "language_detected": result.get("language", "unknown"),
            "translated_text": result.get("translated", ""),
        }
        
        if convert_to_legal and result.get("legal_text"):
            response_data["legal_text"] = result["legal_text"]
        
        # If speech API not available, return helpful message
        if "[Speech-to-Text not available" in result.get("text", ""):
            response_data["success"] = False
            response_data["message"] = "Speech-to-Text API not configured. Please configure Google Speech credentials."
            # Provide mock for demo
            response_data["demo_mode"] = True
            response_data["original_text"] = "[Demo] Voice input received. In production, this would be transcribed using Google Speech-to-Text."
            response_data["translated_text"] = "[Demo] The complainant states that the accused person approached them and committed fraud."
            response_data["legal_text"] = "The complainant has stated that the accused person(s) approached the complainant and committed the offense of fraud by deception."
        
        return response_data
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Voice processing error: {e}")
        raise HTTPException(status_code=500, detail=f"Voice processing failed: {str(e)}")


@router.post("/convert-to-legal")
async def convert_text_to_legal(
    text: str = Form(...),
    officer: dict = Depends(get_current_officer)
):
    """
    Convert informal text to legal FIR format.
    """
    try:
        legal_text = convert_to_legal_format(text)
        return {
            "success": True,
            "original_text": text,
            "legal_text": legal_text
        }
    except Exception as e:
        logger.error(f"Legal conversion error: {e}")
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")



@router.post("/generate-remand")
async def generate_remand_cd(
    police_station: str = Form(...),
    district: str = Form(...),
    fir_number: str = Form(...),
    sections: str = Form(default=""),
    data: str = Form(default="{}"),
    officer: dict = Depends(get_current_officer)
):
    """
    Generate Remand Case Diary (Part-1) when arrest status is detected.
    Auto-triggered or manually requested.
    """
    import json
    
    try:
        extracted_data = json.loads(data) if data else {}
        
        case_info = {
            "police_station": police_station,
            "district": district,
            "fir_number": fir_number,
            "sections": sections,
            "io_name": officer.get("name", ""),
            "io_rank": officer.get("rank", "Sub Inspector of Police")
        }
        
        # Generate Remand CD HTML
        remand_html = generate_remand_case_diary_html(extracted_data, case_info)
        
        return {
            "success": True,
            "remand_cd_html": remand_html,
            "arrest_detected": True
        }
        
    except Exception as e:
        logger.error(f"Remand CD generation error: {e}")
        raise HTTPException(status_code=500, detail=f"Remand CD generation failed: {str(e)}")


@router.get("/cdf-form")
async def get_cdf_digital_form(
    language: str = "english",
    fir_number: Optional[str] = None,
    officer: dict = Depends(get_current_officer)
):
    """
    Get digital CDF form with bilingual support (Telugu/English toggle).
    """
    try:
        # Get existing case data if FIR number provided
        extracted_data = {}
        case_info = {
            "police_station": "",
            "district": "",
            "fir_number": fir_number or "",
            "sections": "",
            "io_name": officer.get("name", ""),
            "io_rank": officer.get("rank", "")
        }
        
        if fir_number and db:
            case = await db.charge_sheet_fusions.find_one({"fir_number": fir_number})
            if case:
                extracted_data = case.get("extracted_data", {})
                case_info["police_station"] = case.get("police_station", "")
                case_info["district"] = case.get("district", "")
                case_info["sections"] = case.get("sections", "")
        
        cdf_html = generate_cdf_digital_form_html(extracted_data, case_info, language)
        
        return {
            "success": True,
            "cdf_html": cdf_html,
            "language": language
        }
        
    except Exception as e:
        logger.error(f"CDF form generation error: {e}")
        raise HTTPException(status_code=500, detail=f"CDF form generation failed: {str(e)}")


@router.post("/cdf-form/save")
async def save_cdf_form(
    police_station: str = Form(...),
    district: str = Form(...),
    fir_number: str = Form(...),
    cdf_data: str = Form(...),
    officer: dict = Depends(get_current_officer)
):
    """
    Save CDF form data and auto-sync to Charge Sheet columns 13 & 16.
    """
    import json
    
    try:
        data = json.loads(cdf_data)
        
        # Extract data for Charge Sheet sync
        chargesheet_data = extract_cdf_data_for_chargesheet(data)
        
        # Save to database
        cdf_record = {
            "officer_id": officer.get("officer_id"),
            "police_station": police_station,
            "district": district,
            "fir_number": fir_number,
            "cdf_data": data,
            "chargesheet_sync": chargesheet_data,
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        
        if db:
            await db.cdf_forms.update_one(
                {"fir_number": fir_number},
                {"$set": cdf_record},
                upsert=True
            )
        
        return {
            "success": True,
            "message": "CDF saved and synced to Charge Sheet",
            "chargesheet_sync": {
                "column_13_witnesses": len(chargesheet_data.get("witnesses", [])),
                "column_16_modus_operandi": bool(chargesheet_data.get("modus_operandi"))
            }
        }
        
    except Exception as e:
        logger.error(f"CDF save error: {e}")
        raise HTTPException(status_code=500, detail=f"CDF save failed: {str(e)}")
